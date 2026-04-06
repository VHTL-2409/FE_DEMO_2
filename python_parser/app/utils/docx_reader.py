"""
docx_reader.py — Extract text from DOCX files preserving paragraph structure.
"""

from __future__ import annotations

from dataclasses import dataclass
from typing import Optional

try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

from .text_normalizer import normalize as normalize_extracted_text


@dataclass
class DocxParagraph:
    """A single paragraph with its index and optional style name."""
    index: int
    text: str
    style: Optional[str] = None


class DocxReader:
    """Read DOCX files and extract structured text."""

    @staticmethod
    def is_available() -> bool:
        """Check whether python-docx is installed."""
        return HAS_DOCX

    def read(self, docx_path: str) -> str:
        """
        Extract full text from DOCX.
        Paragraphs are joined by double newlines to preserve section separation.
        """
        if not HAS_DOCX:
            raise ImportError(
                "python-docx is not installed. Run: pip install python-docx>=1.1.0"
            )

        doc = docx.Document(docx_path)
        paragraphs = []
        for para in doc.paragraphs:
            text = normalize_extracted_text(para.text.strip())
            if text:
                paragraphs.append(text)
        return "\n\n".join(paragraphs)

    def read_with_styles(self, docx_path: str) -> list[DocxParagraph]:
        """
        Extract paragraphs with their paragraph style names.
        Useful for distinguishing headings, body text, etc.
        """
        if not HAS_DOCX:
            raise ImportError(
                "python-docx is not installed. Run: pip install python-docx>=1.1.0"
            )

        result = []
        doc = docx.Document(docx_path)
        for idx, para in enumerate(doc.paragraphs):
            text = normalize_extracted_text(para.text.strip())
            if text:
                style_name = para.style.name if para.style else None
                result.append(DocxParagraph(
                    index=idx,
                    text=text,
                    style=style_name,
                ))
        return result

    def close(self) -> None:
        """No-op for API compatibility with BaseParser.close()."""
        pass
