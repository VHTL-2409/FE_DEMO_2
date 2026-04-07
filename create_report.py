# -*- coding: utf-8 -*-
"""
Tạo BaoCaoNCKH_MAU.docx mới dựa trên dự án FE_DEMO
"""
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_margins(doc, top=2.0, bottom=2.0, left=2.5, right=2.5):
    for section in doc.sections:
        section.top_margin = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin = Cm(left)
        section.right_margin = Cm(right)

def add_heading(doc, text, level=1, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    if level == 1:
        run.font.size = Pt(16)
    elif level == 2:
        run.font.size = Pt(14)
    elif level == 3:
        run.font.size = Pt(13)
    else:
        run.font.size = Pt(12)
    return p

def add_para(doc, text, bold=False, indent=0, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(13)
    return p

def add_bullet(doc, text, indent=1.0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.font.size = Pt(13)
    return p

def create_report():
    doc = Document()
    set_margins(doc)
    
    # ========== MỤC LỤC ==========
    add_heading(doc, "MỤC LỤC", level=1)
    
    toc_items = [
        "MỞ ĐẦU",
        "  1. Bối cảnh và tính cấp thiết của đề tài",
        "  2. Lý do chọn đề tài",
        "  3. Mục đích và ý nghĩa của đề tài",
        "  4. Đối tượng và phạm vi nghiên cứu",
        "CHƯƠNG I: TỔNG QUAN VỀ CÔNG NGHỆ VÀ CÔNG CỤ",
        "  1.1. Giới thiệu Vue.js 3",
        "  1.2. Giới thiệu Vite",
        "  1.3. Giới thiệu Spring Boot",
        "  1.4. Giới thiệu PostgreSQL",
        "  1.5. Giới thiệu FastAPI và AI Service",
        "  1.6. Giới thiệu Docker",
        "CHƯƠNG II: PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG",
        "  2.1. Khảo sát thực trạng",
        "  2.2. Một số hệ thống thi trực tuyến hiện có",
        "  2.3. Các đối tượng tương tác với hệ thống",
        "  2.4. Phân tích các chức năng của hệ thống",
        "CHƯƠNG III: THIẾT KẾ HỆ THỐNG",
        "  3.1. Thiết kế cơ sở dữ liệu",
        "  3.2. Thiết kế API Backend",
        "  3.3. Thiết kế giao diện Frontend",
        "CHƯƠNG IV: VẬN HÀNH VÀ KẾT QUẢ",
        "  4.1. Triển khai Docker",
        "  4.2. Kết quả thực hiện",
        "  4.3. Kiểm thử hệ thống",
        "KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN",
        "TÀI LIỆU THAM KHẢO",
    ]
    
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            run.font.size = Pt(13)
    
    doc.add_page_break()
    
    # ========== MỞ ĐẦU ==========
    add_heading(doc, "MỞ ĐẦU", level=1)
    
    add_heading(doc, "1. Bối cảnh và tính cấp thiết của đề tài", level=2, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para(doc, """Trong những năm gần đây, giáo dục trực tuyến (e-learning) đã phát triển mạnh mẽ trên toàn thế giới, đặc biệt sau đại dịch COVID-19. Hình thức thi trực tuyến trở thành nhu cầu thiết yếu của các trường học, tổ chức giáo dục và doanh nghiệp. Tuy nhiên, các giải pháp thi trực tuyến hiện tại vẫn còn nhiều hạn chế về tính năng giám sát, phát hiện gian lận, và khả năng hỗ trợ AI.""")
    
    add_para(doc, """Xuất phát từ thực tiễn đó, đề tài "Xây dựng hệ thống thi trực tuyến có giám sát với hỗ trợ AI" được nghiên cứu nhằm tạo ra một giải pháp toàn diện, đáp ứng nhu cầu thi trực tuyến hiện đại.""")
    
    add_heading(doc, "2. Lý do chọn đề tài", level=2, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para(doc, """- Nhu cầu thi trực tuyến ngày càng tăng trong giáo dục và doanh nghiệp.""")
    add_para(doc, """- Thiếu các giải pháp thi trực tuyến tích hợp AI để hỗ trợ giám sát và chấm điểm tự động.""")
    add_para(doc, """- Mong muốn xây dựng hệ thống mã nguồn mở, dễ triển khai và mở rộng.""")
    add_para(doc, """- Xu hướng chuyển đổi số trong giáo dục đòi hỏi công cụ thi trực tuyến hiệu quả.""")
    
    add_heading(doc, "3. Mục đích và ý nghĩa của đề tài", level=2, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para(doc, """Mục đích: Xây dựng hệ thống thi trực tuyến với các tính năng quản lý đề thi, tổ chức thi có giám sát, phát hiện gian lận bằng AI, và hỗ trợ AI trong chấm điểm essay.""")
    add_para(doc, """Ý nghĩa: Đóng góp vào việc ứng dụng công nghệ AI vào giáo dục, nâng cao chất lượng và tính công bằng trong kiểm tra đánh giá.""")
    
    add_heading(doc, "4. Đối tượng và phạm vi nghiên cứu", level=2, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para(doc, """Đối tượng: Giáo viên, học sinh, quản trị viên trong môi trường giáo dục.""")
    add_para(doc, """Phạm vi: Hệ thống thi trực tuyến với frontend Vue 3, backend Spring Boot, AI service FastAPI, triển khai bằng Docker.""")
    
    doc.add_page_break()
    
    # ========== CHƯƠNG I ==========
    add_heading(doc, "CHƯƠNG I: TỔNG QUAN VỀ CÔNG NGHỆ VÀ CÔNG CỤ", level=1)
    
    add_heading(doc, "1.1. Giới thiệu Vue.js 3", level=2, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para(doc, """Vue.js là một framework JavaScript tiến bộ để xây dựng giao diện người dùng. Phiên bản Vue 3 mang đến nhiều cải tiến về hiệu suất, Composition API, và TypeScript support. Trong dự án này, Vue 3 được sử dụng kết hợp với Pinia (state management) và Vue Router (điều hướng).""")
    
    add_heading(doc, "1.2. Giới thiệu Vite", level=2, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para(doc, """Vite là một công cụ build tool thế hệ mới cho frontend, cung cấp tốc độ phát triển nhanh nhờ HMR (Hot Module Replacement) và ESM native. Dự án sử dụng Vite với plugin cho Vue 3.""")
    
    add_heading(doc, "1.3. Giới thiệu Spring Boot", level=2, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para(doc, """Spring Boot là framework Java mạnh mẽ giúp xây dựng ứng dụng backend nhanh chóng. Với Java 17, Spring Boot cung cấp RESTful API, JPA/Hibernate cho ORM, Spring Security cho authentication/authorization, và WebSocket cho thông báo real-time.""")
    
    add_heading(doc, "1.4. Giới thiệu PostgreSQL", level=2, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para(doc, """PostgreSQL là hệ quản trị cơ sở dữ liệu quan hệ mã nguồn mở, nổi tiếng về độ tin cậy, tính toàn vẹn dữ liệu và hỗ trợ nhiều tính năng nâng cao. Database "datn" được sử dụng trong dự án với các bảng: users, exams, questions, exam_attempts, submissions, classes, assignments, monitoring_events, fraud_signals.""")
    
    add_heading(doc, "1.5. Giới thiệu FastAPI và AI Service", level=2, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para(doc, """AI Service được xây dựng bằng FastAPI (Python) với các chức năng: OCR nhận diện văn bản từ hình ảnh, phân tích hành vi học sinh khi thi, chấm điểm essay tự động bằng AI, tạo câu hỏi ôn tập tự động, và dự đoán hiệu suất học tập.""")
    
    add_heading(doc, "1.6. Giới thiệu Docker", level=2, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para(doc, """Docker và Docker Compose được sử dụng để đóng gói và triển khai toàn bộ hệ thống, bao gồm: Backend (Spring Boot), Frontend (Vue 3 qua Nginx), AI Service (FastAPI), và PostgreSQL database.""")
    
    doc.add_page_break()
    
    # ========== CHƯƠNG II ==========
    add_heading(doc, "CHƯƠNG II: PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG", level=1)
    
    add_heading(doc, "2.1. Khảo sát thực trạng", level=2, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para(doc, """Khảo sát cho thấy nhu cầu về hệ thống thi trực tuyến ngày càng cao. Các trường học và tổ chức giáo dục cần giải pháp có khả năng: quản lý đề thi đa dạng định dạng, giám sát thi trực tuyến hiệu quả, phát hiện hành vi gian lận, và chấm điểm tự động.""")
    
    add_heading(doc, "2.2. Một số hệ thống thi trực tuyến hiện có", level=2, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para(doc, """- Azota.vn: Nền tảng thi trực tuyến phổ biến tại Việt Nam, có API để import đề thi.""")
    add_para(doc, """- Google Forms, Microsoft Forms: Công cụ thi cơ bản, thiếu tính năng giám sát.""")
    add_para(doc, """- Moodle: Hệ thống LMS mã nguồn mở, phức tạp và khó triển khai.""")
    
    add_heading(doc, "2.3. Các đối tượng tương tác với hệ thống", level=2, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para(doc, """- Admin: Quản lý người dùng, quản lý hệ thống, xem báo cáo thống kê.""")
    add_para(doc, """- Teacher: Tạo đề thi, quản lý lớp học, xem kết quả thi, import đề từ Azota.""")
    add_para(doc, """- Student: Đăng ký, đăng nhập, tham gia thi, xem kết quả.""")
    
    add_heading(doc, "2.4. Phân tích các chức năng của hệ thống", level=2, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    
    add_para(doc, "2.4.1. Các yêu cầu chức năng", bold=True)
    add_para(doc, """a) Quản lý tài khoản: Đăng ký, đăng nhập, quản lý hồ sơ, phân quyền Admin/Teacher/Student.""")
    add_para(doc, """b) Quản lý lớp học: Tạo lớp, thêm học sinh, mã lớp để học sinh tham gia.""")
    add_para(doc, """c) Quản lý đề thi: Tạo đề thi, import từ CSV/Excel/DOCX/PDF, import từ Azota API.""")
    add_para(doc, """d) Tổ chức thi: Bắt đầu thi, nộp bài, tự động nộp khi hết giờ.""")
    add_para(doc, """e) Giám sát thi: Theo dõi hành vi, phát hiện gian lận, cảnh báo giáo viên real-time qua WebSocket.""")
    add_para(doc, """f) AI hỗ trợ: Chấm essay, tạo câu hỏi, phân tích hành vi, dự đoán hiệu suất.""")
    
    add_para(doc, "2.4.2. Sơ đồ Use Case tổng quát", bold=True)
    add_para(doc, """Use Case Đăng nhập: Actor (Student/Teacher/Admin), hệ thống xác thực JWT token.""")
    add_para(doc, """Use Case Tạo đề thi: Actor (Teacher), chọn loại câu hỏi, nhập nội dung, lưu vào database.""")
    add_para(doc, """Use Case Thi trực tuyến: Actor (Student), chọn đề, làm bài, nộp bài, xem kết quả.""")
    add_para(doc, """Use Case Giám sát: Actor (Teacher/Admin), xem thời gian thực, nhận cảnh báo.""")
    
    doc.add_page_break()
    
    # ========== CHƯƠNG III ==========
    add_heading(doc, "CHƯƠNG III: THIẾT KẾ HỆ THỐNG", level=1)
    
    add_heading(doc, "3.1. Thiết kế cơ sở dữ liệu", level=2, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    
    add_para(doc, "3.1.1. Mô hình quan hệ (ERD)", bold=True)
    add_para(doc, """Các thực thể chính: User (id, username, email, password, role), Exam (id, title, duration, start_time, created_by), Question (id, exam_id, content, type, options, correct_answer), ExamAttempt (id, exam_id, student_id, start_time, end_time, status), Submission (id, attempt_id, question_id, answer), Class (id, name, code, teacher_id), MonitoringEvent (id, attempt_id, event_type, timestamp, details).""")
    
    add_para(doc, "3.1.2. Mô hình vật lý", bold=True)
    add_para(doc, """Bảng users: id (PK), username, email, password_hash, role, created_at, updated_at.""")
    add_para(doc, """Bảng exams: id (PK), title, description, duration_minutes, start_time, end_time, teacher_id (FK), status.""")
    add_para(doc, """Bảng questions: id (PK), exam_id (FK), content, question_type, options (JSON), correct_answer, points.""")
    add_para(doc, """Bảng exam_attempts: id (PK), exam_id (FK), student_id (FK), start_time, end_time, status, ip_address.""")
    add_para(doc, """Bảng submissions: id (PK), attempt_id (FK), question_id (FK), answer, is_correct, created_at.""")
    add_para(doc, """Bảng classes: id (PK), name, code, teacher_id (FK), created_at.""")
    add_para(doc, """Bảng class_students: class_id (FK), student_id (FK), joined_at.""")
    add_para(doc, """Bảng monitoring_events: id (PK), attempt_id (FK), event_type, severity, details (JSON), created_at.""")
    add_para(doc, """Bảng fraud_signals: id (PK), attempt_id (FK), signal_type, risk_score, detected_at.""")
    
    add_heading(doc, "3.2. Thiết kế API Backend", level=2, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para(doc, """Backend Spring Boot cung cấp RESTful API trên cổng 8082:""")
    add_para(doc, """- Auth: POST /api/auth/login, /register, /refresh""")
    add_para(doc, """- Users: GET/PUT /api/profile, PUT /password""")
    add_para(doc, """- Classes: CRUD /api/classes, POST /join, POST /students/bulk""")
    add_para(doc, """- Exams: CRUD /api/exams, POST /import (CSV/Excel/DOCX/Azota)""")
    add_para(doc, """- Questions: CRUD /api/questions, /bulk""")
    add_para(doc, """- Submissions: POST /start, /submit, GET /attempts""")
    add_para(doc, """- Monitoring: WebSocket /ws/alerts, POST /events, GET /risk-score""")
    add_para(doc, """- AI: POST /api/ai/evaluate-essay, /generate-questions, /analyze-behavior, /predict-performance""")
    
    add_heading(doc, "3.3. Thiết kế giao diện Frontend", level=2, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para(doc, """Frontend Vue 3 với Router:""")
    add_para(doc, """- /login, /register: Trang đăng nhập/đăng ký""")
    add_para(doc, """- /teacher/dashboard: Trang chủ giáo viên với thống kê""")
    add_para(doc, """- /teacher/exams: Quản lý đề thi, tạo mới, import""")
    add_para(doc, """- /teacher/classes: Quản lý lớp học""")
    add_para(doc, """- /teacher/monitoring: Giám sát thi real-time""")
    add_para(doc, """- /student/exams: Danh sách kỳ thi""")
    add_para(doc, """- /student/exam/:id: Trang thi với timer, câu hỏi""")
    add_para(doc, """- /student/results: Xem kết quả các kỳ thi""")
    add_para(doc, """- /admin/dashboard: Thống kê hệ thống""")
    add_para(doc, """- /admin/users: Quản lý người dùng""")
    
    doc.add_page_break()
    
    # ========== CHƯƠNG IV ==========
    add_heading(doc, "CHƯƠNG IV: VẬN HÀNH VÀ KẾT QUẢ", level=1)
    
    add_heading(doc, "4.1. Triển khai Docker", level=2, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para(doc, """Hệ thống được triển khai bằng Docker Compose với các service:""")
    add_para(doc, """- app (Backend Spring Boot): Cổng 8082""")
    add_para(doc, """- frontend (Nginx + Vue build): Cổng 8080""")
    add_para(doc, """- ai-service (FastAPI): Cổng 8090""")
    add_para(doc, """- db (PostgreSQL): Cổng 5432""")
    add_para(doc, """- nginx (Reverse Proxy): Cổng 80, 443""")
    
    add_heading(doc, "4.2. Kết quả thực hiện", level=2, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para(doc, "4.2.1. Giao diện đăng nhập/đăng ký", bold=True)
    add_para(doc, """Thiết kế theo phong cách hiện đại, responsive, sử dụng design system với CSS variables. Form validation, thông báo lỗi, loading states.""")
    
    add_para(doc, "4.2.2. Trang Dashboard", bold=True)
    add_para(doc, """Hiển thị thống kê tổng quan: số lượng kỳ thi, học sinh, tỷ lệ hoàn thành, biểu đồ ECharts.""")
    
    add_para(doc, "4.2.3. Quản lý đề thi", bold=True)
    add_para(doc, """Hỗ trợ tạo đề thi với nhiều loại câu hỏi: trắc nghiệm, điền từ, essay. Import đề từ CSV, Excel, DOCX, PDF, hoặc từ Azota API.""")
    
    add_para(doc, "4.2.4. Trang thi trực tuyến", bold=True)
    add_para(doc, """Giao diện thi với timer đếm ngược, hiển thị câu hỏi, lưu tạm câu trả lời, nộp bài tự động khi hết giờ. Hiển thị số câu đã trả lời.""")
    
    add_para(doc, "4.2.5. Giám sát real-time", bold=True)
    add_para(doc, """WebSocket kết nối giữa frontend và backend để gửi cảnh báo. Hiển thị danh sách học sinh đang thi, trạng thái, điểm rủi ro. Cảnh báo khi phát hiện hành vi bất thường.""")
    
    add_heading(doc, "4.3. Kiểm thử hệ thống", level=2, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para(doc, """- Unit tests cho backend (JUnit, Mockito)""")
    add_para(doc, """- Integration tests cho API endpoints""")
    add_para(doc, """- Kiểm thử chức năng: đăng nhập, tạo đề, thi, nộp bài, giám sát""")
    add_para(doc, """- Kiểm thử hiệu năng: đồng thời nhiều học sinh thi""")
    add_para(doc, """- Kiểm thử bảo mật: JWT authentication, CORS, SQL injection""")
    
    doc.add_page_break()
    
    # ========== KẾT LUẬN ==========
    add_heading(doc, "KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", level=1)
    
    add_para(doc, "Kết luận:", bold=True)
    add_para(doc, """Đề tài đã hoàn thành việc xây dựng hệ thống thi trực tuyến với các tính năng chính: quản lý tài khoản đa vai trò, tạo và import đề thi đa dạng định dạng, tổ chức thi trực tuyến với giám sát real-time, phát hiện gian lận bằng AI, và chấm điểm essay tự động.""")
    
    add_para(doc, "Hướng phát triển:", bold=True)
    add_para(doc, """- Tích hợp thêm các phương thức thi (thi tự luận online, thi vấn đáp qua video call).""")
    add_para(doc, """- Cải thiện thuật toán phát hiện gian lận bằng Deep Learning.""")
    add_para(doc, """- Phát triển ứng dụng di động (React Native).""")
    add_para(doc, """- Tích hợp hệ thống LMS bên thứ ba.""")
    add_para(doc, """- Mở rộng AI service với nhiều mô hình ngôn ngữ lớn hơn.""")
    
    doc.add_page_break()
    
    # ========== TÀI LIỆU THAM KHẢO ==========
    add_heading(doc, "TÀI LIỆU THAM KHẢO", level=1)
    
    refs = [
        "1. Vue.js Documentation. https://vuejs.org/",
        "2. Spring Boot Reference Documentation. https://spring.io/projects/spring-boot",
        "3. PostgreSQL Documentation. https://www.postgresql.org/docs/",
        "4. FastAPI Documentation. https://fastapi.tiangolo.com/",
        "5. Docker Documentation. https://docs.docker.com/",
        "6. Vite Documentation. https://vitejs.dev/",
        "7. Pinia Documentation. https://pinia.vuejs.org/",
        "8. JWT Official Website. https://jwt.io/",
        "9. Azota API Documentation. https://azota.vn/",
        "10. OpenAI API Documentation. https://platform.openai.com/",
    ]
    
    for ref in refs:
        add_para(doc, ref)
    
    # Lưu file
    output_path = r"C:\Users\Administrator\Desktop\FE_DEMO\BaoCaoNCKH_MAU_new.docx"
    doc.save(output_path)
    print(f"Đã tạo file: {output_path}")
    return output_path

if __name__ == "__main__":
    create_report()
