# -*- coding: utf-8 -*-
from docx import Document
import sys
import codecs

def extract_docx(filepath, output_path):
    doc = Document(filepath)
    
    with codecs.open(output_path, 'w', 'utf-8') as f:
        for para in doc.paragraphs:
            f.write(para.text + '\n')
        
        # Check for tables
        for i, table in enumerate(doc.tables):
            f.write(f"\n=== TABLE {i+1} ===\n")
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                f.write(" | ".join(row_text) + '\n')

if __name__ == "__main__":
    filepath = sys.argv[1]
    output_path = sys.argv[2]
    extract_docx(filepath, output_path)
    print(f"Extracted to {output_path}")
