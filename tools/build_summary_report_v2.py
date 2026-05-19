from __future__ import annotations

import re
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
BAOCao_DIR = ROOT / "baocao"
SOURCE_DOCX = BAOCao_DIR / "baocao_long_test_format_anh.docx"
TEMPLATE_DOCX = BAOCao_DIR / "4. Mau Bao cao tóm tắt.docx"
OUTPUT_DOCX = BAOCao_DIR / "Bao-cao-tom-tat-do-an-thi-truc-tuyen-giam-sat-gian-lan-thong-minh.docx"


def clear_body(document: Document) -> None:
    body = document._body._element
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def apply_styles(document: Document) -> None:
    for name, size, bold, italic in [
        ("Normal", 13, False, False),
        ("Heading 1", 16, True, False),
        ("Heading 2", 14, True, False),
        ("Heading 3", 13, True, False),
        ("Caption", 12, False, True),
    ]:
        if name not in document.styles:
            continue
        style = document.styles[name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.italic = italic


def set_run_font(run, *, size: int = 13, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic


def add_paragraph(
    document: Document,
    text: str = "",
    *,
    style: str = "Normal",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_cm: float | None = 0.75,
    size: int = 13,
    bold: bool = False,
    italic: bool = False,
) -> object:
    p = document.add_paragraph(style=style)
    p.alignment = align
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(0)
    if first_line_cm is not None:
        p.paragraph_format.first_line_indent = Cm(first_line_cm)
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def add_heading(document: Document, text: str, level: int, *, page_break: bool = False) -> None:
    if page_break:
        document.add_page_break()
    p = document.add_paragraph(style=f"Heading {level}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, size=16 if level == 1 else 14 if level == 2 else 13, bold=True)


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for p in header_cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run_font(run, size=11, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            for p in cells[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    set_run_font(run, size=11)
    document.add_paragraph()


def extract_figures(source: Path, *, predicate) -> list[dict[str, object]]:
    doc = Document(str(source))
    figures: list[dict[str, object]] = []
    paras = doc.paragraphs
    for i, para in enumerate(paras):
        blips = para._p.xpath(".//a:blip")
        if not blips:
            continue
        caption = ""
        for j in range(i + 1, min(i + 4, len(paras))):
            txt = paras[j].text.strip().replace("\t", " ")
            if txt:
                caption = txt
                break
        raw = " ".join(caption.split())
        if not raw or not predicate(raw):
            continue
        clean = re.sub(r"^Hình\s+\d+\.\d+\.\s*", "", raw, flags=re.I)
        clean = re.sub(r"(?i)usecase", "Use Case", clean)
        clean = re.sub(r"(?i)use case", "Use Case", clean)
        rid = blips[0].get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
        figures.append({"caption": clean, "blob": doc.part.related_parts[rid].blob})
    return figures


def add_figure(document: Document, chapter_no: int, figure_no: int, caption: str, blob: bytes, *, width_in: float) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(BytesIO(blob), width=Inches(width_in))

    cap = document.add_paragraph(style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(6)
    run = cap.add_run(f"Hình {chapter_no}.{figure_no}. {caption}")
    set_run_font(run, size=12, italic=True)


def build_report(document: Document, versions: dict[str, str]) -> None:
    apply_styles(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(80)
    title.add_run("BÁO CÁO TÓM TẮT ĐỒ ÁN").font.size = Pt(20)
    set_run_font(title.runs[0], size=20, bold=True)

    sub = document.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Hệ thống thi trực tuyến với cơ chế giám sát và phát hiện gian lận thông minh")
    set_run_font(run, size=16, bold=True)

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("FE_DEMO")
    set_run_font(run, size=13, bold=True)

    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run("Phiên bản tóm tắt, lấy ảnh từ báo cáo dài của đồ án")
    set_run_font(run, size=12, italic=True)

    document.add_page_break()

    add_heading(document, "MỞ ĐẦU", 1)
    add_heading(document, "1. Lý do chọn đề tài", 2)
    add_paragraph(document, "Thi trực tuyến càng phổ biến thì áp lực giám sát gian lận càng lớn. Nếu chỉ dừng ở phát đề và nộp bài, hệ thống thiếu cơ chế theo dõi hành vi bất thường, khó phát hiện đổi tab, mất focus, copy-paste hay thao tác devtools trong lúc làm bài.")
    add_paragraph(document, "FE_DEMO xử lý đúng khoảng trống đó bằng kiến trúc nhiều tầng: frontend Vue 3, backend Spring Boot, dịch vụ AI FastAPI và PostgreSQL. Cách làm này cho phép hệ thống vừa hỗ trợ thi trực tuyến, vừa giám sát realtime, vừa lưu vết và phân tích rủi ro.")

    add_heading(document, "2. Mục tiêu và phạm vi", 2)
    add_paragraph(document, "Mục tiêu là xây dựng luồng thi hoàn chỉnh cho ba vai trò chính: sinh viên, giáo viên và quản trị viên. Hệ thống cần đáp ứng đăng nhập, quản lý đề thi, tham gia thi, giám sát realtime, chấm điểm rủi ro, quản lý lớp học và dashboard quản trị.")
    add_paragraph(document, "Phạm vi tóm tắt tập trung vào phần phân tích chức năng, thiết kế hệ thống và kết quả triển khai đã có trong repo. Báo cáo không đi sâu mọi màn hình nhỏ, mà nhấn vào kiến trúc, use case chính và ảnh chụp tiêu biểu.")

    add_heading(document, "3. Phương pháp thực hiện", 2)
    add_paragraph(document, "Cách làm dựa trên phân tích mã nguồn frontend/backend/AI service và tài liệu UML trong repo. Phần minh họa lấy trực tiếp từ file báo cáo dài và screenshot của hệ thống để bám sát triển khai thật.")
    add_paragraph(document, "Bố cục báo cáo gồm tổng quan công nghệ, phân tích hệ thống, thiết kế và xây dựng, rồi tới vận hành và kiểm thử. Đây là bản rút gọn, dùng được như bản nộp tóm tắt của đồ án.")

    add_heading(document, "4. Cấu trúc báo cáo", 2)
    add_paragraph(document, "Chương I: Tổng quan công nghệ và môi trường phát triển.", first_line_cm=0.0)
    add_paragraph(document, "Chương II: Phân tích hệ thống và use case.", first_line_cm=0.0)
    add_paragraph(document, "Chương III: Thiết kế và xây dựng hệ thống.", first_line_cm=0.0)
    add_paragraph(document, "Chương IV: Vận hành và kiểm thử.", first_line_cm=0.0)

    add_heading(document, "CHƯƠNG I: TỔNG QUAN CÔNG NGHỆ VÀ MÔI TRƯỜNG PHÁT TRIỂN", 1, page_break=True)
    add_heading(document, "1.1. Tổng quan hệ thống", 2)
    add_paragraph(document, "FE_DEMO là hệ thống thi trực tuyến có giám sát thông minh, được thiết kế để hỗ trợ cả thao tác nghiệp vụ lẫn theo dõi hành vi bất thường của thí sinh. Ở mức tổng thể, hệ thống chia rõ giao diện người dùng, lớp nghiệp vụ, lớp AI và lớp dữ liệu để dễ mở rộng.")
    add_paragraph(document, "Điểm mạnh là luồng realtime qua WebSocket và lớp phân tích AI cho camera, OCR, chat và đánh giá rủi ro. Nhờ vậy giáo viên có thể theo dõi diễn biến phòng thi, còn hệ thống có đủ dữ liệu để lưu nhật ký và cảnh báo sau thi.")

    add_heading(document, "1.2. Frontend", 2)
    add_paragraph(document, f"Frontend nằm trong `demo/`, dùng Vue {versions['vue']} với Vite {versions['vite']}. Các thành phần như Pinia {versions['pinia']}, Vue Router {versions['vue_router']}, PrimeVue {versions['primevue']}, ECharts {versions['echarts']} và xlsx {versions['xlsx']} cho thấy giao diện giàu tương tác và có dashboard tốt.")
    add_table(document, ["Thành phần", "Phiên bản", "Vai trò"], [
        ["Vue", versions["vue"], "Framework giao diện"],
        ["Vite", versions["vite"], "Dev server và bundler"],
        ["Pinia", versions["pinia"], "Quản lý state"],
        ["Vue Router", versions["vue_router"], "Điều hướng"],
        ["PrimeVue", versions["primevue"], "Component UI"],
        ["ECharts", versions["echarts"], "Biểu đồ và dashboard"],
        ["xlsx", versions["xlsx"], "Import/export dữ liệu"],
    ])

    add_heading(document, "1.3. Backend", 2)
    add_paragraph(document, f"Backend nằm trong `BE/`, xây trên Spring Boot {versions['spring_boot']} và Java {versions['java']}. Lớp backend xử lý xác thực, phân quyền, đề thi, lớp học, phiên thi, giám sát và các API admin.")
    add_table(document, ["Thành phần", "Vai trò"], [
        ["Spring Security", "Xác thực và phân quyền"],
        ["Spring Web", "REST API"],
        ["Spring WebSocket", "Kết nối realtime"],
        ["Spring Data JPA", "Truy cập dữ liệu"],
        ["JWT", "Token đăng nhập và refresh"],
        ["Mail", "Xác thực email và quên mật khẩu"],
    ])

    add_heading(document, "1.4. AI Service", 2)
    add_paragraph(document, "AI service tách riêng thành microservice FastAPI để xử lý tác vụ nặng như OCR, phân tích hình ảnh camera, chấm tự luận, sinh câu hỏi và chat hỗ trợ. Cách tách riêng này giúp backend gọn hơn và AI có thể scale độc lập.")
    add_table(document, ["Thành phần", "Vai trò"], [
        ["FastAPI", "Khung API cho AI service"],
        ["uvicorn", "ASGI server"],
        ["OpenCV", "Xử lý ảnh và video"],
        ["pytesseract", "Nhận dạng ký tự"],
        ["Pillow", "Xử lý ảnh"],
        ["NumPy", "Tính toán số học"],
        ["openai", "Kết nối mô hình LLM"],
        ["slowapi", "Giới hạn tần suất"],
    ])

    add_heading(document, "1.5. Triển khai", 2)
    add_paragraph(document, "Toàn bộ hệ thống được container hóa, có thể chạy qua Docker Compose với các service tách biệt cho frontend, backend, AI service, parser và PostgreSQL. Cấu trúc này phù hợp với môi trường phát triển lẫn demo nội bộ.")
    add_paragraph(document, "Mô hình nhiều service giúp debug từng khối riêng, đồng thời giữ luồng dữ liệu rõ ràng: trình duyệt gửi yêu cầu đến backend, backend gọi AI khi cần, còn trạng thái và log được lưu trong cơ sở dữ liệu trung tâm.")

    add_heading(document, "CHƯƠNG II: PHÂN TÍCH HỆ THỐNG VÀ USE CASE", 1, page_break=True)
    add_heading(document, "2.1. Tác nhân hệ thống", 2)
    add_paragraph(document, "Hệ thống có ba vai trò chính là sinh viên, giáo viên và quản trị viên. Ngoài ra còn có backend, AI service và WebSocket server để bảo đảm truyền dữ liệu, giám sát và phân tích diễn ra liên tục.")
    add_table(document, ["Tác nhân", "Nhiệm vụ chính", "Phạm vi tương tác"], [
        ["Sinh viên", "Thi, xem kết quả, tham gia lớp", "Khu vực sinh viên"],
        ["Giáo viên", "Tạo đề, giám sát, xem thống kê", "Khu vực giáo viên"],
        ["Quản trị viên", "Quản lý user và dashboard", "Khu vực admin"],
        ["AI service", "OCR, phân tích camera, chấm điểm", "API nội bộ"],
        ["Backend", "Xử lý nghiệp vụ và lưu dữ liệu", "API trung tâm"],
    ])

    usecase_figures = extract_figures(
        SOURCE_DOCX,
        predicate=lambda raw: ("use case" in raw.lower() or "usecase" in raw.lower())
        and "hoạt động" not in raw.lower()
        and "tuần tự" not in raw.lower(),
    )
    if len(usecase_figures) < 12:
        raise RuntimeError(f"Thiếu use case figure: {len(usecase_figures)}")
    usecase_groups = [
        ("2.2. Use case tổng quát và xác thực", [0, 1, 2], "Nhóm use case này mô tả lớp truy cập nền tảng: từ tổng quan hệ thống tới đăng nhập và khôi phục mật khẩu."),
        ("2.3. Use case sinh viên", [3, 4, 5], "Sinh viên là nhóm người dùng chính trong luồng làm bài. Các sơ đồ dưới đây cho thấy cách vào phòng thi, quản lý lớp học và xem kết quả sau thi."),
        ("2.4. Use case giáo viên", [6, 7, 8, 9], "Giáo viên quản lý gần như toàn bộ vòng đời kỳ thi: dashboard, lớp học, tạo đề thi, giám sát và điều phối các kỳ thi đang mở."),
        ("2.5. Use case quản trị", [10, 11], "Khối quản trị tập trung vào dashboard hệ thống và quản lý tài khoản người dùng."),
    ]
    fig_no = 1
    for heading, indexes, intro in usecase_groups:
        add_heading(document, heading, 2)
        add_paragraph(document, intro)
        for idx in indexes:
            figure = usecase_figures[idx]
            add_figure(document, 2, fig_no, str(figure["caption"]), figure["blob"], width_in=6.2)
            fig_no += 1

    add_heading(document, "CHƯƠNG III: THIẾT KẾ VÀ XÂY DỰNG HỆ THỐNG", 1, page_break=True)
    design_figures = extract_figures(SOURCE_DOCX, predicate=lambda raw: raw.lower().startswith("hình 3."))
    if len(design_figures) < 7:
        raise RuntimeError(f"Thiếu design figure: {len(design_figures)}")

    add_heading(document, "3.1. Thiết kế kiến trúc hệ thống", 2)
    add_paragraph(document, "Hệ thống được thiết kế theo kiến trúc microservices với Frontend (Vue 3 + Vite), Backend (Spring Boot), AI Service (FastAPI) và Python Parser. PostgreSQL là database tập trung, còn Nginx đóng vai trò reverse proxy và load balancer.")
    add_paragraph(document, "Frontend giao tiếp với backend qua REST API và WebSocket/STOMP; backend gọi AI service qua REST API; toàn bộ service cùng nằm trong một Docker network để giao tiếp theo service name.")
    add_figure(document, 3, 1, str(design_figures[0]["caption"]), design_figures[0]["blob"], width_in=6.35)

    add_heading(document, "3.2. Thiết kế cơ sở dữ liệu", 2)
    add_paragraph(document, "CSDL dùng PostgreSQL theo mô hình quan hệ. Các bảng chính gồm users, roles, exams, questions, exam_attempts, answers, monitoring_events, fraud_signals, risk_snapshots và assignments.")
    add_table(document, ["Nhóm bảng", "Vai trò"], [
        ["Người dùng", "users, roles, refresh_tokens"],
        ["Thi", "exams, exam_attempts, exam_events"],
        ["Bài làm", "answers, submissions"],
        ["Giám sát", "monitoring_events, fraud_signals, risk_scores, audit_logs"],
        ["Lớp học", "classes, class_students"],
        ["Import", "import_jobs, import_job_issues"],
    ])
    add_figure(document, 3, 2, str(design_figures[1]["caption"]), design_figures[1]["blob"], width_in=6.25)

    add_heading(document, "3.3. Thiết kế API", 2)
    add_paragraph(document, "Backend cung cấp RESTful API theo nhóm chức năng, dùng JWT cho authentication và phân quyền theo vai trò. Hệ thống cũng hỗ trợ refresh token để gia hạn phiên đăng nhập mà không làm gián đoạn trải nghiệm.")
    add_table(document, ["Nhóm API", "Ví dụ endpoint", "Mục đích"], [
        ["Auth", "/api/auth/*, /api/users/me", "Đăng nhập và hồ sơ"],
        ["Exam", "/api/exams, /api/exams/{id}", "Quản lý đề thi"],
        ["Attempt", "/api/attempts/*", "Phiên làm bài"],
        ["Monitoring", "/api/v1/proctor/*", "Sự kiện và điểm rủi ro"],
        ["Admin", "/api/admin/*", "Dashboard và user"],
    ])

    add_heading(document, "3.4. Thiết kế giao diện người dùng", 2)
    add_paragraph(document, "Giao diện dùng PrimeVue 4, TailwindCSS và design tokens để giữ đồng nhất về màu sắc, typography, spacing và shadow. Frontend tối ưu cho desktop và tablet, đồng thời tách route rõ theo vai trò người dùng.")
    for idx, figure in enumerate(design_figures[2:7], start=3):
        add_figure(document, 3, idx, str(figure["caption"]), figure["blob"], width_in=5.95)

    add_heading(document, "3.5. Thiết kế module giám sát và AI", 2)
    add_paragraph(document, "Module giám sát chạy theo pipeline sự kiện: client gửi telemetry qua WebSocket/STOMP, backend lưu vào monitoring_events và gọi AI service để phân tích camera hoặc hành vi. Kết quả được quy đổi sang risk score 0-100 và đẩy lên dashboard giám sát.")
    add_table(document, ["Nhóm tín hiệu", "Trọng số / ghi chú"], [
        ["DEVTOOLS_OPEN", "0.9"],
        ["PRINT_SCREEN, MULTI_MONITOR", "0.85"],
        ["COPY_PASTE, EXIT_FULLSCREEN", "0.8 - 0.75"],
        ["DUPLICATE_IP, HEARTBEAT_STALE", "0.7"],
        ["TAB_SWITCH, WINDOW_BLUR", "0.5 - 0.4"],
        ["RIGHT_CLICK, IDLE_TIME", "0.35 - 0.3"],
    ])

    add_heading(document, "CHƯƠNG IV: VẬN HÀNH VÀ KIỂM THỬ", 1, page_break=True)
    result_figures = extract_figures(SOURCE_DOCX, predicate=lambda raw: raw.lower().startswith("hình 4."))
    if len(result_figures) < 6:
        raise RuntimeError(f"Thiếu result figure: {len(result_figures)}")

    add_heading(document, "4.1. Môi trường triển khai", 2)
    add_paragraph(document, "Hệ thống hỗ trợ chạy local bằng Docker Compose với PostgreSQL, backend, AI service, Python Parser và frontend preview. Ở chế độ production, Nginx đảm nhận reverse proxy và HTTPS.")
    add_table(document, ["Thành phần", "Cổng / vai trò"], [
        ["Frontend", "4173 / 5173"],
        ["Backend", "8082"],
        ["AI service", "8090"],
        ["Python Parser", "8000"],
        ["PostgreSQL", "5432"],
    ])

    add_heading(document, "4.2. Kết quả triển khai", 2)
    add_paragraph(document, "Các module chính đã chạy được thành chuỗi nghiệp vụ khép kín: đăng nhập, tạo đề, vào phòng thi, giám sát realtime, lưu dấu vết vi phạm, xem kết quả và quản trị tài khoản.")
    add_paragraph(document, "Phần AI service giúp dự án vượt khỏi mức hệ thống thi online thông thường. Nó cho phép kiểm tra camera, phân tích trạng thái bất thường và tạo thêm lớp dữ liệu cho dashboard giáo viên cũng như dashboard admin.")
    for idx, figure in enumerate(result_figures[:8], start=1):
        add_figure(document, 4, idx, str(figure["caption"]), figure["blob"], width_in=5.95)

    add_heading(document, "4.3. Kiểm thử nhanh", 2)
    add_table(document, ["Kịch bản", "Kết quả"], [
        ["Đăng nhập và phân quyền", "Đạt"],
        ["Sinh viên vào phòng thi", "Đạt"],
        ["Lưu nháp và nộp bài", "Đạt"],
        ["Giáo viên theo dõi realtime", "Đạt"],
        ["Admin xem dashboard và user", "Đạt"],
    ])

    add_heading(document, "KẾT LUẬN", 1, page_break=True)
    add_paragraph(document, "FE_DEMO đã đạt mục tiêu chính là xây dựng hệ thống thi trực tuyến có giám sát, có phân vai rõ ràng và có lớp phân tích gian lận thông minh. Bộ use case, thiết kế hệ thống và kết quả chạy thực tế đều cho thấy hệ thống bao phủ được cả người học, người dạy và quản trị viên.")
    add_paragraph(document, "Hướng phát triển tiếp theo là tăng độ chính xác của mô hình AI, hoàn thiện audit log, bổ sung báo cáo thống kê sâu hơn và tối ưu trải nghiệm phòng thi trên nhiều thiết bị.")


def main() -> None:
    if not TEMPLATE_DOCX.exists():
        raise FileNotFoundError(TEMPLATE_DOCX)
    document = Document(str(TEMPLATE_DOCX))
    clear_body(document)
    versions = {
        "vue": "^3.5.13",
        "vite": "^6.2.0",
        "pinia": "^3.0.4",
        "vue_router": "^4.5.1",
        "primevue": "^4.5.4",
        "echarts": "^6.0.0",
        "xlsx": "^0.18.5",
        "spring_boot": "4.0.3",
        "java": "17",
    }
    build_report(document, versions)
    document.save(str(OUTPUT_DOCX))
    print(f"Saved: {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
