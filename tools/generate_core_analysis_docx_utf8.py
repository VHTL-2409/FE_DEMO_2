from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Phan_tich_code_core_phat_hien_gian_lan_giam_sat_AI.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_font(run, name: str = "Arial", size: float = 9, bold: bool = False) -> None:
    run.bold = bold
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    set_font(run, size=9, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, True)
        set_cell_shading(table.rows[0].cells[index], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value)
    if widths:
        for table_row in table.rows:
            for index, width in enumerate(widths):
                table_row.cells[index].width = Inches(width)
    return table


def add_code(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    set_font(run, name="Consolas", size=8)
    run.font.color.rgb = RGBColor(48, 48, 48)


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(2)
    set_font(paragraph.add_run(text), size=10)


def add_number(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(2)
    set_font(paragraph.add_run(text), size=10)


def read_source_snippet(relative_path: str, start: int, end: int) -> str:
    path = ROOT / relative_path
    lines = path.read_text(encoding="utf-8", errors="replace").splitlines()
    selected = []
    in_block_comment = False
    for number in range(start, min(end, len(lines)) + 1):
        raw = lines[number - 1].rstrip()
        stripped = raw.strip()
        if not stripped:
            continue
        if stripped.startswith("#"):
            continue
        if stripped.startswith("//"):
            continue
        if stripped.startswith("/*") or stripped.startswith("/**"):
            in_block_comment = True
            continue
        if in_block_comment:
            if "*/" in stripped:
                in_block_comment = False
            continue
        if stripped.startswith("*"):
            continue
        selected.append(f"{number:>4}: {raw}")
    return "\n".join(selected)


def add_analysis_block(doc: Document, title: str, source: str, snippet: str, analysis: list[str]) -> None:
    doc.add_heading(title, level=2)
    doc.add_paragraph(f"Nguồn code: {source}")
    add_code(doc, snippet)
    for item in analysis:
        add_bullet(doc, item)


def add_code_sample_sections(doc: Document) -> None:
    doc.add_heading("8. Phân Tích Theo Code Mẫu Trích Từ Dự Án", level=1)
    doc.add_paragraph(
        "Phần này trích các đoạn code thật từ dự án để phân tích trực tiếp. "
        "Mỗi code mẫu được rút gọn theo các dòng quan trọng, bỏ bớt comment để dễ đọc trong Word."
    )

    add_analysis_block(
        doc,
        "8.1 Code mẫu: ProctorAnalyzer.analyze_frame",
        "ai-service/app/proctor.py:318-430",
        read_source_snippet("ai-service/app/proctor.py", 318, 430),
        [
            "Input chính là image_base64 và metadata. Metadata chứa attempt_id/session_id, trạng thái camera, mic, thời điểm capture.",
            "Hàm kiểm tra camera trước. Nếu camera tắt, hệ thống không cố phân tích ảnh mà trả response NO_CAMERA để backend ghi nhận signal.",
            "Ảnh được decode thành RGB và grayscale. RGB dùng cho DNN/MediaPipe, grayscale dùng cho Haar cascade, brightness, blur và occlusion.",
            "Pipeline xử lý theo thứ tự: detect face -> track eyes -> brightness/variance -> spoofing -> occlusion -> position -> gaze -> quality gate -> temporal tracking.",
            "Output không chỉ có signals mà còn có diagnostics và visual_overlay. Đây là bằng chứng để giáo viên xem vì sao hệ thống báo gian lận.",
        ],
    )

    add_analysis_block(
        doc,
        "8.2 Code mẫu: Sinh signal trong analyze_frame",
        "ai-service/app/proctor.py:380-610",
        read_source_snippet("ai-service/app/proctor.py", 380, 610),
        [
            "Các signal được sinh theo rule rõ ràng: NO_MIC, FACE_NOT_DETECTED, MULTIPLE_FACES, FACE_SPOOFING_SUSPECTED, EYES_NOT_DETECTED, GAZE_OFF_SCREEN, VERY_LOW_LIGHTING, VERY_BLURRY_FRAME.",
            "Một số signal cần stability và quality_gate. Ví dụ GAZE_OFF_SCREEN chỉ phát khi gaze_valid=true, còn face issue cần face_issue_stable.",
            "Severity được gán theo mức nguy hiểm: MULTIPLE_FACES và FACE_SPOOFING_SUSPECTED là CRITICAL; FACE_NOT_DETECTED/NO_MIC/VERY_LOW_LIGHTING là HIGH.",
            "Evidence đi kèm signal chứa lý do, khuyến nghị, qualityGate, detectionMethod. Backend dùng phần này để lưu evidence JSON.",
        ],
    )

    add_analysis_block(
        doc,
        "8.3 Code mẫu: Phân tích hướng nhìn _analyze_gaze",
        "ai-service/app/proctor.py:1539-1603",
        read_source_snippet("ai-service/app/proctor.py", 1539, 1603),
        [
            "Hàm chỉ chạy khi có đúng một khuôn mặt và có dữ liệu mắt. Điều này tránh suy luận gaze sai khi không đủ điều kiện.",
            "normalized_pupil của hai mắt được lấy trung bình. Tâm mắt là 0.5; lệch khỏi tâm theo ngưỡng ngang/dọc thì suy ra LEFT/RIGHT/UP/DOWN.",
            "gaze_confidence là tổ hợp của pupil_confidence và tracking_confidence. Nếu chỉ bắt được một mắt thì bị trừ confidence.",
            "Hàm chỉ tạo gaze_off_screen_candidate. Signal thật được xác nhận ở _update_tracking_state sau khi lệch hướng đủ số frame hoặc đủ thời gian.",
        ],
    )

    add_analysis_block(
        doc,
        "8.4 Code mẫu: Bộ nhớ trạng thái _update_tracking_state",
        "ai-service/app/proctor.py:1605-1847",
        read_source_snippet("ai-service/app/proctor.py", 1605, 1847),
        [
            "tracking_key thường lấy từ attempt_id/session_id. Nhờ vậy mỗi bài thi có state riêng, không lẫn dữ liệu giữa học sinh.",
            "State lưu face_status, closed_since_ms, off_screen_since_ms, streak, gaze_history và last_signal_emit_ms.",
            "Face, eye missing, face turned, eye closure và gaze off screen chỉ được coi là ổn định nếu đạt required_frames hoặc required_ms.",
            "Khi face_count không bằng 1, state liên quan gaze/eye được reset. Đây là xử lý quan trọng vì gaze chỉ có ý nghĩa khi có đúng một khuôn mặt.",
            "attention_score bị hạ nếu mắt nhắm lâu hoặc gaze off screen. Đây là metric tổng hợp cho dashboard.",
        ],
    )

    add_analysis_block(
        doc,
        "8.5 Code mẫu: Detect mặt bằng DNN/Haar fallback",
        "ai-service/app/proctor.py:2079-2199",
        read_source_snippet("ai-service/app/proctor.py", 2079, 2199),
        [
            "DNN được ưu tiên vì ổn định hơn Haar. Nếu DNN không thấy mặt, Haar vẫn được chạy để giảm bỏ sót.",
            "DNN nhận ảnh BGR resize 300x300 theo chuẩn model OpenCV Caffe, sau đó convert bounding box về kích thước frame gốc.",
            "Các face boxes được merge bằng IoU hoặc center-close để tránh đếm trùng một khuôn mặt.",
            "detection_method được trả về diagnostics, giúp debug khi kết quả đến từ DNN, HAAR_FALLBACK hoặc DNN+HAAR.",
        ],
    )

    add_analysis_block(
        doc,
        "8.6 Code mẫu: Bridge backend gọi AI - AiAssistService.analyzeFrame",
        "BE/src/main/java/com/example/demo/service/AiAssistService.java:128-176",
        read_source_snippet("BE/src/main/java/com/example/demo/service/AiAssistService.java", 128, 176),
        [
            "Backend nhận frame từ frontend, validate payload rồi publish trước một bản AI_PENDING để giáo viên vẫn thấy camera đang gửi.",
            "Nếu AI service bật, backend gọi /proctor/analyze/frame, nhận response và collect AI signals.",
            "Signal AI được cluster trước khi ghi DB để tránh mỗi frame tạo một FraudSignal giống nhau.",
            "Nếu gọi AI lỗi, backend trả AI_UNAVAILABLE nhưng không làm chết luồng camera. Đây là thiết kế fallback tốt cho giám sát realtime.",
        ],
    )

    add_analysis_block(
        doc,
        "8.7 Code mẫu: Ghi signal AI vào backend",
        "BE/src/main/java/com/example/demo/service/AiAssistService.java:358-430",
        read_source_snippet("BE/src/main/java/com/example/demo/service/AiAssistService.java", 358, 430),
        [
            "safeBridgeAiSignals đọc danh sách signals trong response AI và bỏ qua signal không thuộc nhóm camera/gaze/spoofing hợp lệ.",
            "Mỗi signal được chuẩn hóa type, confidence, severity và evidence trước khi gọi FraudSignalService.recordServerSignal.",
            "Sau khi ghi signal, backend gọi riskScoringService.recomputeRisk để cập nhật điểm rủi ro ngay.",
            "Đây là đoạn code nối AI service với hệ thống chấm điểm chính của backend.",
        ],
    )

    add_analysis_block(
        doc,
        "8.8 Code mẫu: Ingest event trình duyệt - ExamEventService.ingestBatchInternal",
        "BE/src/main/java/com/example/demo/service/ExamEventService.java:95-148",
        read_source_snippet("BE/src/main/java/com/example/demo/service/ExamEventService.java", 95, 148),
        [
            "Frontend gửi nhiều event theo batch. Backend kiểm tra attempt, quyền học sinh, trạng thái bài thi và fingerprint.",
            "Mỗi event có sequenceNo để chống gửi lặp. Ngoài ra còn có duplicate burst và signal rate-limit.",
            "Event hợp lệ được lưu thành ExamEvent, MonitoringEvent và FraudSignal. Sau đó backend recomputeRisk một lần cho cả batch.",
            "Thiết kế này giảm số lần tính risk và chống spam event như TAB_SWITCH liên tục.",
        ],
    )

    add_analysis_block(
        doc,
        "8.9 Code mẫu: Chuẩn hóa FraudSignal",
        "BE/src/main/java/com/example/demo/service/FraudSignalService.java:43-108",
        read_source_snippet("BE/src/main/java/com/example/demo/service/FraudSignalService.java", 43, 108),
        [
            "recordFromEvent dùng cho signal đến từ browser event. recordServerSignal dùng cho AI hoặc backend detection.",
            "Cả hai đường đều dùng descriptorFor để chuẩn hóa signalType, category, displayMessage, riskImpact và severity.",
            "Evidence được serialize JSON để lưu lại bối cảnh: payload, browserContext, telemetry hoặc evidence từ AI.",
            "Nhờ lớp này, RiskScoringService không cần biết signal đến từ frontend, AI hay backend rule.",
        ],
    )

    add_analysis_block(
        doc,
        "8.10 Code mẫu: Từ điển signal descriptorFor",
        "BE/src/main/java/com/example/demo/service/FraudSignalService.java:115-230",
        read_source_snippet("BE/src/main/java/com/example/demo/service/FraudSignalService.java", 115, 230),
        [
            "descriptorFor là bảng mapping nghiệp vụ. Nó quyết định một event/signal thuộc category nào, severity gì và điểm ảnh hưởng mặc định bao nhiêu.",
            "Ví dụ TAB_SWITCH thuộc SCREEN_LEAVE, DEVTOOLS_OPEN thuộc TECHNICAL, MULTIPLE_FACES thuộc AI_CAMERA.",
            "Nếu signal mới không được bổ sung vào descriptor, nó sẽ rơi vào OTHER, severity LOW và risk thấp.",
            "Khi bảo trì hệ thống, đây là nơi cần cập nhật đầu tiên nếu thêm loại gian lận mới.",
        ],
    )

    add_analysis_block(
        doc,
        "8.11 Code mẫu: Tính điểm rủi ro - RiskScoringService.recomputeRisk",
        "BE/src/main/java/com/example/demo/service/RiskScoringService.java:221-311",
        read_source_snippet("BE/src/main/java/com/example/demo/service/RiskScoringService.java", 221, 311),
        [
            "Hàm lấy toàn bộ FraudSignal của attempt rồi tính điểm theo category.",
            "screenLeave, clipboard, technical, identity, heartbeat, visualIdentity đều có cap riêng.",
            "behaviorScore bị giới hạn 70, totalRisk bị giới hạn 100 để tránh một nhóm hành vi áp đảo toàn bộ kết quả.",
            "Sau khi tính score, backend cập nhật ExamAttempt, đồng bộ ProctorFlag, lưu RiskScoreLog nếu cần và gửi realtime notification.",
        ],
    )

    add_analysis_block(
        doc,
        "8.12 Code mẫu: Cluster điểm theo thời gian",
        "BE/src/main/java/com/example/demo/service/RiskScoringService.java:358-447",
        read_source_snippet("BE/src/main/java/com/example/demo/service/RiskScoringService.java", 358, 447),
        [
            "computeCategoryScores gom signal theo category, bỏ OTHER, rồi gọi computeClusteredScore.",
            "computeClusteredScore sắp xếp signal theo thời gian. Trong cùng cửa sổ dedup, chỉ lấy điểm cao nhất.",
            "IDENTITY là trường hợp đặc biệt: lấy max trong session thay vì cộng nhiều lần.",
            "Cơ chế này làm risk score ổn định hơn khi học sinh hoặc browser gửi lặp cùng một loại event.",
        ],
    )

    add_analysis_block(
        doc,
        "8.13 Code mẫu: Dashboard realtime - useExamMonitoring.applyRealtimeEvent",
        "demo/src/composables/useExamMonitoring.js:42-116",
        read_source_snippet("demo/src/composables/useExamMonitoring.js", 42, 116),
        [
            "Frontend nhận payload realtime từ WebSocket, chuẩn hóa type và attemptId.",
            "Dashboard card được patch trước, sau đó alert panel và system timeline được cập nhật.",
            "Các event như FRAUD_SIGNAL_RECORDED, RISK_UPDATED, AI_CAMERA_SIGNAL, ATTEMPT_PAUSED đều đi qua cùng một hàm.",
            "Nhờ tách addAlertFromEvent và addSystemEventFromEvent, UI dễ mở rộng khi backend thêm event mới.",
        ],
    )


def build_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    doc.styles["Normal"].font.size = Pt(10.5)
    doc.styles["Heading 1"].font.size = Pt(16)
    doc.styles["Heading 2"].font.size = Pt(13)
    doc.styles["Heading 3"].font.size = Pt(11.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("PHÂN TÍCH CODE CORE\nPHÁT HIỆN GIAN LẬN - GIÁM SÁT - AI")
    set_font(run, size=20, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(
        subtitle.add_run("Dự án: Hệ thống thi trực tuyến với cơ chế giám sát và phát hiện gian lận"),
        size=12,
    )

    created = doc.add_paragraph()
    created.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(created.add_run(f"Ngày lập tài liệu: {datetime.now().strftime('%d/%m/%Y %H:%M')}"), size=10)

    doc.add_paragraph()
    add_table(doc, ["Thành phần", "Công nghệ", "Vai trò trong hệ thống"], [
        ["ai-service", "Python FastAPI, OpenCV, MediaPipe, PIL/Numpy", "Phân tích frame camera, nhận diện mặt, mắt, gaze, giả mạo ảnh, hành vi paste/tab/idle."],
        ["BE", "Spring Boot Java", "Nhận sự kiện, lưu FraudSignal, tính điểm rủi ro, tạo flag, đẩy realtime cho giáo viên."],
        ["demo", "Vue/Vite", "Thu thập tín hiệu trình duyệt, gửi heartbeat/frame, hiển thị dashboard giám sát realtime."],
    ])
    doc.add_page_break()

    doc.add_heading("Mục Lục", level=1)
    for item in [
        "1. Tóm tắt kiến trúc và luồng dữ liệu",
        "2. Bản đồ file core",
        "3. Phân tích nhóm AI service",
        "4. Phân tích nhóm Backend phát hiện gian lận và chấm điểm rủi ro",
        "5. Phân tích nhóm Frontend giám sát",
        "6. Bảng tín hiệu gian lận và điểm rủi ro",
        "7. Nhận xét kỹ thuật, rủi ro và đề xuất cải tiến",
        "8. Phân tích theo code mẫu trích từ dự án",
        "9. Phụ lục pseudocode các hàm core",
        "10. Kết luận",
    ]:
        doc.add_paragraph(item)
    doc.add_page_break()

    doc.add_heading("1. Tóm Tắt Kiến Trúc Và Luồng Dữ Liệu", level=1)
    doc.add_paragraph(
        "Hệ thống phát hiện gian lận được thiết kế theo mô hình nhiều lớp. Frontend thu thập sự kiện từ trình duyệt và camera; "
        "backend chuẩn hóa thành FraudSignal, tính điểm rủi ro; ai-service phân tích hình ảnh/hành vi để tạo tín hiệu AI; "
        "dashboard giáo viên nhận realtime để giám sát."
    )
    add_table(doc, ["Bước", "Nguồn", "Hàm/API core", "Kết quả"], [
        ["1", "Student UI", "useExamProctoring.start(), event batch, heartbeat, frame upload", "Bắt đầu phiên giám sát, gửi fingerprint, trạng thái camera/mic và sự kiện trình duyệt."],
        ["2", "Backend BE", "ExamEventService.ingestBatch(), processHeartbeat()", "Lưu ExamEvent, MonitoringEvent, tạo FraudSignal từ event rõ ràng."],
        ["3", "AI Bridge", "AiAssistService.analyzeFrame()", "Nhận frame, publish trạng thái pending, gọi ai-service, lọc/cluster tín hiệu AI."],
        ["4", "AI service", "ProctorAnalyzer.analyze_frame()", "Phân tích mặt, mắt, gaze, ánh sáng, blur, spoofing và trả danh sách signals."],
        ["5", "Risk engine", "RiskScoringService.recomputeRisk()", "Cộng điểm theo category, dedup theo cửa sổ thời gian, tạo RiskScore và ProctorFlag."],
        ["6", "Realtime dashboard", "TeacherAlertGateway + useExamMonitoring.applyRealtimeEvent()", "Giáo viên thấy risk score, alert, timeline và frame camera mới."],
    ])
    add_code(doc, "Student Browser -> BE -> FraudSignalService -> RiskScoringService -> WebSocket Dashboard\nStudent Camera Frame -> BE AiAssistService -> ai-service ProctorAnalyzer -> BE FraudSignal/Risk -> Teacher Dashboard")

    doc.add_heading("2. Bản Đồ File Core", level=1)
    add_table(doc, ["File", "Dòng chính", "Nội dung core"], [
        ["ai-service/app/proctor.py", "318, 863, 935, 1539, 1605, 2079", "Pipeline AI camera: decode ảnh, detect face, track eyes, phân tích gaze, quality gate, stateful stability, lọc signal."],
        ["ai-service/app/main.py", "191, 217", "Endpoint /proctor/analyze/frame và /proctor/analyze/behavior."],
        ["BE/.../AiAssistService.java", "128, 358, 374, 534, 1024", "Bridge backend gọi ai-service, publish frame realtime, lưu ảnh bằng chứng, ghi FraudSignal/FraudWarning."],
        ["BE/.../ExamEventService.java", "95, 158, 249, 388", "Ingest event batch, heartbeat, dedupe/rate-limit, chuyển event thành fraud signal."],
        ["BE/.../FraudSignalService.java", "43, 87, 115", "Chuẩn hóa signal, gán category/severity/riskImpact, lưu evidence JSON."],
        ["BE/.../RiskScoringService.java", "221, 358, 392, 450, 534, 570", "Chấm điểm rủi ro theo category, cluster/dedup, tính contribution, đồng bộ ProctorFlag."],
        ["BE/.../MonitoringService.java", "48, 242, 300+", "Timeline, hành động giám thị, device status, camera status/alerts."],
        ["BE/.../MLRiskScoringService.java", "124, 245, 440, 531", "Phân tích rủi ro kiểu ML/hybrid, hiện rule-based là nguồn chính khi ML disabled."],
        ["demo/src/composables/useExamProctoring.js", "20, 38, 60", "Khởi tạo phiên proctoring, đồng bộ risk state."],
        ["demo/src/composables/useExamMonitoring.js", "42, 141", "Nhận realtime event, patch dashboard, fallback polling."],
        ["demo/src/composables/useAiCameraDashboard.js", "160", "Dashboard AI camera: cập nhật status, active signals, alert list."],
    ])

    doc.add_heading("3. Phân Tích Nhóm AI Service", level=1)
    doc.add_paragraph(
        "Đây là lớp xử lý nhận thức máy tính. Các hàm core nằm chủ yếu trong ProctorAnalyzer. "
        "Cách triển khai hiện tại kết hợp deep learning face detector, Haar fallback, MediaPipe FaceMesh, "
        "heuristic về ánh sáng/blur/gaze và state theo attempt để giảm báo nhầm."
    )
    add_table(doc, ["Hàm", "Vị trí", "Vai trò", "Cách hoạt động", "Ý nghĩa giám sát"], [
        ["ProctorAnalyzer.__init__", "ai-service/app/proctor.py:48", "Khởi tạo cascade, DNN face detector, landmark detector, spoofing detector và tracking_state.", "Thiết lập ngưỡng từ biến môi trường: face threshold, signal dedup, required frames/ms cho face/gaze/eye closure.", "Nếu OpenCV/MediaPipe thiếu, hệ thống vẫn trả diagnostics thay vì crash; DNN có fallback về Haar."],
        ["analyze_frame", "ai-service/app/proctor.py:318", "Hàm lõi nhất của AI camera. Nhận image_base64 và metadata, trả metrics, signals, warnings, visual_overlay, diagnostics.", "Kiểm tra camera/mic -> decode ảnh -> detect face -> track eyes -> đo brightness/variance -> spoofing -> occlusion -> face position -> gaze -> quality gate -> temporal tracking -> sinh signal -> lọc signal.", "Có state theo attempt/session để chỉ báo khi trạng thái ổn định qua nhiều frame, giảm false positive."],
        ["_detect_face_locations", "ai-service/app/proctor.py:2079", "Chọn backend detect mặt: DNN nếu sẵn sàng, Haar fallback nếu DNN không thấy mặt hoặc lỗi.", "Ưu tiên DNN; nếu không có model hoặc DNN fail thì dùng Haar cascade, profile cascade và ảnh flip.", "Trả detection_method trong diagnostics để backend/UI biết frame được phân tích bằng mô hình nào."],
        ["_detect_faces_dnn", "ai-service/app/proctor.py:2099", "Chạy OpenCV DNN Caffe/TensorFlow trên input 300x300, lọc confidence vượt ngưỡng.", "Scale box về kích thước frame và merge các box trùng/lệch gần nhau.", "Ngưỡng confidence mặc định 0.55, có thể chỉnh qua AI_SERVICE_FACE_DETECTION_THRESHOLD."],
        ["_track_eyes", "ai-service/app/proctor.py:935", "Theo dõi mắt trong vùng mặt.", "Ưu tiên FaceMesh/landmark nếu có, fallback bằng Haar eye cascade và phân tích ROI mắt.", "Là đầu vào cho gaze và phát hiện EYES_NOT_DETECTED/EYES_CLOSED_PROLONGED."],
        ["_analyze_gaze", "ai-service/app/proctor.py:1539", "Suy luận hướng nhìn từ normalized pupil.", "Tính avg_x/avg_y, offset so với tâm 0.5, so với threshold ngang/dọc để phát hiện gaze_off_screen_candidate.", "Không phát signal ngay; _update_tracking_state xác nhận qua nhiều frame/thời gian."],
        ["_update_tracking_state", "ai-service/app/proctor.py:1605", "Bộ nhớ trạng thái theo attempt/session.", "Xác nhận ổn định cho no face, multiple faces, eye missing, face turned, gaze off screen, prolonged eye closure.", "Biến phân tích từng ảnh thành giám sát liên tục."],
        ["_filter_frame_signals", "ai-service/app/proctor.py:863", "Lọc, ưu tiên và dedup các signal sinh từ cùng frame hoặc nhiều frame gần nhau.", "Chuẩn hóa type/severity, xếp priority, chặn spam một loại signal trong khoảng AI_SERVICE_SIGNAL_DEDUP_SECONDS.", "Giúp risk score không bị tăng giả tạo do FE gửi frame liên tục."],
        ["_build_quality_gate", "ai-service/app/proctor.py:745", "Đánh giá độ tin cậy của frame trước khi cho phép signal nhạy.", "Dựa trên kích thước mặt, ánh sáng, độ nét, eye tracking confidence, vị trí mặt.", "GAZE_OFF_SCREEN chỉ phát khi quality_gate.gaze_valid=true."],
        ["analyze_behavior", "ai-service/app/proctor.py:2361", "Phân tích hành vi phi hình ảnh.", "Dựa vào paste_length, tab_switch_count, idle_seconds, typing_intervals.", "Sinh COPY_PASTE_DETECTED, TAB_SWITCH_CLUSTER, IDLE_EXCESSIVE, TYPING_ANOMALY."],
    ], widths=[1.4, 1.3, 2.0, 2.6, 2.0])

    doc.add_heading("3.1 Pipeline analyze_frame Chi Tiết", level=2)
    for step in [
        "Nhận metadata camera/mic. Nếu camera tắt, trả response NO_CAMERA ngay.",
        "Decode base64 thành ảnh RGB bằng PIL, đồng thời tạo ảnh grayscale để OpenCV xử lý nhanh.",
        "Detect face bằng DNN trước, Haar fallback sau.",
        "Track mắt và pupil trong vùng mặt.",
        "Tính brightness và variance để phát hiện ánh sáng yếu/quá sáng hoặc frame mờ.",
        "Phân tích spoofing, occlusion, face position và gaze.",
        "Quality gate + tracking state xác nhận tín hiệu qua nhiều frame/thời gian.",
        "Tạo signal chuẩn hóa, lọc/dedup, trả diagnostics và overlay cho UI.",
    ]:
        add_number(doc, step)

    doc.add_heading("4. Phân Tích Backend Phát Hiện Gian Lận Và Chấm Điểm", level=1)
    doc.add_paragraph(
        "Backend là nguồn chuẩn hóa và ra quyết định. AI chỉ tạo tín hiệu; quyết định risk score, flag, timeline, realtime "
        "đều nằm trong BE để dễ audit và cấu hình."
    )
    add_table(doc, ["Hàm", "Vị trí", "Nhóm", "Cách hoạt động", "Tại sao core"], [
        ["AiAssistService.analyzeFrame", "BE/.../AiAssistService.java:128", "Cầu nối camera AI", "Validate attemptId/image size, publish frame pending, gọi /proctor/analyze/frame, collect signals, cluster/dedup, lưu ảnh bằng chứng, ghi FraudSignal/FraudWarning, publish realtime.", "Nếu AI service down vẫn publish frame với AI_UNAVAILABLE để giáo viên không mất stream camera."],
        ["AiAssistService.safeBridgeAiSignals", "BE/.../AiAssistService.java:358", "Ghi tín hiệu AI vào fraud pipeline", "Duyệt signals từ AI response, chuẩn hóa signalType/severity/confidence/evidence rồi gọi FraudSignalService.recordServerSignal.", "Đảm bảo AI signal trở thành cùng một loại dữ liệu với browser signal."],
        ["AiAssistService.clusterAiCameraSignalsForRecording", "BE/.../AiAssistService.java:534", "Giảm spam signal camera", "Gom cụm signal theo attempt/type/risk, chỉ record khi vượt dedup hoặc có trạng thái mới đáng ghi.", "Camera frame gửi liên tục; nếu không cluster, risk score sẽ tăng bất thường."],
        ["ExamEventService.ingestBatchInternal", "BE/.../ExamEventService.java:95", "Ingest sự kiện trình duyệt", "Kiểm tra quyền, attempt active, fingerprint/IP, duyệt event theo sequence, bỏ duplicate, rate-limit theo signal, lưu ExamEvent/MonitoringEvent/FraudSignal.", "Cửa vào chính của TAB_SWITCH, WINDOW_BLUR, COPY_PASTE, DEVTOOLS_OPEN."],
        ["ExamEventService.processHeartbeat", "BE/.../ExamEventService.java:158/196", "Heartbeat và thiết bị", "Cập nhật lastHeartbeatAt, cameraOn, micOn, fingerprint/IP; nếu camera/mic chuyển từ on sang off thì record NO_CAMERA/NO_MIC.", "Giúp phát hiện tắt camera/mic trong lúc thi."],
        ["FraudSignalService.recordFromEvent", "BE/.../FraudSignalService.java:43", "Chuyển ExamEvent thành FraudSignal", "Lấy descriptor theo eventType, normalize confidence, build evidence gồm browserContext/telemetry/payload, lưu DB.", "Tạo chuẩn dữ liệu thống nhất để risk engine không phụ thuộc trực tiếp vào FE payload."],
        ["FraudSignalService.descriptorFor", "BE/.../FraudSignalService.java:115", "Từ điển signal", "Map event/signal thành signalType, category, display message, riskImpact, severity, default confidence.", "Đây là bảng cấu hình nghiệp vụ quan trọng."],
        ["RiskScoringService.recomputeRisk", "BE/.../RiskScoringService.java:221", "Chấm điểm canonical", "Lấy FraudSignal của attempt, compute category scores, cap từng category, tính totalRisk, resolve level, cập nhật attempt, sync flag, lưu snapshot, notify realtime.", "Quyết định CLEAN/SUSPICIOUS/HIGH_RISK/CRITICAL."],
        ["RiskScoringService.computeClusteredScore", "BE/.../RiskScoringService.java:392", "Dedup theo cửa sổ thời gian", "Sắp xếp signal theo thời gian, chia cluster theo category window, mỗi cluster chỉ lấy điểm cao nhất, sau đó cap category.", "Ngăn spam cùng loại event làm phình điểm."],
        ["RiskScoringService.syncRiskFlag", "BE/.../RiskScoringService.java:534", "Tạo/cập nhật flag review", "Nếu score >= highRiskMin thì mở ProctorFlag HIGH_RISK/CRITICAL.", "Biến điểm rủi ro thành hành động review cho giáo viên."],
        ["MonitoringService.timeline", "BE/.../MonitoringService.java:242", "Gộp timeline", "Gộp ExamEvent, MonitoringEvent, FraudSignal, FraudWarning, RiskScoreLog rồi sort theo thời gian.", "Giáo viên có toàn cảnh diễn biến phiên thi."],
        ["MLRiskScoringService.analyzeRisk", "BE/.../MLRiskScoringService.java:124", "ML/hybrid", "Trích xuất feature từ FraudSignal, tính rule-based score, nếu ML enabled thì kết hợp ML score theo confidence.", "Hiện mặc định ML disabled, rule-based là nguồn canonical."],
    ], widths=[1.6, 1.2, 1.4, 3.0, 2.2])

    doc.add_heading("4.1 Công Thức Risk Score", level=2)
    add_table(doc, ["Category", "Signal tiêu biểu", "Dedup/cách cộng", "Cap", "Ý nghĩa"], [
        ["SCREEN_LEAVE", "TAB_SWITCH, WINDOW_BLUR, EXIT_FULLSCREEN", "Cluster 60s, mỗi cluster lấy điểm cao nhất", "35", "Rời màn hình hoặc thoát fullscreen."],
        ["CLIPBOARD", "COPY_PASTE", "Cluster 30s", "25", "Sao chép/dán nội dung."],
        ["TECHNICAL", "RIGHT_CLICK, PRINT_SCREEN, DEVTOOLS_OPEN", "Cluster 120s", "25", "Dùng công cụ hoặc thao tác kỹ thuật đáng nghi."],
        ["IDENTITY", "DUPLICATE_IP, IP_CHANGED, DEVICE_FINGERPRINT_CHANGED", "Lấy max trong session", "30", "Bất thường định danh, thiết bị, IP."],
        ["HEARTBEAT", "HEARTBEAT_STALE, NETWORK_INSTABILITY, SESSION_RECOVERY", "Theo session", "10", "Mất kết nối hoặc phục hồi phiên."],
        ["VISUAL_IDENTITY", "NO_CAMERA, MULTIPLE_FACES, FACE_NOT_DETECTED, GAZE_OFF_SCREEN", "Cluster camera AI 30s", "40", "Bất thường khuôn mặt, mắt, camera/mic."],
    ])
    add_code(doc, "behaviorScore = min(70, screenLeave + clipboard + technical + heartbeat)\ntotalRisk = min(100, behaviorScore + identityScore + visualIdentityScore)\nLevel: 0-19 CLEAN, 20-49 SUSPICIOUS, 50-74 HIGH_RISK, 75-100 CRITICAL")

    doc.add_heading("5. Phân Tích Frontend Giám Sát", level=1)
    add_table(doc, ["Hàm", "Vị trí", "Cách hoạt động", "Vai trò"], [
        ["useExamProctoring.start", "demo/src/composables/useExamProctoring.js:38", "Gửi startProctoringSession với attemptId/sessionId, fingerprint, userAgent, timezone, camera/mic, viewport.", "Khởi tạo phiên giám sát, giúp backend có baseline thiết bị/IP và trạng thái ban đầu."],
        ["useExamProctoring.syncRisk", "demo/src/composables/useExamProctoring.js:20", "Chuẩn hóa payload riskScore/level/flag thành riskState local.", "Giữ UI đồng bộ dù backend trả field theo nhiều format."],
        ["useExamMonitoring.applyRealtimeEvent", "demo/src/composables/useExamMonitoring.js:42", "Nhận event websocket, xác định type/attemptId, patch card, thêm alert và system event.", "Hàm core của dashboard realtime giáo viên."],
        ["useExamMonitoring.connect", "demo/src/composables/useExamMonitoring.js:141", "Subscribe /topic/exams/{id}/alerts và per-attempt topics, fallback polling nếu websocket không kết nối.", "Đảm bảo dashboard vẫn cập nhật khi realtime lỗi."],
        ["useAiCameraDashboard.applyRealtimeEvent", "demo/src/composables/useAiCameraDashboard.js:160", "Cập nhật cameraStatuses, recentAlerts, activeSignals, criticalSignals từ AI camera realtime.", "Hiển thị trạng thái camera AI từng học sinh."],
    ], widths=[2.0, 1.7, 3.2, 2.3])

    doc.add_heading("6. Bảng Tín Hiệu Gian Lận Quan Trọng", level=1)
    add_table(doc, ["Signal", "Category", "Severity", "Điểm", "Điều kiện/ý nghĩa"], [
        ["NO_CAMERA", "AI_CAMERA/VISUAL_IDENTITY", "HIGH", "20", "Camera bị tắt hoặc track disabled/ended."],
        ["NO_MIC", "AI_CAMERA/VISUAL_IDENTITY", "HIGH", "20", "Micro bị tắt."],
        ["FACE_NOT_DETECTED", "AI_CAMERA/VISUAL_IDENTITY", "HIGH", "20", "Không thấy khuôn mặt ổn định qua nhiều frame."],
        ["MULTIPLE_FACES", "AI_CAMERA/VISUAL_IDENTITY", "CRITICAL", "25", "Nhiều khuôn mặt trong khung hình."],
        ["FACE_SPOOFING_SUSPECTED", "AI_SPOOFING/VISUAL_IDENTITY", "CRITICAL", "25", "Nghi ảnh in, màn hình phát lại hoặc giả mạo khuôn mặt."],
        ["GAZE_OFF_SCREEN", "GAZE_TRACKING/VISUAL_IDENTITY", "HIGH", "12", "Nhìn ra ngoài màn hình đủ lâu và quality gate hợp lệ."],
        ["EYES_NOT_DETECTED", "EYE_TRACKING/VISUAL_IDENTITY", "MEDIUM", "8", "Không nhận diện được mắt khi có một mặt."],
        ["VERY_LOW_LIGHTING", "AI_CAMERA/VISUAL_IDENTITY", "HIGH", "15", "Ánh sáng quá yếu, ảnh hưởng xác thực."],
        ["VERY_BLURRY_FRAME", "AI_CAMERA/VISUAL_IDENTITY", "HIGH", "15", "Frame quá mờ, khó giám sát."],
        ["TAB_SWITCH", "SCREEN_LEAVE", "MEDIUM", "10", "Chuyển tab trong lúc thi."],
        ["WINDOW_BLUR", "SCREEN_LEAVE", "MEDIUM", "8", "Cửa sổ mất focus."],
        ["EXIT_FULLSCREEN", "SCREEN_LEAVE", "HIGH", "15", "Thoát toàn màn hình."],
        ["COPY_PASTE", "CLIPBOARD", "HIGH", "10", "Sao chép/dán nội dung."],
        ["DEVTOOLS_OPEN", "TECHNICAL", "HIGH", "22", "Mở công cụ lập trình."],
        ["PRINT_SCREEN", "TECHNICAL", "HIGH", "15", "Chụp màn hình."],
        ["DUPLICATE_IP/IP_FINGERPRINT_GRAPH", "IDENTITY", "HIGH", "25-30", "Nhiều thí sinh/IP/fingerprint có liên hệ đáng ngờ."],
    ], widths=[1.8, 1.8, 1.0, 0.8, 3.6])

    doc.add_heading("7. Nhận Xét Kỹ Thuật", level=1)
    doc.add_heading("7.1 Điểm Mạnh", level=2)
    for text in [
        "Thiết kế tách lớp rõ: AI service chỉ phân tích, backend quyết định risk/flag, frontend hiển thị và thu thập tín hiệu.",
        "Risk scoring minh bạch, có bảng điểm/cap/dedup nên dễ giải thích với giáo viên và hội đồng.",
        "AI camera có quality gate và temporal state, không phụ thuộc vào một frame đơn lẻ.",
        "Có fallback khi thiếu DNN/MediaPipe hoặc AI service down, giúp hệ thống không bị dừng toàn bộ.",
        "Evidence JSON lưu nhiều thông tin: browser context, telemetry, diagnostics, overlay, ảnh bằng chứng nếu có.",
        "Realtime dashboard có fallback polling, giảm rủi ro mất cập nhật khi WebSocket lỗi.",
    ]:
        add_bullet(doc, text)

    doc.add_heading("7.2 Rủi Ro/Khoảng Trống", level=2)
    for text in [
        "Một số comment trong source Python/Java đang bị lỗi encoding ký tự tiếng Việt; không ảnh hưởng runtime nhưng ảnh hưởng bảo trì và báo cáo.",
        "MLRiskScoringService hiện chủ yếu là rule-based/statistical fallback; callMLService còn TODO, vì vậy không nên trình bày là ML model đã huấn luyện hoàn chỉnh nếu chưa có external ML service.",
        "Heuristic spoofing/occlusion/lighting có thể false positive trong môi trường ánh sáng yếu, webcam kém hoặc học sinh đeo kính.",
        "Camera frame base64 gửi qua HTTP có giới hạn kích thước nhưng vẫn tốn băng thông; cần cân nhắc nén/chọn tần suất gửi.",
        "Một số hành vi nâng cao như sync behavior, heartbeat-derived signal, fingerprint change auto signal đang được để no-op ở giai đoạn 1.",
        "Risk score phụ thuộc mạnh vào mapping descriptor và RiskSignalScoreProperties; sai cấu hình có thể làm score lệch.",
    ]:
        add_bullet(doc, text)

    doc.add_heading("7.3 Đề Xuất Cải Tiến", level=2)
    for text in [
        "Chuẩn hóa encoding source sang UTF-8 và sửa comment tiếng Việt để tài liệu kỹ thuật đọc được trực tiếp.",
        "Tách bảng signal/score/category thành file cấu hình hoặc bảng DB có version để audit thay đổi điểm rủi ro.",
        "Bổ sung unit test cho RiskScoringService với các case: spam tab switch, nhiều signal cùng window, visual identity cap, contribution delta.",
        "Bổ sung integration test cho luồng AI frame: frame -> ai-service -> FraudSignal -> risk update -> websocket payload.",
        "Nếu muốn gọi là AI/ML đầy đủ, cần hoàn thiện external ML service, dữ liệu train, nhãn gian lận, metrics precision/recall và cơ chế rollback model.",
        "Thêm calibration UI cho camera trước khi thi để giảm signal ánh sáng/blur do môi trường, không phải do gian lận.",
        "Lưu audit trail cho các quyết định teacher review: xác nhận đúng/sai signal để tạo dataset cải thiện model.",
    ]:
        add_bullet(doc, text)

    add_code_sample_sections(doc)

    doc.add_heading("9. Phụ Lục: Pseudocode Các Hàm Core", level=1)
    doc.add_heading("9.1 ProctorAnalyzer.analyze_frame", level=2)
    add_code(doc, "metadata -> camera/mic state\nif camera_off: return NO_CAMERA signal\nimage = decode_base64(image_base64)\nface_count, face_locations = detect_face_locations(image)\neye_tracking = track_eyes(image, face_locations)\nquality = build_quality_gate(face, eyes, brightness, variance)\ngaze = analyze_gaze(eye_tracking)\nstability = update_tracking_state(attempt/session, face, eyes, gaze)\nsignals = build_signals(face, spoofing, occlusion, lighting, blur, gaze, mic)\nsignals = filter_frame_signals(signals, tracking_key, sample_time)\nreturn metrics + signals + overlay + diagnostics")

    doc.add_heading("9.2 RiskScoringService.recomputeRisk", level=2)
    add_code(doc, "signals = findByAttemptOrderByCreatedAtAsc(attempt)\ncategoryScores = computeCategoryScores(signals)\nscore = min(100, min(70, screenLeave + clipboard + technical + heartbeat) + identity + visualIdentity)\nlevel = resolveLevel(score)\nupdate attempt.riskScore/riskLevel/suspicious\nflag = syncRiskFlag(attempt, score, level)\nif snapshot needed: save RiskScoreLog\nnotify latest FraudSignal + risk update\nreturn RiskScoreResponse")

    doc.add_heading("9.3 ExamEventService.ingestBatchInternal", level=2)
    add_code(doc, "attempt = requireAttempt(attemptId)\nensure student owns attempt and attempt active\nnormalize fingerprint; run identity anomaly detection\nfor each event item with sequence:\n  skip if duplicate sequence or duplicate burst or signal rate-limited\n  save ExamEvent and legacy MonitoringEvent\n  FraudSignalService.recordFromEvent(...)\nrisk = RiskScoringService.recomputeRisk(attempt)\nreturn accepted/dropped counts + risk score/level")

    doc.add_heading("10. Kết Luận", level=1)
    doc.add_paragraph(
        "Các hàm core nhất của dự án tập trung vào ba trục: ProctorAnalyzer.analyze_frame ở ai-service để sinh tín hiệu thị giác; "
        "ExamEventService/FraudSignalService ở backend để chuẩn hóa và ghi nhận tín hiệu; RiskScoringService.recomputeRisk để tính điểm và tạo flag; "
        "các composable frontend để khởi tạo phiên, nhận realtime và hiển thị dashboard. Kiến trúc hiện tại phù hợp với hệ thống thi trực tuyến cần giải thích được quyết định phát hiện gian lận, "
        "vì phần chấm điểm chính là rule-based minh bạch, còn AI camera đóng vai trò nguồn tín hiệu bổ sung có diagnostics/evidence."
    )
    return doc


def main() -> None:
    doc = build_document()
    if OUT.exists():
        OUT.unlink()
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
