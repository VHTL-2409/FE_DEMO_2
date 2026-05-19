from __future__ import annotations

import re
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
DIAGRAM_ROOT = ROOT / "docs" / "uml-diagrams"
PLANTUML_ROOT = DIAGRAM_ROOT / "plantuml_2"
MANIFEST = PLANTUML_ROOT / "_images" / "FIGURE_NAMES.md"
RENDER_SCRIPT = DIAGRAM_ROOT / "render-plantuml.ps1"
PLANTUML_JAR = ROOT / "plantuml-1.2026.2.jar"
OUTPUT_DOCX = ROOT / "baocao" / "Phu-luc-hinh-anh-plantuml_2-A4-split.docx"

A4_PORTRAIT = (21.0, 29.7)
A4_LANDSCAPE = (29.7, 21.0)
MARGIN_CM = 1.0
CAPTION_SPACE_IN = 0.45
RENDER_SCALE = 2


def read_manifest() -> list[dict[str, str]]:
    pattern = re.compile(r"^\|\s*(\d+)\s*\|\s*`([^`]+)`\s*\|\s*(.*?)\s*\|\s*`([^`]+)`\s*\|")
    rows: list[dict[str, str]] = []
    for line in MANIFEST.read_text(encoding="utf-8").splitlines():
        match = pattern.match(line)
        if not match:
            continue
        rows.append({
            "index": match.group(1),
            "image": match.group(2),
            "title": match.group(3).strip(),
            "source": match.group(4),
        })
    if not rows:
        raise RuntimeError(f"No figure rows found in {MANIFEST}")
    return rows


def add_scale_directive(text: str) -> str:
    if re.search(r"(?m)^scale\s+\d+", text):
        return text

    lines = text.splitlines()
    insert_at = 1
    for index, line in enumerate(lines[:20]):
        stripped = line.strip()
        if stripped.startswith("title "):
            insert_at = index + 1
            break
        if stripped.startswith("!theme"):
            insert_at = index + 1

    lines.insert(insert_at, f"scale {RENDER_SCALE}")
    return "\n".join(lines) + "\n"


def prepare_scaled_sources(rows: list[dict[str, str]], input_root: Path) -> None:
    for row in rows:
        source = ROOT / row["source"]
        if not source.exists():
            raise FileNotFoundError(source)
        rel_source = source.relative_to(PLANTUML_ROOT)
        target = input_root / rel_source
        target.parent.mkdir(parents=True, exist_ok=True)
        text = source.read_text(encoding="utf-8-sig", errors="ignore")
        target.write_text(add_scale_directive(text), encoding="utf-8", newline="\n")


def render_scaled_images(input_root: Path, output_root: Path) -> None:
    if not PLANTUML_JAR.exists():
        raise FileNotFoundError(PLANTUML_JAR)

    input_arg = input_root.relative_to(DIAGRAM_ROOT).as_posix()
    output_arg = output_root.relative_to(DIAGRAM_ROOT).as_posix()
    command = [
        "powershell",
        "-ExecutionPolicy",
        "Bypass",
        "-File",
        str(RENDER_SCRIPT),
        "-InputDir",
        input_arg,
        "-OutputDir",
        output_arg,
        "-Format",
        "png",
        "-PlantUmlJar",
        str(PLANTUML_JAR),
    ]
    subprocess.run(command, cwd=ROOT, check=True)


def set_style_defaults(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)

    caption = document.styles["Caption"]
    caption.font.name = "Times New Roman"
    caption.font.size = Pt(10)
    caption.font.italic = False


def set_section(section, orientation: WD_ORIENT) -> None:
    if orientation == WD_ORIENT.LANDSCAPE:
        width_cm, height_cm = A4_LANDSCAPE
    else:
        width_cm, height_cm = A4_PORTRAIT

    section.orientation = orientation
    section.page_width = Cm(width_cm)
    section.page_height = Cm(height_cm)
    section.top_margin = Cm(MARGIN_CM)
    section.bottom_margin = Cm(MARGIN_CM)
    section.left_margin = Cm(MARGIN_CM)
    section.right_margin = Cm(MARGIN_CM)
    section.header_distance = Cm(0.5)
    section.footer_distance = Cm(0.5)


def available_box(orientation: WD_ORIENT) -> tuple[float, float]:
    if orientation == WD_ORIENT.LANDSCAPE:
        width_cm, height_cm = A4_LANDSCAPE
    else:
        width_cm, height_cm = A4_PORTRAIT

    max_width_in = (width_cm - MARGIN_CM * 2) / 2.54
    max_height_in = (height_cm - MARGIN_CM * 2) / 2.54 - CAPTION_SPACE_IN
    return max_width_in, max_height_in


def fit_size(width_px: int, height_px: int, orientation: WD_ORIENT) -> tuple[float, float, float]:
    max_width_in, max_height_in = available_box(orientation)
    scale = min(max_width_in / width_px, max_height_in / height_px)
    width_in = width_px * scale
    height_in = height_px * scale
    return width_in, height_in, scale


def choose_orientation(width_px: int, height_px: int) -> WD_ORIENT:
    portrait = fit_size(width_px, height_px, WD_ORIENT.PORTRAIT)
    landscape = fit_size(width_px, height_px, WD_ORIENT.LANDSCAPE)
    portrait_area = portrait[0] * portrait[1]
    landscape_area = landscape[0] * landscape[1]
    return WD_ORIENT.LANDSCAPE if landscape_area > portrait_area else WD_ORIENT.PORTRAIT


def clean_caption(title: str) -> str:
    return re.sub(r"^Hình\s*-\s*", "", title).strip()


def add_figure_page(document: Document, section, image_path: Path, figure_no: int, title: str) -> None:
    with Image.open(image_path) as image:
        width_px, height_px = image.size

    orientation = choose_orientation(width_px, height_px)
    set_section(section, orientation)
    width_in, _, _ = fit_size(width_px, height_px, orientation)

    image_paragraph = document.add_paragraph()
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.paragraph_format.space_before = Pt(0)
    image_paragraph.paragraph_format.space_after = Pt(2)
    image_paragraph.paragraph_format.keep_together = True
    run = image_paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width_in))

    caption = document.add_paragraph(style="Caption")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(0)
    caption.paragraph_format.space_after = Pt(0)
    caption.paragraph_format.line_spacing = 1.0
    caption.paragraph_format.keep_together = True
    run = caption.add_run(f"Hình 1.{figure_no}. {clean_caption(title)}")
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    run.font.bold = False


def image_path_for_row(output_root: Path, row: dict[str, str]) -> Path:
    original = ROOT / row["image"]
    rel_parent = original.relative_to(PLANTUML_ROOT / "_images").parent
    return output_root / rel_parent / original.name


def build_document(rows: list[dict[str, str]], image_root: Path) -> None:
    document = Document()
    set_style_defaults(document)
    document.core_properties.title = "Phụ lục hình ảnh PlantUML 2 A4"
    document.core_properties.subject = "Biểu đồ hoạt động và biểu đồ tuần tự FE_DEMO"

    section = document.sections[0]
    for index, row in enumerate(rows, start=1):
        if index > 1:
            section = document.add_section(WD_SECTION_START.NEW_PAGE)
        image_path = image_path_for_row(image_root, row)
        if not image_path.exists():
            raise FileNotFoundError(image_path)
        add_figure_page(document, section, image_path, index, row["title"])

    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_DOCX)


def main() -> None:
    rows = read_manifest()
    with tempfile.TemporaryDirectory(prefix=".tmp-plantuml2-a4-", dir=DIAGRAM_ROOT) as temp_dir:
        temp_root = Path(temp_dir)
        input_root = temp_root / "input"
        output_root = temp_root / "output"
        prepare_scaled_sources(rows, input_root)
        render_scaled_images(input_root, output_root)
        build_document(rows, output_root)
    print(f"created {OUTPUT_DOCX}")
    print(f"figures {len(rows)}")


if __name__ == "__main__":
    main()
