from __future__ import annotations

import json
import re
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
BAOCao_DIR = ROOT / "baocao"
SOURCE_DOCX = BAOCao_DIR / "baocao_long_test_format_anh.docx"
OUTPUT_DOCX = BAOCao_DIR / "Bao-cao-tom-tat-do-an-thi-truc-tuyen-giam-sat-gian-lan-thong-minh.docx"


def find_template() -> Path:
    for path in BAOCao_DIR.iterdir():
        if path.suffix.lower() == ".docx" and path.name.startswith("4. "):
            return path
    raise FileNotFoundError("Khong tim thay file mau 4. Mau Bao cao tom tat.docx")


def clear_body(document: Document) -> None:
    body = document._body._element
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def set_fonts(document: Document) -> None:
    styles = document.styles
    for name, size, bold, italic in [
        ("Normal", 13, False, False),
        ("Heading 1", 16, True, False),
        ("Heading 2", 14, True, False),
        ("Heading 3", 13, True, False),
        ("Heading 4", 12, True, False),
        ("Caption", 12, False, True),
        ("Title", 20, True, False),
    ]:
        if name not in styles:
            continue
        style = styles[name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.italic = italic


def set_run_font(run, size: int = 13, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic


def add_paragraph(
    document: Document,
    text: str = "",
    *,
    style: str = "Normal",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_cm: float | None = 0.75,
    size: int = 13,
    bold: bool = False,
    italic: bool = False,
) -> object:
    paragraph = document.add_paragraph(style=style)
    paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.line_spacing = 1.5
    fmt.space_after = Pt(6)
    fmt.space_before = Pt(0)
    if first_line_cm is not None:
        fmt.first_line_indent = Pt(first_line_cm * 28.3465)
    if text:
        run = paragraph.add_run(text)
        set_run_font(run, size=size, bold=bold, italic=italic)
    return paragraph


def add_heading(document: Document, text: str, level: int, *, page_break: bool = False) -> None:
    if page_break:
        document.add_page_break()
    paragraph = document.add_paragraph(style=f"Heading {level}")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_run_font(run, size=16 if level == 1 else 14 if level == 2 else 13, bold=True)


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr[idx].text = header
        for p in hdr[idx].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                set_run_font(r, size=11, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            for p in cells[idx].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    set_run_font(r, size=11)
    document.add_paragraph()


def extract_usecase_figures(source: Path) -> list[dict[str, object]]:
    doc = Document(str(source))
    paras = doc.paragraphs
    figures: list[dict[str, object]] = []
    for i, para in enumerate(paras):
        blips = para._p.xpath('.//a:blip')
        if not blips:
            continue
        caption = ""
        for j in range(i + 1, min(i + 4, len(paras))):
            text = paras[j].text.strip().replace("\t", " ")
            if text:
                caption = text
                break
        lower = caption.lower()
        if ("use case" not in lower and "usecase" not in lower) or "hoạt động" in lower or "tuần tự" in lower:
            continue
        caption = re.sub(r"^Hình\s+\d+\.\d+\.\s*", "", caption)
        caption = re.sub(r"(?i)usecase", "Use Case", caption)
        caption = re.sub(r"(?i)use case", "Use Case", caption)
        rid = blips[0].get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
        blob = doc.part.related_parts[rid].blob
        figures.append({"caption": caption, "blob": blob})
    return figures


def extract_figures_by_prefix(
    source: Path,
    *,
    prefix: str,
    include: tuple[str, ...] = (),
    exclude: tuple[str, ...] = (),
) -> list[dict[str, object]]:
    doc = Document(str(source))
    paras = doc.paragraphs
    figures: list[dict[str, object]] = []
    for i, para in enumerate(paras):
        blips = para._p.xpath('.//a:blip')
        if not blips:
            continue
        caption = ""
        for j in range(i + 1, min(i + 4, len(paras))):
            text = paras[j].text.strip().replace("\t", " ")
            if text:
                caption = text
                break
        lower = caption.lower()
        if include and not any(token.lower() in lower for token in include):
            continue
        if exclude and any(token.lower() in lower for token in exclude):
            continue
        if not lower.startswith(prefix.lower()):
            continue
        caption = re.sub(r"^Hình\s+\d+\.\d+\.\s*", "", caption)
        caption = re.sub(r"(?i)usecase", "Use Case", caption)
        caption = re.sub(r"(?i)use case", "Use Case", caption)
        rid = blips[0].get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
        blob = doc.part.related_parts[rid].blob
        figures.append({"caption": caption, "blob": blob})
    return figures


def add_figure_chapter(document: Document, chapter_no: int, number: int, caption: str, blob: bytes, *, width_in: float = 6.2) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run()
    run.add_picture(BytesIO(blob), width=Inches(width_in))

    caption_p = document.add_paragraph(style="Caption")
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_p.paragraph_format.keep_together = True
    caption_p.paragraph_format.space_after = Pt(6)
    run = caption_p.add_run(f"Hình {chapter_no}.{number}. {caption}")
    set_run_font(run, size=12, italic=True)


def add_figure(document: Document, number: int, caption: str, blob: bytes, *, width_in: float = 6.2) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run()
    run.add_picture(BytesIO(blob), width=Inches(width_in))

    caption_p = document.add_paragraph(style="Caption")
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_p.paragraph_format.keep_together = True
    caption_p.paragraph_format.space_after = Pt(6)
    run = caption_p.add_run(f"Hình 2.{number}. {caption}")
    set_run_font(run, size=12, italic=True)


def build_document(document: Document, figures: list[dict[str, object]], versions: dict[str, str]) -> None:
    set_fonts(document)

    # Title page
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(80)
    title.paragraph_format.space_after = Pt(12)
    run = title.add_run("BÁO CÁO TÓM TẮT ĐỒ ÁN")
    set_run_font(run, size=20, bold=True)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(12)
    run = subtitle.add_run("Hệ thống thi trực tuyến với cơ chế giám sát và phát hiện gian lận thông minh")
    set_run_font(run, size=16, bold=True)

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(4)
    run = meta.add_run("FE_DEMO")
    set_run_font(run, size=13, bold=True)

    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_after = Pt(0)
    run = note.add_run("Phiên bản tóm tắt, dùng ảnh use case trích từ báo cáo dài")
    set_run_font(run, size=12, italic=True)

    document.add_page_break()

    add_heading(document, "MỞ ĐẦU", 1)
    add_heading(document, "1. Lý do chọn đề tài", 2)
    add_paragraph(
        document,
        "Thi trực tuyến ngày càng phổ biến, nhưng bài toán giám sát gian lận vẫn là điểm nghẽn lớn. Nếu chỉ dừng ở phát đề và nộp bài, hệ thống sẽ thiếu cơ chế theo dõi hành vi bất thường, khó phát hiện đổi tab, mất focus, copy-paste hay các tình huống gian lận phức tạp trong quá trình làm bài.",
    )
    add_paragraph(
        document,
        "Dự án FE_DEMO xử lý đúng khoảng trống đó bằng kiến trúc nhiều tầng: frontend Vue 3, backend Spring Boot, dịch vụ AI FastAPI và cơ sở dữ liệu PostgreSQL. Cách làm này cho phép hệ thống vừa hỗ trợ thi trực tuyến, vừa giám sát realtime, vừa lưu vết và phân tích rủi ro một cách rõ ràng.",
    )
    add_heading(document, "2. Mục tiêu và phạm vi", 2)
    add_paragraph(
        document,
        "Mục tiêu là xây dựng luồng thi hoàn chỉnh cho ba vai trò chính: sinh viên, giáo viên và quản trị viên. Hệ thống cần đáp ứng đăng nhập, tạo và quản lý đề thi, tham gia thi, giám sát realtime, chấm điểm rủi ro, quản lý lớp học và dashboard quản trị.",
    )
    add_paragraph(
        document,
        "Phạm vi tóm tắt tập trung vào phần phân tích chức năng và các use case chính đã triển khai trong đồ án. Báo cáo không đi sâu vào mọi màn hình chi tiết, mà nhấn vào cấu trúc nghiệp vụ, các vai trò, và luồng chính của hệ thống.",
    )
    add_heading(document, "3. Phương pháp thực hiện", 2)
    add_paragraph(
        document,
        "Cách làm dựa trên phân tích yêu cầu từ repo, đối chiếu mã nguồn frontend/backend/AI service, rồi gom lại thành nhóm chức năng và use case. Phần minh họa lấy trực tiếp từ file báo cáo dài của đồ án để đảm bảo hình ảnh khớp chính hệ thống.",
    )
    add_paragraph(
        document,
        "Bố cục báo cáo gồm tổng quan công nghệ, phân tích chức năng theo vai trò, các sơ đồ use case chính và phần kết quả triển khai. Đây là bản rút gọn, phù hợp để nộp kèm hoặc dùng làm bản tóm tắt cho đồ án.",
    )
    add_heading(document, "4. Cấu trúc báo cáo", 2)
    add_paragraph(document, "Chương I: Tổng quan công nghệ và môi trường phát triển.", first_line_cm=0.0)
    add_paragraph(document, "Chương II: Phân tích hệ thống và sơ đồ use case.", first_line_cm=0.0)
    add_paragraph(document, "Chương III: Kết quả triển khai, kiểm thử và đánh giá.", first_line_cm=0.0)

    add_heading(document, "CHƯƠNG I: TỔNG QUAN CÔNG NGHỆ VÀ MÔI TRƯỜNG PHÁT TRIỂN", 1, page_break=True)
    add_heading(document, "1.1. Tổng quan hệ thống", 2)
    add_paragraph(
        document,
        "FE_DEMO là hệ thống thi trực tuyến có giám sát thông minh, được thiết kế để hỗ trợ cả thao tác nghiệp vụ lẫn theo dõi hành vi bất thường của thí sinh. Ở mức tổng thể, hệ thống chia rõ giao diện người dùng, lớp nghiệp vụ, lớp AI và lớp dữ liệu để có thể mở rộng độc lập.",
    )
    add_paragraph(
        document,
        "Điểm nổi bật là luồng realtime qua WebSocket và lớp phân tích AI cho camera, OCR và đánh giá rủi ro. Nhờ vậy giáo viên có thể theo dõi diễn biến phòng thi, còn hệ thống có đủ dữ liệu để lưu nhật ký, cảnh báo và tổng hợp sau thi.",
    )

    add_heading(document, "1.2. Frontend", 2)
    add_paragraph(
        document,
        f"Frontend nằm trong thư mục `demo/`, dùng Vue {versions['vue']} với Vite {versions['vite']} làm nền chính. Các thành phần như Pinia {versions['pinia']}, Vue Router {versions['vue_router']}, PrimeVue {versions['primevue']}, ECharts {versions['echarts']} và xlsx {versions['xlsx']} cho thấy giao diện được thiết kế theo hướng giàu tương tác và có khả năng xử lý dữ liệu tốt.",
    )
    add_table(
        document,
        ["Thành phần", "Phiên bản", "Vai trò"],
        [
            ["Vue", versions["vue"], "Framework giao diện"],
            ["Vite", versions["vite"], "Dev server và bundler"],
            ["Pinia", versions["pinia"], "Quản lý state"],
            ["Vue Router", versions["vue_router"], "Điều hướng"],
            ["PrimeVue", versions["primevue"], "Component UI"],
            ["ECharts", versions["echarts"], "Biểu đồ và dashboard"],
            ["xlsx", versions["xlsx"], "Import/export dữ liệu"],
        ],
    )

    add_heading(document, "1.3. Backend", 2)
    add_paragraph(
        document,
        f"Backend nằm trong thư mục `BE/`, xây trên Spring Boot {versions['spring_boot']} và Java {versions['java']}. Lớp backend xử lý xác thực, phân quyền, đề thi, lớp học, phiên thi, giám sát và các API admin.",
    )
    add_table(
        document,
        ["Thành phần", "Vai trò"],
        [
            ["Spring Security", "Xác thực và phân quyền"],
            ["Spring Web", "REST API"],
            ["Spring WebSocket", "Kết nối realtime"],
            ["Spring Data JPA", "Truy cập dữ liệu"],
            ["JWT", "Token đăng nhập và refresh"],
            ["Mail", "Xác thực email và quên mật khẩu"],
        ],
    )

    add_heading(document, "1.4. AI Service", 2)
    add_paragraph(
        document,
        "AI service tách riêng thành microservice FastAPI để xử lý tác vụ nặng như OCR, phân tích hình ảnh camera, chấm tự luận, sinh câu hỏi và chat hỗ trợ. Cách tách riêng này giúp backend gọn hơn và AI có thể scale độc lập.",
    )
    add_table(
        document,
        ["Thành phần", "Vai trò"],
        [
            ["FastAPI", "Khung API cho AI service"],
            ["uvicorn", "ASGI server"],
            ["OpenCV", "Xử lý ảnh và video"],
            ["pytesseract", "Nhận dạng ký tự"],
            ["Pillow", "Xử lý ảnh"],
            ["NumPy", "Tính toán số học"],
            ["openai", "Kết nối mô hình LLM"],
            ["slowapi", "Giới hạn tần suất"],
        ],
    )

    add_heading(document, "1.5. Triển khai", 2)
    add_paragraph(
        document,
        "Toàn bộ hệ thống được triển khai theo kiểu container hóa, có thể chạy qua Docker Compose với các service tách biệt cho frontend, backend, AI service, parser và PostgreSQL. Cấu trúc này phù hợp với môi trường phát triển lẫn demo nội bộ.",
    )
    add_paragraph(
        document,
        "Mô hình nhiều service giúp debug từng khối riêng, đồng thời giữ luồng dữ liệu rõ ràng: trình duyệt gửi yêu cầu đến backend, backend gọi AI khi cần, còn dữ liệu trạng thái và log được lưu trong cơ sở dữ liệu trung tâm.",
    )

    add_heading(document, "CHƯƠNG II: PHÂN TÍCH HỆ THỐNG VÀ USE CASE", 1, page_break=True)
    add_heading(document, "2.1. Tác nhân hệ thống", 2)
    add_paragraph(
        document,
        "Hệ thống có ba vai trò chính là sinh viên, giáo viên và quản trị viên. Ngoài ra còn có các thành phần kỹ thuật như backend, AI service và WebSocket server để bảo đảm việc truyền dữ liệu, giám sát và phân tích diễn ra liền mạch.",
    )
    add_table(
        document,
        ["Tác nhân", "Nhiệm vụ chính", "Phạm vi tương tác"],
        [
            ["Sinh viên", "Thi, xem kết quả, tham gia lớp", "Khu vực sinh viên"],
            ["Giáo viên", "Tạo đề, giám sát, xem thống kê", "Khu vực giáo viên"],
            ["Quản trị viên", "Quản lý user và dashboard", "Khu vực admin"],
            ["AI service", "OCR, phân tích camera, chấm điểm", "API nội bộ"],
            ["Backend", "Xử lý nghiệp vụ và lưu dữ liệu", "API trung tâm"],
        ],
    )

    groups = [
        ("2.2. Use case tổng quát và xác thực", [
            "Nhóm use case này mô tả lớp truy cập nền tảng: từ tổng quan hệ thống cho tới đăng nhập và khôi phục mật khẩu. Đây là lớp nền để mọi vai trò đi vào các chức năng nghiệp vụ phía sau.",
            0, 1, 2
        ]),
        ("2.3. Use case sinh viên", [
            "Sinh viên là nhóm người dùng chính trong luồng làm bài. Các sơ đồ dưới đây cho thấy cách vào phòng thi, làm bài, quản lý lớp học và xem kết quả sau thi.",
            3, 4, 5
        ]),
        ("2.4. Use case giáo viên", [
            "Giáo viên quản lý gần như toàn bộ vòng đời kỳ thi: dashboard, lớp học, tạo đề thi, giám sát và điều phối các kỳ thi đang mở.",
            6, 7, 8, 9
        ]),
        ("2.5. Use case quản trị", [
            "Khối quản trị tập trung vào dashboard hệ thống và quản lý tài khoản người dùng. Đây là phần bảo đảm vận hành và phân quyền ở mức nền tảng.",
            10, 11
        ]),
    ]

    fig_no = 1
    for heading, payload in groups:
        add_heading(document, heading, 2)
        add_paragraph(document, payload[0])
        for idx in payload[1:]:
            figure = figures[idx]
            add_figure(document, fig_no, str(figure["caption"]), figure["blob"])
            fig_no += 1

    add_heading(document, "CHƯƠNG III: KẾT QUẢ TRIỂN KHAI VÀ ĐÁNH GIÁ", 1, page_break=True)
    add_heading(document, "3.1. Kết quả triển khai", 2)
    add_paragraph(
        document,
        "Các module chính của đồ án đã hình thành thành chuỗi nghiệp vụ khép kín: đăng nhập, tạo đề, vào phòng thi, giám sát realtime, lưu dấu vết vi phạm, xem kết quả và quản trị tài khoản. Giao diện đã được tách theo từng vai trò nên dễ kiểm thử và dễ mở rộng.",
    )
    add_paragraph(
        document,
        "Phần AI service giúp dự án vượt khỏi mức một hệ thống thi online thông thường. Nó cho phép kiểm tra camera, phân tích trạng thái bất thường và tạo thêm lớp dữ liệu cho dashboard giáo viên cũng như dashboard admin.",
    )

    add_heading(document, "3.2. Kiểm thử nhanh", 2)
    add_table(
        document,
        ["Kịch bản", "Kết quả"],
        [
            ["Đăng nhập và phân quyền", "Đạt"],
            ["Sinh viên vào phòng thi", "Đạt"],
            ["Lưu nháp và nộp bài", "Đạt"],
            ["Giáo viên theo dõi realtime", "Đạt"],
            ["Admin xem dashboard và user", "Đạt"],
        ],
    )

    add_heading(document, "3.3. Đánh giá", 2)
    add_paragraph(
        document,
        "Điểm mạnh lớn nhất của đồ án là tách lớp rõ, có realtime, có AI và có lộ trình nghiệp vụ đủ cho cả ba vai trò. Báo cáo tóm tắt này chỉ giữ phần cốt lõi, nhưng vẫn phản ánh được cấu trúc thật của hệ thống.",
    )
    add_paragraph(
        document,
        "Hướng phát triển tiếp theo có thể là tăng độ chính xác của mô hình AI, hoàn thiện audit log, bổ sung báo cáo thống kê sâu hơn và tối ưu trải nghiệm phòng thi trên nhiều thiết bị.",
    )

    add_heading(document, "KẾT LUẬN", 1, page_break=True)
    add_paragraph(
        document,
        "Đồ án FE_DEMO đã đạt mục tiêu chính là xây dựng hệ thống thi trực tuyến có giám sát, có phân vai rõ ràng và có lớp phân tích gian lận thông minh. Bộ use case trong báo cáo cho thấy hệ thống bao phủ được cả người học, người dạy và quản trị viên, đủ làm nền cho bản triển khai thực tế và cho các vòng mở rộng sau này.",
    )


def main() -> None:
    template = find_template()
    document = Document(str(template))
    clear_body(document)

    figures = extract_usecase_figures(SOURCE_DOCX)
    if len(figures) < 12:
        raise RuntimeError(f"Khong du use case images: chi tim thay {len(figures)}")

    versions = {
        "vue": "^3.5.13",
        "vite": "^6.2.0",
        "pinia": "^3.0.4",
        "vue_router": "^4.5.1",
        "primevue": "^4.5.4",
        "echarts": "^6.0.0",
        "xlsx": "^0.18.5",
        "spring_boot": "4.0.3",
        "java": "17",
    }

    build_document(document, figures, versions)
    document.save(str(OUTPUT_DOCX))
    print(f"Saved: {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
