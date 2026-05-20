from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Phan_tich_code_core_phat_hien_gian_lan_giam_sat_AI.docx"


def set_font(run, name: str = "Arial", size: float = 10, bold: bool = False, color: str | None = None) -> None:
    run.bold = bold
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def cell_text(cell, text: str, bold: bool = False, fill: str | None = None) -> None:
    cell.text = ""
    run = cell.paragraphs[0].add_run(str(text))
    set_font(run, size=9, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    if fill:
        shade(cell, fill)


def table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell_text(t.rows[0].cells[i], h, bold=True, fill="D9EAF7")
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cell_text(cells[i], value)
    if widths:
        for row in t.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    return t


def para(doc: Document, text: str = "", bold: bool = False, size: float = 10.5) -> None:
    p = doc.add_paragraph()
    set_font(p.add_run(text), size=size, bold=bold)


def bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(1)
    set_font(p.add_run(text), size=10)


def code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_font(run, name="Consolas", size=8, color="303030")


def snippet(path: str, start: int, end: int, keep: list[str] | None = None, max_lines: int = 24) -> str:
    lines = (ROOT / path).read_text(encoding="utf-8", errors="replace").splitlines()
    selected: list[str] = []
    in_comment = False
    for number in range(start, min(end, len(lines)) + 1):
        raw = lines[number - 1].rstrip()
        s = raw.strip()
        if not s:
            continue
        if s.startswith("#") or s.startswith("//"):
            continue
        if s.startswith("/*") or s.startswith("/**"):
            in_comment = True
            continue
        if in_comment:
            if "*/" in s:
                in_comment = False
            continue
        if s.startswith("*"):
            continue
        if keep and not any(token in raw for token in keep):
            continue
        selected.append(f"{number}: {raw}")
        if len(selected) >= max_lines:
            break
    return "\n".join(selected)


def setup_doc() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    doc.styles["Normal"].font.size = Pt(10.5)
    doc.styles["Heading 1"].font.size = Pt(16)
    doc.styles["Heading 2"].font.size = Pt(13)
    doc.styles["Heading 3"].font.size = Pt(11.5)
    return doc


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("PHÂN TÍCH CODE CORE\nPHÁT HIỆN GIAN LẬN - GIÁM SÁT - AI"), size=20, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("Bản dễ theo dõi: phân tích theo luồng nghiệp vụ và code thật trích từ dự án"), size=12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run(f"Ngày tạo: {datetime.now().strftime('%d/%m/%Y %H:%M')}"), size=10)
    doc.add_paragraph()
    table(doc, ["Cách đọc tài liệu", "Ý nghĩa"], [
        ["Luồng 1", "Camera AI: từ frame ảnh đến FraudSignal."],
        ["Luồng 2", "Sự kiện trình duyệt: từ tab switch/copy paste đến điểm rủi ro."],
        ["Luồng 3", "Risk scoring: từ danh sách FraudSignal đến riskScore, riskLevel, flag."],
        ["Luồng 4", "Dashboard realtime: từ backend notification đến màn hình giáo viên."],
    ], widths=[1.6, 5.8])
    doc.add_page_break()


def add_overview(doc: Document) -> None:
    doc.add_heading("1. Tổng Quan Dễ Nhìn", level=1)
    para(doc, "Thay vì đọc từng file rời rạc, nên nhìn hệ thống theo 4 luồng dữ liệu chính:")
    para(
        doc,
        "Lưu ý: các đoạn code trong tài liệu là code thật được trích trực tiếp từ file nguồn của dự án, có giữ số dòng ở đầu mỗi dòng. "
        "Tài liệu chỉ rút gọn bớt comment/dòng không liên quan để dễ theo dõi, không viết lại thành code minh họa.",
        bold=True,
    )
    table(doc, ["Luồng", "Input", "Hàm core", "Output", "Ai sử dụng"], [
        ["1. Camera AI", "Ảnh base64 + metadata camera/mic", "AiAssistService.analyzeFrame -> ProctorAnalyzer.analyze_frame", "AI signals: NO_CAMERA, MULTIPLE_FACES, GAZE_OFF_SCREEN...", "Backend risk engine, dashboard camera"],
        ["2. Browser event", "TAB_SWITCH, WINDOW_BLUR, COPY_PASTE, DEVTOOLS_OPEN...", "ExamEventService.ingestBatchInternal", "ExamEvent + FraudSignal", "RiskScoringService"],
        ["3. Risk scoring", "Danh sách FraudSignal của attempt", "RiskScoringService.recomputeRisk", "riskScore, riskLevel, ProctorFlag, RiskScoreLog", "Teacher dashboard, timeline"],
        ["4. Realtime UI", "WebSocket payload từ backend", "useExamMonitoring.applyRealtimeEvent", "Card/alert/timeline cập nhật tức thì", "Giáo viên giám sát"],
    ], widths=[1.5, 2.1, 2.6, 2.3, 1.8])
    code(doc, "Student camera/event\n   -> Backend nhận dữ liệu\n   -> AI hoặc FraudSignalService chuẩn hóa signal\n   -> RiskScoringService tính điểm\n   -> WebSocket cập nhật dashboard giáo viên")


def add_flow_camera(doc: Document) -> None:
    doc.add_heading("2. Luồng 1 - Camera AI", level=1)
    para(doc, "Mục tiêu: từ một frame camera của học sinh, hệ thống xác định có dấu hiệu gian lận thị giác hay không.")
    table(doc, ["Bước", "File/Hàm", "Vai trò"], [
        ["1", "BE/AiAssistService.analyzeFrame", "Nhận frame từ frontend, validate, publish pending, gọi AI service."],
        ["2", "ai-service/ProctorAnalyzer.analyze_frame", "Phân tích ảnh: mặt, mắt, gaze, ánh sáng, blur, spoofing."],
        ["3", "BE/AiAssistService.safeBridgeAiSignals", "Chuyển AI signals thành FraudSignal trong DB."],
        ["4", "RiskScoringService.recomputeRisk", "Tính lại điểm rủi ro sau khi có signal AI."],
    ], widths=[0.7, 3.0, 5.2])

    doc.add_heading("2.1 Code thật: backend nhận frame", level=2)
    code(doc, snippet("BE/src/main/java/com/example/demo/service/AiAssistService.java", 128, 176, keep=[
        "analyzeFrame", "validateFrameRequest", "buildFrameFallbackResponse", "publishCameraFrameInternal",
        "postJson", "collectAiCameraWarningSignals", "clusterAiCameraSignalsForRecording",
        "safeBridgeAiSignals", "safeRecordAiCameraWarnings", "AI_UNAVAILABLE"
    ], max_lines=22))
    table(doc, ["Dòng/ý chính", "Giải thích dễ hiểu"], [
        ["validateFrameRequest(request)", "Chặn request thiếu attemptId, thiếu ảnh, ảnh quá lớn."],
        ["publishCameraFrameInternal(..., \"ai_pending\")", "Giáo viên vẫn thấy frame đang được nhận, dù AI chưa phân tích xong."],
        ["postJson(\"/proctor/analyze/frame\", ...)", "Gọi Python ai-service để phân tích ảnh."],
        ["clusterAiCameraSignalsForRecording(...)", "Gom cụm signal camera để tránh mỗi frame đều ghi một cảnh báo giống nhau."],
        ["safeBridgeAiSignals(...)", "Ghi signal AI thành FraudSignal để risk engine xử lý."],
        ["catch Exception -> AI_UNAVAILABLE", "Nếu AI service lỗi, backend vẫn trả trạng thái rõ ràng thay vì làm hỏng luồng camera."],
    ], widths=[2.8, 5.6])

    doc.add_heading("2.2 Code thật: AI phân tích frame", level=2)
    code(doc, snippet("ai-service/app/proctor.py", 318, 430, keep=[
        "def analyze_frame", "_resolve_camera_state", "_build_camera_off_response", "_decode_image",
        "_detect_face_locations", "_track_eyes", "average_brightness", "variance",
        "_detect_spoofing_deep", "_detect_occlusion", "_analyze_face_position", "_analyze_gaze",
        "_build_quality_gate", "_update_tracking_state", "signals ="
    ], max_lines=28))
    table(doc, ["Nhóm xử lý", "Code tương ứng", "Ý nghĩa"], [
        ["Camera/mic", "_resolve_camera_state, _resolve_mic_state", "Nếu camera/mic tắt, sinh signal NO_CAMERA/NO_MIC."],
        ["Ảnh", "_decode_image -> RGB/gray", "Chuẩn bị ảnh cho OpenCV/MediaPipe."],
        ["Mặt", "_detect_face_locations", "Đếm số mặt: 0 mặt, 1 mặt, nhiều mặt."],
        ["Mắt", "_track_eyes", "Xác định mắt, pupil, eye state."],
        ["Chất lượng ảnh", "average_brightness, variance", "Phát hiện ánh sáng yếu/quá sáng, ảnh mờ."],
        ["Gaze", "_analyze_gaze + _update_tracking_state", "Chỉ kết luận nhìn ra ngoài màn hình khi đủ nhiều frame/thời gian."],
        ["Output", "signals, warnings, visual_overlay, diagnostics", "Trả cả kết quả và bằng chứng để dashboard hiển thị."],
    ], widths=[1.6, 2.6, 4.3])

    doc.add_heading("2.3 Code thật: sinh signal camera", level=2)
    code(doc, snippet("ai-service/app/proctor.py", 390, 575, keep=[
        "signals.append", "FACE_NOT_DETECTED", "MULTIPLE_FACES", "FACE_SPOOFING_SUSPECTED",
        "EYES_NOT_DETECTED", "GAZE_OFF_SCREEN", "VERY_LOW_LIGHTING", "VERY_BLURRY_FRAME",
        "quality_gate", "stability"
    ], max_lines=30))
    table(doc, ["Signal", "Điều kiện chính", "Mức độ"], [
        ["FACE_NOT_DETECTED", "Không thấy mặt và trạng thái đã ổn định qua nhiều frame", "HIGH"],
        ["MULTIPLE_FACES", "Có hơn 1 khuôn mặt trong frame", "CRITICAL"],
        ["FACE_SPOOFING_SUSPECTED", "Heuristic/DL nghi ảnh giả, replay, deepfake", "CRITICAL"],
        ["EYES_NOT_DETECTED", "Có 1 mặt nhưng không thấy mắt, quality gate đủ tin cậy", "MEDIUM"],
        ["GAZE_OFF_SCREEN", "Hướng nhìn lệch khỏi màn hình đủ lâu, gaze_valid=true", "HIGH"],
        ["VERY_LOW_LIGHTING", "Độ sáng quá thấp", "HIGH"],
        ["VERY_BLURRY_FRAME", "Variance thấp, ảnh quá mờ", "HIGH"],
    ], widths=[2.1, 4.7, 1.2])


def add_flow_browser(doc: Document) -> None:
    doc.add_heading("3. Luồng 2 - Sự Kiện Trình Duyệt", level=1)
    para(doc, "Mục tiêu: các hành vi như chuyển tab, mất focus, copy/paste, mở DevTools được gửi từ frontend về backend và chuẩn hóa thành FraudSignal.")
    doc.add_heading("3.1 Code thật: ingest batch", level=2)
    code(doc, snippet("BE/src/main/java/com/example/demo/service/ExamEventService.java", 95, 148, keep=[
        "ingestBatchInternal", "requireAttempt", "ensureStudentOwnsAttempt", "ensureAttemptActive",
        "normalizeFingerprint", "identityAnomalyService", "existsByAttemptAndSequenceNo",
        "isDuplicateBurst", "isSignalRateLimited", "saveExamEvent", "saveLegacyMonitoringEvent",
        "recordFromEvent", "recomputeRisk"
    ], max_lines=30))
    table(doc, ["Cơ chế", "Vì sao cần"], [
        ["sequenceNo", "Tránh một batch gửi lại làm nhân đôi event."],
        ["isDuplicateBurst", "Chặn event giống nhau trong vài giây, ví dụ blur/tab switch bị spam."],
        ["isSignalRateLimited", "Giới hạn số signal cùng loại trong cửa sổ thời gian."],
        ["saveExamEvent", "Lưu lịch sử event thô để audit."],
        ["FraudSignalService.recordFromEvent", "Chuyển event thành signal chuẩn cho risk engine."],
        ["riskScoringService.recomputeRisk", "Mỗi batch hợp lệ xong thì cập nhật riskScore."],
    ], widths=[2.3, 5.8])

    doc.add_heading("3.2 Code thật: chuẩn hóa FraudSignal", level=2)
    code(doc, snippet("BE/src/main/java/com/example/demo/service/FraudSignalService.java", 43, 108, keep=[
        "recordFromEvent", "recordServerSignal", "descriptorFor", "normalizeConfidence",
        "buildEvidence", "FraudSignal.builder", "signalType", "category", "displayMessage",
        "riskImpact", "confidence", "severity", "evidence", "createdAt"
    ], max_lines=32))
    table(doc, ["Field trong FraudSignal", "Nguồn", "Ý nghĩa"], [
        ["signalType", "descriptorFor(eventType)", "Tên chuẩn: TAB_SWITCH, NO_CAMERA, MULTIPLE_FACES..."],
        ["category", "descriptorFor", "Nhóm tính điểm: SCREEN_LEAVE, TECHNICAL, AI_CAMERA..."],
        ["riskImpact", "descriptor/risk properties", "Điểm cơ sở hoặc mức ảnh hưởng nghiệp vụ."],
        ["confidence", "Frontend/AI hoặc default", "Độ tin cậy của signal."],
        ["severity", "descriptor hoặc AI severity", "LOW/MEDIUM/HIGH/CRITICAL."],
        ["evidence", "payload/browserContext/telemetry/AI diagnostics", "Bằng chứng để review sau này."],
    ], widths=[2.0, 2.2, 4.3])


def add_flow_risk(doc: Document) -> None:
    doc.add_heading("4. Luồng 3 - Tính Điểm Rủi Ro", level=1)
    para(doc, "Mục tiêu: biến nhiều FraudSignal rời rạc thành một điểm rủi ro dễ hiểu cho giáo viên.")
    doc.add_heading("4.1 Code thật: recomputeRisk", level=2)
    code(doc, snippet("BE/src/main/java/com/example/demo/service/RiskScoringService.java", 221, 311, keep=[
        "recomputeRisk", "findByAttemptOrderByCreatedAtAsc", "computeCategoryScores",
        "screenLeaveScore", "clipboardScore", "technicalScore", "identityScore",
        "heartbeatScore", "visualIdentityScore", "behaviorScore", "totalRisk",
        "resolveLevel", "setRiskScore", "syncRiskFlag", "RiskScoreLog", "notify"
    ], max_lines=34))
    table(doc, ["Bước", "Code/biến", "Giải thích"], [
        ["1", "signals = findByAttempt...", "Lấy toàn bộ signal của bài thi theo thời gian."],
        ["2", "computeCategoryScores", "Gom signal theo nhóm và tính điểm từng nhóm."],
        ["3", "screenLeaveScore...visualIdentityScore", "Áp cap để mỗi nhóm không vượt quá điểm tối đa."],
        ["4", "behaviorScore = min(70,...)", "Hành vi browser/technical/heartbeat tối đa 70 điểm."],
        ["5", "totalRisk = min(100,...)", "Tổng điểm tối đa 100."],
        ["6", "resolveLevel(totalRisk)", "Đổi điểm thành CLEAN/SUSPICIOUS/HIGH_RISK/CRITICAL."],
        ["7", "syncRiskFlag", "Nếu rủi ro cao thì tạo flag review cho giáo viên."],
        ["8", "notifyRiskUpdated", "Đẩy realtime cho dashboard."],
    ], widths=[0.7, 2.8, 5.2])

    doc.add_heading("4.2 Công thức risk score", level=2)
    code(doc, "screenLeaveScore = min(cluster(SCREEN_LEAVE), 35)\nclipboardScore   = min(cluster(CLIPBOARD), 25)\ntechnicalScore   = min(cluster(TECHNICAL), 25)\nidentityScore    = min(max(IDENTITY), 30)\nvisualScore      = min(cluster(VISUAL_IDENTITY), 40)\nbehaviorScore    = min(70, screenLeave + clipboard + technical + heartbeat)\ntotalRisk        = min(100, behaviorScore + identityScore + visualScore)")
    table(doc, ["Điểm", "Mức", "Ý nghĩa"], [
        ["0-19", "CLEAN", "Chưa có dấu hiệu đáng kể."],
        ["20-49", "SUSPICIOUS", "Có dấu hiệu cần theo dõi."],
        ["50-74", "HIGH_RISK", "Cần giáo viên review."],
        ["75-100", "CRITICAL", "Nguy cơ gian lận nghiêm trọng."],
    ], widths=[1.4, 1.8, 4.8])

    doc.add_heading("4.3 Code thật: cluster/dedup điểm", level=2)
    code(doc, snippet("BE/src/main/java/com/example/demo/service/RiskScoringService.java", 358, 447, keep=[
        "computeCategoryScores", "computeClusteredScore", "signalsByCategory",
        "categoryDedupSeconds", "sortSignalsByCreatedAt", "IDENTITY", "clusterMaxScore",
        "secondsSinceClusterStart", "totalScore", "categoryCap"
    ], max_lines=30))
    table(doc, ["Vấn đề", "Cách code xử lý"], [
        ["Một signal bị gửi lặp liên tục", "Trong cùng dedup window chỉ lấy điểm cao nhất."],
        ["Nhiều signal khác nhau trong cùng category", "Chia thành cluster theo thời gian, mỗi cluster cộng một lần."],
        ["Signal định danh như IP/fingerprint", "Không cộng lặp, chỉ lấy điểm cao nhất trong session."],
        ["Một category quá nhiều signal", "categoryCap giới hạn điểm tối đa."],
    ], widths=[3.2, 5.0])


def add_flow_realtime(doc: Document) -> None:
    doc.add_heading("5. Luồng 4 - Dashboard Realtime", level=1)
    para(doc, "Mục tiêu: khi backend ghi signal hoặc cập nhật riskScore, giáo viên thấy thay đổi gần như ngay lập tức.")
    doc.add_heading("5.1 Code thật: applyRealtimeEvent", level=2)
    code(doc, snippet("demo/src/composables/useExamMonitoring.js", 42, 116, keep=[
        "function applyRealtimeEvent", "const type", "attemptId", "store.patchAttemptFromRealtime",
        "addAlertFromEvent", "addSystemEventFromEvent", "alertTypes", "store.addAlert",
        "signalType", "severity", "riskScore"
    ], max_lines=32))
    table(doc, ["Phần UI", "Được cập nhật bởi", "Kết quả"], [
        ["Dashboard card", "store.patchAttemptFromRealtime", "Điểm risk, trạng thái attempt, học sinh được cập nhật."],
        ["Alert panel", "addAlertFromEvent", "Thêm cảnh báo mới: fraud signal, risk update, AI camera signal."],
        ["Timeline/system event", "addSystemEventFromEvent", "Hiển thị pause/resume/warning/stop/submitted."],
        ["Fallback polling", "connect/startPolling", "Nếu WebSocket lỗi, dashboard vẫn đồng bộ định kỳ."],
    ], widths=[2.2, 2.8, 3.0])


def add_reference_tables(doc: Document) -> None:
    doc.add_heading("6. Bảng Tra Cứu Nhanh", level=1)
    doc.add_heading("6.1 Signal quan trọng", level=2)
    table(doc, ["Signal", "Nguồn", "Điểm", "Khi nào xuất hiện"], [
        ["NO_CAMERA", "AI/Heartbeat", "20", "Camera tắt hoặc track ended/disabled."],
        ["NO_MIC", "AI/Heartbeat", "20", "Micro tắt."],
        ["FACE_NOT_DETECTED", "AI camera", "20", "Không phát hiện mặt ổn định."],
        ["MULTIPLE_FACES", "AI camera", "25", "Có nhiều mặt trong frame."],
        ["FACE_SPOOFING_SUSPECTED", "AI spoofing", "25", "Nghi ảnh giả, replay, deepfake."],
        ["GAZE_OFF_SCREEN", "AI gaze", "12", "Nhìn lệch khỏi màn hình đủ lâu."],
        ["TAB_SWITCH", "Browser", "10", "Chuyển tab."],
        ["WINDOW_BLUR", "Browser", "8", "Cửa sổ mất focus."],
        ["EXIT_FULLSCREEN", "Browser", "15", "Thoát fullscreen."],
        ["COPY_PASTE", "Browser", "10", "Copy/paste nội dung."],
        ["DEVTOOLS_OPEN", "Browser", "22", "Mở DevTools."],
        ["DUPLICATE_IP", "Backend identity", "25", "Trùng IP đáng ngờ."],
    ], widths=[2.0, 1.8, 0.9, 4.0])

    doc.add_heading("6.2 File cần nhớ", level=2)
    table(doc, ["File", "Nên đọc khi muốn hiểu"], [
        ["ai-service/app/proctor.py", "AI camera: mặt, mắt, gaze, spoofing, quality gate, stability."],
        ["BE/.../AiAssistService.java", "Backend gọi AI service và ghi signal AI."],
        ["BE/.../ExamEventService.java", "Nhận event trình duyệt và heartbeat."],
        ["BE/.../FraudSignalService.java", "Chuẩn hóa signal, category, severity, evidence."],
        ["BE/.../RiskScoringService.java", "Công thức risk score, cluster, cap, flag."],
        ["demo/src/composables/useExamMonitoring.js", "Dashboard realtime giáo viên."],
    ], widths=[3.0, 5.2])


def add_summary(doc: Document) -> None:
    doc.add_heading("7. Kết Luận Ngắn Gọn", level=1)
    bullet(doc, "Nếu cần hiểu phát hiện gian lận bằng camera, bắt đầu từ AiAssistService.analyzeFrame rồi sang ProctorAnalyzer.analyze_frame.")
    bullet(doc, "Nếu cần hiểu phát hiện gian lận từ trình duyệt, bắt đầu từ ExamEventService.ingestBatchInternal.")
    bullet(doc, "Nếu cần hiểu vì sao học sinh bị đánh điểm rủi ro, đọc RiskScoringService.recomputeRisk và computeClusteredScore.")
    bullet(doc, "Nếu cần hiểu giáo viên thấy cảnh báo thế nào, đọc useExamMonitoring.applyRealtimeEvent.")
    para(doc, "Cốt lõi của hệ thống là: mọi nguồn dữ liệu đều được chuẩn hóa thành FraudSignal, sau đó RiskScoringService là nơi duy nhất tính riskScore/riskLevel để đảm bảo dễ giải thích và dễ kiểm soát.")


def build() -> Document:
    doc = setup_doc()
    add_cover(doc)
    add_overview(doc)
    add_flow_camera(doc)
    add_flow_browser(doc)
    add_flow_risk(doc)
    add_flow_realtime(doc)
    add_reference_tables(doc)
    add_summary(doc)
    return doc


def main() -> None:
    doc = build()
    if OUT.exists():
        OUT.unlink()
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
