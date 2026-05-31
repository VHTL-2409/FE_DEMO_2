from __future__ import annotations

from datetime import datetime
from pathlib import Path
import re

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Phan_tich_chi_tiet_theo_ham_service_gian_lan_va_ai_service.docx"
FONT = "Times New Roman"
FONT_SIZE = Pt(13)
CODE_FONT_SIZE = Pt(10)


FILES = [
    (
        "Backend service phát hiện gian lận",
        "BE/src/main/java/com/example/demo/service/FraudSignalTypeNormalizer.java",
        "Chuẩn hóa tên signal gian lận để các service dùng cùng một bộ thuật ngữ.",
    ),
    (
        "Backend service phát hiện gian lận",
        "BE/src/main/java/com/example/demo/service/FraudSignalService.java",
        "Ghi nhận FraudSignal từ event hoặc tín hiệu hệ thống, xác định severity, confidence, riskImpact và evidence.",
    ),
    (
        "Backend service phát hiện gian lận",
        "BE/src/main/java/com/example/demo/service/RiskScoringService.java",
        "Tính điểm rủi ro theo nhóm tín hiệu, cửa sổ chống trùng, giới hạn điểm, cấp rủi ro và proctor flag.",
    ),
    (
        "Backend service phát hiện gian lận",
        "BE/src/main/java/com/example/demo/service/FraudWarningService.java",
        "Sinh cảnh báo gian lận từ signal hoặc pattern, chống trùng cảnh báo và hỗ trợ giáo viên review.",
    ),
    (
        "Backend service phát hiện gian lận",
        "BE/src/main/java/com/example/demo/service/FraudAnalysisService.java",
        "Phân tích hậu kiểm: giống đáp án, thời gian làm bài, thống kê điểm, hành vi, IP và pattern nghi vấn.",
    ),
    (
        "Backend service phát hiện gian lận",
        "BE/src/main/java/com/example/demo/service/MLRiskScoringService.java",
        "Trích xuất feature, tính anomaly score và kết hợp điểm rule-based với điểm rủi ro theo ML.",
    ),
    (
        "Backend service phát hiện gian lận",
        "BE/src/main/java/com/example/demo/service/FraudModelTrainer.java",
        "Huấn luyện và đánh giá mô hình gian lận từ dữ liệu attempt và fraud signal.",
    ),
    (
        "Backend service phát hiện gian lận",
        "BE/src/main/java/com/example/demo/service/AnswerSimilarityService.java",
        "So sánh đáp án giữa các thí sinh để phát hiện pattern giống nhau bất thường.",
    ),
    (
        "Backend service phát hiện gian lận",
        "BE/src/main/java/com/example/demo/service/DeviceFingerprintService.java",
        "Tạo và chuẩn hóa fingerprint thiết bị để phát hiện đổi thiết bị hoặc dùng chung thiết bị/IP.",
    ),
    (
        "Backend service phát hiện gian lận",
        "BE/src/main/java/com/example/demo/service/ProctorEvidenceImageService.java",
        "Lưu ảnh bằng chứng từ camera AI và tạo URL evidence cho warning hoặc signal.",
    ),
    (
        "Backend service phát hiện gian lận",
        "BE/src/main/java/com/example/demo/service/ProctorFlagService.java",
        "Truy vấn và review flag proctoring được sinh bởi risk scoring.",
    ),
    ("AI service", "ai-service/app/main.py", "FastAPI entrypoint: expose API OCR, proctor frame/behavior, generator, evaluator, analytics và chat."),
    ("AI service", "ai-service/app/schemas.py", "Pydantic schema cho request/response AI, đặc biệt FrameAnalysis và BehaviorAnalysis."),
    ("AI service", "ai-service/app/proctor.py", "Lõi AI proctoring: decode frame, detect mặt/mắt, gaze, spoofing, quality gate và signal."),
    ("AI service", "ai-service/app/ocr.py", "OCR engine đọc text từ ảnh/PDF đầu vào cho các chức năng AI phụ trợ."),
    ("AI service", "ai-service/app/analytics.py", "AI analytics dự đoán hiệu suất, chất lượng câu hỏi và phân tích dữ liệu học tập."),
    ("AI service", "ai-service/app/evaluator.py", "AI chấm hoặc tư vấn đánh giá bài tự luận theo tiêu chí."),
    ("AI service", "ai-service/app/generator.py", "AI sinh câu hỏi từ prompt hoặc nội dung đầu vào."),
    ("AI service", "ai-service/app/chat.py", "AI chat assistant dùng model provider để trả lời hội thoại."),
    ("AI service", "ai-service/app/openai_client.py", "Wrapper tạo OpenAI client theo cấu hình môi trường."),
    ("AI service", "ai-service/app/model_key_resolver.py", "Resolve model, key và provider cho các chức năng AI theo environment."),
]


SIGNAL_TERMS = {
    "TAB_SWITCH": "rời tab",
    "WINDOW_BLUR": "mất focus cửa sổ",
    "EXIT_FULLSCREEN": "thoát fullscreen",
    "COPY_PASTE": "copy/paste",
    "RIGHT_CLICK": "click phải",
    "PRINT_SCREEN": "chụp màn hình",
    "DEVTOOLS": "mở devtools",
    "NO_CAMERA": "camera tắt",
    "NO_MIC": "micro tắt",
    "FACE_NOT_DETECTED": "không thấy mặt",
    "MULTIPLE_FACES": "nhiều mặt",
    "GAZE_AWAY": "nhìn ra ngoài",
    "SPOOFING": "nghi giả mạo camera",
    "DEVICE_CHANGE": "đổi thiết bị",
    "DUPLICATE_IP": "trùng IP",
    "IDENTITY": "nhóm định danh",
    "SCREEN_LEAVE": "nhóm rời màn hình",
    "CLIPBOARD": "nhóm clipboard",
    "TECHNICAL": "nhóm kỹ thuật",
    "HEARTBEAT": "nhóm heartbeat",
    "VISUAL_IDENTITY": "nhóm định danh bằng hình ảnh",
}


def apply_font(run, size: Pt = FONT_SIZE, font: str = FONT) -> None:
    run.font.name = font
    run.font.size = size
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run._element.rPr.rFonts.set(qn("w:cs"), font)


def para(doc: Document, text: str, style: str | None = None, *, bold: bool = False):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    apply_font(run)
    return p


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def cell_text(cell, text, *, bold: bool = False, color=None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    apply_font(run)


def header(table, columns: list[str]) -> None:
    for i, title in enumerate(columns):
        cell = table.rows[0].cells[i]
        cell_text(cell, title, bold=True, color=(255, 255, 255))
        shade(cell, "1F4E79")


def read_lines(path: Path) -> list[str]:
    raw = path.read_bytes()
    for enc in ("utf-8-sig", "utf-8", "cp1258", "cp1252"):
        try:
            return raw.decode(enc).splitlines()
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace").splitlines()


def file_kind(path: Path) -> str:
    name = path.name.lower()
    if name.endswith(".java"):
        return "java"
    if name.endswith(".py"):
        return "python"
    return "text"


def python_symbol(line: str):
    m = re.match(r"^(\s*)(class|def|async def)\s+([A-Za-z_][\w_]*)\s*(\(.*)?", line)
    if not m:
        return None
    indent = len(m.group(1).replace("\t", "    "))
    return {
        "kind": m.group(2),
        "name": m.group(3),
        "indent": indent,
        "signature": line.strip(),
    }


def java_symbol(line: str):
    s = line.strip()
    if re.match(r"(public|private|protected)?\s*(class|record|enum)\s+", s):
        m = re.search(r"\b(class|record|enum)\s+([A-Za-z_][\w_]*)", s)
        if m:
            return {"kind": m.group(1), "name": m.group(2), "signature": s}
    if re.match(r"(public|private|protected)\s+.*\)\s*(\{|$)", s) and "(" in s:
        if any(x in s for x in (" if ", " for ", " while ", " switch ")):
            return None
        before_paren = s.split("(", 1)[0].strip()
        name = before_paren.split()[-1]
        return {"kind": "method", "name": name, "signature": re.sub(r"\{\s*$", "", s)}
    return None


def extract_python_blocks(lines: list[str]) -> list[dict]:
    candidates = []
    for idx, line in enumerate(lines):
        sym = python_symbol(line)
        if sym:
            sym["start"] = idx + 1
            candidates.append(sym)
    blocks = []
    for i, sym in enumerate(candidates):
        start_idx = sym["start"] - 1
        end = len(lines)
        for j in range(i + 1, len(candidates)):
            nxt = candidates[j]
            if nxt["indent"] <= sym["indent"]:
                end = nxt["start"] - 1
                break
        sym["end"] = end
        sym["code"] = lines[start_idx:end]
        blocks.append(sym)
    return blocks


def extract_java_blocks(lines: list[str]) -> list[dict]:
    blocks = []
    idx = 0
    while idx < len(lines):
        sym = java_symbol(lines[idx])
        if not sym:
            idx += 1
            continue
        brace_count = 0
        seen_open = False
        end = idx
        scan = idx
        while scan < len(lines):
            text = strip_java_strings(lines[scan])
            brace_count += text.count("{")
            if "{" in text:
                seen_open = True
            brace_count -= text.count("}")
            if seen_open and brace_count <= 0:
                end = scan
                break
            scan += 1
        if not seen_open:
            end = idx
        sym["start"] = idx + 1
        sym["end"] = end + 1
        sym["code"] = lines[idx : end + 1]
        blocks.append(sym)
        idx = end + 1
    return blocks


def strip_java_strings(line: str) -> str:
    # Good enough for brace counting in regular service files.
    line = re.sub(r'"(?:\\.|[^"\\])*"', '""', line)
    line = re.sub(r"'(?:\\.|[^'\\])*'", "''", line)
    return line


def code_excerpt(lines: list[str], limit: int = 80) -> str:
    if len(lines) <= limit:
        return "\n".join(lines)
    head = lines[:45]
    tail = lines[-25:]
    omitted = len(lines) - len(head) - len(tail)
    return "\n".join(head + [f"... // lược bớt {omitted} dòng ở giữa để tài liệu dễ theo dõi"] + tail)


def add_code_block(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.2)
    for line_no, line in enumerate(text.splitlines() or [""], 1):
        if line_no > 1:
            p.add_run("\n")
        run = p.add_run(line)
        apply_font(run, CODE_FONT_SIZE, "Courier New")


def terms_in_code(code: list[str]) -> list[str]:
    text = "\n".join(code)
    found = []
    for term, meaning in SIGNAL_TERMS.items():
        if term in text:
            found.append(f"{term} ({meaning})")
    return found


def infer_purpose(block: dict, rel: str, role: str) -> str:
    name = block["name"]
    lower = name.lower()
    if block["kind"] in ("class", "record", "enum"):
        return f"Khai báo {block['kind']} `{name}`. Thành phần này là khung chính để {role[0].lower() + role[1:]}"
    if "record" in lower or "save" in lower or "ingest" in lower:
        return "Nhận dữ liệu đầu vào, chuẩn hóa và lưu xuống database để hệ thống có bằng chứng/phân tích tiếp."
    if "risk" in lower or "score" in lower:
        return "Tính toán hoặc trả về điểm rủi ro, cấp rủi ro và dữ liệu giải thích cho attempt."
    if "warning" in lower or "alert" in lower:
        return "Tạo, đọc hoặc review cảnh báo gian lận dựa trên signal và pattern đã phát hiện."
    if "analyze" in lower or "analysis" in lower:
        return "Phân tích dữ liệu đầu vào để phát hiện bất thường hoặc tạo kết luận hỗ trợ giám sát."
    if "feature" in lower or "extract" in lower:
        return "Trích xuất đặc trưng từ attempt/signal để phục vụ scoring hoặc mô hình ML."
    if "normalize" in lower:
        return "Chuẩn hóa dữ liệu đầu vào để các phần còn lại của hệ thống xử lý nhất quán."
    if "detect" in lower or "track" in lower:
        return "Phát hiện hoặc theo dõi đặc trưng trong dữ liệu, đặc biệt là ảnh camera, mắt, mặt hoặc hành vi."
    if "build" in lower or "toresponse" in lower or lower.startswith("to"):
        return "Dựng DTO/response/payload từ entity hoặc kết quả phân tích để trả cho caller."
    if "init" in lower or lower == "__init__":
        return "Khởi tạo dependency, cấu hình và trạng thái cần thiết trước khi service xử lý request."
    if "health" in lower or "ready" in lower or "status" in lower:
        return "Trả trạng thái sẵn sàng của service hoặc dependency để backend/frontend kiểm tra khả dụng."
    return f"Thực hiện một phần logic trong `{rel}`; hàm này phục vụ vai trò chung của file: {role}"


def infer_inputs_outputs(block: dict, kind: str) -> tuple[str, str]:
    signature = block["signature"]
    if kind == "python":
        params = "()"
        if "(" in signature and ")" in signature:
            params = signature[signature.find("(") : signature.rfind(")") + 1]
        output = "Giá trị trả về được xác định bởi `return` trong hàm; thường là dict/DTO/list hoặc trạng thái xử lý."
        if "->" in signature:
            output = "Kiểu trả về khai báo sau `->` trong signature."
        return params, output
    if kind == "java":
        if "(" in signature and ")" in signature:
            params = signature[signature.find("(") : signature.rfind(")") + 1]
            before = signature.split("(", 1)[0].split()
            output = before[-2] if len(before) >= 2 else "void/constructor"
            return params, output
    return "Không có tham số rõ ràng.", "Không có kiểu trả về rõ ràng."


def infer_flow(block: dict) -> list[str]:
    code = "\n".join(block["code"])
    flow = []
    if re.search(r"\b(find|findBy|Repository|query)\b", code):
        flow.append("Đọc dữ liệu cần thiết từ repository hoặc nguồn dữ liệu liên quan.")
    if re.search(r"\bif\b|\belif\b|\bswitch\b", code):
        flow.append("Kiểm tra điều kiện nghiệp vụ để phân nhánh theo trạng thái, loại signal hoặc ngưỡng.")
    if re.search(r"\bfor\b|\bwhile\b|\.stream\(\)", code):
        flow.append("Duyệt danh sách dữ liệu để tính toán, gom nhóm hoặc tạo kết quả.")
    if re.search(r"\bsave\(|\.save\(", code):
        flow.append("Lưu entity hoặc kết quả xử lý để phục vụ audit, cảnh báo hoặc phân tích sau.")
    if re.search(r"builder\(\)|\.build\(\)", code):
        flow.append("Dựng DTO/entity bằng builder để dữ liệu trả về có cấu trúc rõ ràng.")
    if re.search(r"return\b", code):
        flow.append("Trả kết quả cho caller để controller/service khác tiếp tục pipeline.")
    if not flow:
        flow.append("Thực thi logic trực tiếp theo signature; chủ yếu chuẩn bị dữ liệu hoặc cấu hình.")
    return flow


def created_items(block: dict, kind: str) -> list[str]:
    items = []
    seen = set()
    for line in block["code"]:
        s = line.strip()
        if kind == "java":
            m = re.match(
                r"(?:final\s+)?([A-Z][A-Za-z0-9_<>, ?]+|int|long|double|boolean|String|var)\s+([a-zA-Z_][\w_]*)\s*=",
                s,
            )
            if m:
                typ, name = m.group(1), m.group(2)
                key = ("var", name)
                if key not in seen:
                    seen.add(key)
                    items.append(f"`{name}`: biến kiểu `{typ}` được tạo để giữ dữ liệu trung gian cho bước xử lý sau.")
            b = re.search(r"([A-Za-z0-9_]+)\.builder\(\)", s)
            if b:
                name = b.group(1)
                key = ("builder", name)
                if key not in seen:
                    seen.add(key)
                    items.append(f"`{name}.builder()`: tạo object dạng builder, thường là DTO/entity trả về hoặc entity sẽ lưu database.")
            if ".save(" in s:
                key = ("save", s)
                if key not in seen:
                    seen.add(key)
                    items.append("Entity được lưu vào database qua repository, tạo dữ liệu bền vững cho audit/phân tích/cảnh báo.")
        elif kind == "python":
            m = re.match(r"([a-zA-Z_][\w_]*)\s*=", s)
            if m and "==" not in s:
                name = m.group(1)
                key = ("var", name)
                if key not in seen:
                    seen.add(key)
                    items.append(f"`{name}`: biến được tạo để lưu trạng thái, dữ liệu request, kết quả AI hoặc response trung gian.")
            if "dict(" in s or s.startswith("{") or "{" in s and ":" in s:
                key = ("dict", len(items))
                if key not in seen:
                    seen.add(key)
                    items.append("Một `dict`/object JSON được tạo để gom dữ liệu trả về hoặc evidence cho backend.")
            if "np.array" in s or "Image.open" in s or "cv2." in s:
                key = ("image", len(items))
                if key not in seen:
                    seen.add(key)
                    items.append("Dữ liệu ảnh/mảng ảnh được tạo hoặc biến đổi để phục vụ phân tích AI camera/OCR.")
    if not items:
        return ["Hàm này chủ yếu điều phối hoặc trả dữ liệu; không tạo object nghiệp vụ lớn, nhưng có thể dùng biến cục bộ trong từng nhánh xử lý."]
    return items[:12]


def split_code_phases(block: dict) -> list[tuple[str, str, list[str]]]:
    phases: list[tuple[str, str, list[str]]] = []
    current_title = "Khởi tạo và chuẩn bị dữ liệu"
    current_lines: list[str] = []

    def title_for(line: str) -> str | None:
        s = line.strip()
        low = s.lower()
        if not s:
            return None
        if s.startswith(("if ", "if(")):
            return "Kiểm tra điều kiện và rẽ nhánh nghiệp vụ"
        if s.startswith(("else", "elif ")):
            return "Xử lý nhánh thay thế"
        if s.startswith(("for ", "for(", "while ")):
            return "Duyệt dữ liệu và tổng hợp kết quả"
        if ".stream()" in s or ".filter(" in s or ".map(" in s or "Collectors" in s:
            return "Biến đổi danh sách bằng stream/pipeline"
        if ".save(" in s or "repository.save" in low:
            return "Lưu dữ liệu xuống database"
        if ".builder()" in s or ".build()" in s:
            return "Dựng DTO/entity kết quả"
        if s.startswith("return "):
            return "Trả kết quả cho caller"
        if "throw " in s or s.startswith("raise "):
            return "Dừng xử lý khi dữ liệu không hợp lệ"
        if "cv2" in s or "mediapipe" in low or "dnn" in low or "face" in low and "detect" in low:
            return "Phân tích ảnh bằng AI/computer vision"
        if "openai" in low or "client." in low or "model" in low:
            return "Gọi model AI hoặc chuẩn bị prompt/model"
        return None

    for line in block["code"]:
        next_title = title_for(line)
        if next_title and current_lines and next_title != current_title:
            phases.append((current_title, explain_phase(current_title, current_lines), current_lines))
            current_lines = []
            current_title = next_title
        elif next_title:
            current_title = next_title
        current_lines.append(line)
    if current_lines:
        phases.append((current_title, explain_phase(current_title, current_lines), current_lines))
    return phases[:10]


def explain_phase(title: str, lines: list[str]) -> str:
    text = "\n".join(lines)
    low = text.lower()
    if title == "Khởi tạo và chuẩn bị dữ liệu":
        return "Khối này lấy dữ liệu đầu vào, khai báo biến trung gian hoặc chuẩn bị dependency để các bước chính phía sau có đủ dữ liệu xử lý."
    if title == "Kiểm tra điều kiện và rẽ nhánh nghiệp vụ":
        return "Khối này kiểm tra trạng thái, quyền truy cập, dữ liệu thiếu, threshold hoặc loại signal. Ý nghĩa là tránh xử lý sai và phân loại đúng tình huống."
    if title == "Xử lý nhánh thay thế":
        return "Khối này xử lý trường hợp điều kiện trước không đúng, giúp hàm có hành vi rõ ràng cho nhiều trạng thái dữ liệu."
    if title == "Duyệt dữ liệu và tổng hợp kết quả":
        return "Khối này đi qua nhiều phần tử như signal, attempt, warning, answer hoặc feature để tính toán/tổng hợp dữ liệu cuối."
    if title == "Biến đổi danh sách bằng stream/pipeline":
        return "Khối này lọc, map, group hoặc collect danh sách dữ liệu. Đây thường là phần tạo thống kê hoặc response từ entity."
    if title == "Lưu dữ liệu xuống database":
        return "Khối này tạo dữ liệu bền vững trong database. Kết quả lưu lại thường là bằng chứng, signal, warning, flag, log hoặc model metadata."
    if title == "Dựng DTO/entity kết quả":
        return "Khối này gán từng field vào builder để tạo object có cấu trúc. Object đó có thể được trả về API, lưu DB hoặc gửi sang service khác."
    if title == "Trả kết quả cho caller":
        return "Khối này kết thúc hàm và trả dữ liệu đã xử lý. Đây là contract đầu ra mà controller/service khác phụ thuộc vào."
    if title == "Dừng xử lý khi dữ liệu không hợp lệ":
        return "Khối này bảo vệ hệ thống khỏi dữ liệu sai, thiếu quyền hoặc trạng thái không hợp lệ bằng cách ném lỗi/dừng hàm."
    if title == "Phân tích ảnh bằng AI/computer vision":
        return "Khối này biến đổi hoặc phân tích ảnh camera/OCR bằng OpenCV, DNN, Haar cascade hoặc landmark để tạo kết quả nhận diện."
    if title == "Gọi model AI hoặc chuẩn bị prompt/model":
        return "Khối này chuẩn bị model, prompt hoặc client để gọi tác vụ AI như sinh câu hỏi, chấm bài, chat hoặc phân tích."
    if "risk" in low or "score" in low:
        return "Khối này liên quan tính điểm hoặc cấp độ rủi ro."
    return "Khối này thực hiện một bước logic trong hàm, góp phần tạo kết quả cuối cùng."


def business_meaning(block: dict, rel: str) -> str:
    name = block["name"].lower()
    file_name = Path(rel).name
    if "RiskScoringService" in rel:
        return "Ý nghĩa nghiệp vụ: chuyển các dấu hiệu rời rạc thành điểm rủi ro có thể giải thích, giúp giáo viên biết thí sinh đang ở mức CLEAN, SUSPICIOUS, HIGH_RISK hay CRITICAL."
    if "FraudSignalService" in rel:
        return "Ý nghĩa nghiệp vụ: biến event kỹ thuật thành FraudSignal chuẩn hóa để toàn hệ thống có cùng cách hiểu về một hành vi nghi vấn."
    if "FraudWarningService" in rel:
        return "Ý nghĩa nghiệp vụ: tạo cảnh báo dễ đọc cho giáo viên, tránh trùng lặp và lưu trạng thái review."
    if "FraudAnalysisService" in rel:
        return "Ý nghĩa nghiệp vụ: phân tích sau kỳ thi để tìm pattern khó nhìn thấy realtime như giống đáp án, làm bài quá nhanh hoặc trùng IP."
    if "MLRiskScoringService" in rel or "FraudModelTrainer" in rel:
        return "Ý nghĩa nghiệp vụ: bổ sung lớp đánh giá dựa trên feature/anomaly/model để tăng khả năng phát hiện hành vi bất thường ngoài rule cố định."
    if "AnswerSimilarityService" in rel:
        return "Ý nghĩa nghiệp vụ: phát hiện dấu hiệu trao đổi đáp án hoặc sao chép bằng cách so sánh mẫu trả lời giữa nhiều thí sinh."
    if "DeviceFingerprintService" in rel:
        return "Ý nghĩa nghiệp vụ: nhận diện thiết bị/ngữ cảnh truy cập để phát hiện đổi thiết bị, dùng chung thiết bị hoặc dấu hiệu tổ chức gian lận."
    if "ProctorEvidenceImageService" in rel:
        return "Ý nghĩa nghiệp vụ: lưu ảnh bằng chứng để cảnh báo không chỉ là kết luận mà có dữ liệu chứng minh."
    if "proctor.py" in rel:
        return "Ý nghĩa nghiệp vụ: phân tích camera để phát hiện không có mặt, nhiều mặt, nhìn ra ngoài, che mặt, ánh sáng kém hoặc nghi giả mạo."
    if file_name in {"main.py", "schemas.py"}:
        return "Ý nghĩa nghiệp vụ: định nghĩa API và contract dữ liệu để backend/frontend gọi đúng chức năng AI."
    if any(x in name for x in ["generate", "evaluate", "chat", "ocr", "analytics"]):
        return "Ý nghĩa nghiệp vụ: cung cấp năng lực AI phụ trợ như OCR, sinh câu hỏi, đánh giá tự luận, chat hoặc phân tích học tập."
    return "Ý nghĩa nghiệp vụ: hỗ trợ pipeline AI/phát hiện gian lận bằng cách xử lý dữ liệu đúng cấu trúc và đúng ngữ cảnh."


def risks_or_notes(block: dict) -> list[str]:
    code = "\n".join(block["code"]).lower()
    notes = []
    if "threshold" in code or "min" in code or "cap" in code:
        notes.append("Có ngưỡng/cap cấu hình; khi thay đổi cần kiểm tra tác động tới mức cảnh báo.")
    if "dedup" in code or "window" in code:
        notes.append("Có cơ chế chống trùng theo thời gian; giúp tránh spam signal/warning.")
    if "confidence" in code:
        notes.append("Có dùng độ tin cậy; cần giải thích rõ confidence không phải xác suất tuyệt đối.")
    if "exception" in code or "throw" in code or "except" in code:
        notes.append("Có xử lý lỗi; cần đảm bảo lỗi không làm mất dữ liệu giám sát quan trọng.")
    if "cv2" in code or "mediapipe" in code or "dnn" in code:
        notes.append("Phụ thuộc thư viện thị giác máy tính; môi trường thiếu dependency sẽ ảnh hưởng chất lượng detect.")
    if "openai" in code or "model" in code:
        notes.append("Phụ thuộc model/provider; cần cấu hình API key/model đúng khi deploy.")
    return notes or ["Không có ghi chú đặc biệt ngoài việc cần giữ đúng contract đầu vào/đầu ra của hàm."]


def file_summary(rel: str, role: str, blocks: list[dict]) -> str:
    names = ", ".join(b["name"] for b in blocks[:8])
    more = "" if len(blocks) <= 8 else f", ... và {len(blocks) - 8} thành phần khác"
    return f"File này có {len(blocks)} lớp/hàm/phương thức được phân tích. Vai trò chính: {role} Các thành phần nổi bật: {names}{more}."


def add_file_section(doc: Document, group: str, rel: str, role: str) -> None:
    path = ROOT / rel
    para(doc, rel, "Heading 1")
    para(doc, f"Nhóm: {group}")
    para(doc, f"Vai trò: {role}")
    if not path.exists():
        para(doc, "Không tìm thấy file trong workspace.")
        return
    lines = read_lines(path)
    kind = file_kind(path)
    blocks = extract_java_blocks(lines) if kind == "java" else extract_python_blocks(lines)
    para(doc, file_summary(rel, role, blocks))

    if blocks:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        header(table, ["Thành phần", "Loại", "Dòng", "Vai trò ngắn"])
        for block in blocks:
            row = table.add_row().cells
            cell_text(row[0], block["name"])
            cell_text(row[1], block["kind"])
            cell_text(row[2], f"{block['start']}-{block['end']}")
            cell_text(row[3], infer_purpose(block, rel, role))

    for block in blocks:
        para(doc, f"{block['name']} ({block['start']}-{block['end']})", "Heading 2")
        para(doc, f"Khai báo: {block['signature']}")
        para(doc, f"Chức năng của hàm/lớp: {infer_purpose(block, rel, role)}")
        para(doc, business_meaning(block, rel))
        inputs, outputs = infer_inputs_outputs(block, kind)
        para(doc, f"Đầu vào: {inputs}")
        para(doc, f"Đầu ra: {outputs}")
        found_terms = terms_in_code(block["code"])
        if found_terms:
            para(doc, "Signal/khái niệm liên quan: " + "; ".join(found_terms))
        para(doc, "Những thứ được tạo ra trong code:")
        for item in created_items(block, kind):
            para(doc, f"- {item}")
        para(doc, "Luồng xử lý chính:")
        for item in infer_flow(block):
            para(doc, f"- {item}")
        para(doc, "Giải thích code theo khối:")
        for title, description, phase_lines in split_code_phases(block):
            para(doc, f"- {title}: {description}")
            sample = "\n".join(phase_lines[:8])
            if len(phase_lines) > 8:
                sample += f"\n... // còn {len(phase_lines) - 8} dòng trong khối này"
            add_code_block(doc, sample)
        para(doc, "Ghi chú khi giải thích/bảo vệ:")
        for item in risks_or_notes(block):
            para(doc, f"- {item}")
        para(doc, "Code gốc của hàm/lớp:")
        add_code_block(doc, code_excerpt(block["code"]))


def build_doc() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    for style_name in ("Normal", "Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[style_name]
        style.font.name = FONT
        style.font.size = FONT_SIZE
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style._element.rPr.rFonts.set(qn("w:cs"), FONT)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("PHÂN TÍCH THEO HÀM: SERVICE PHÁT HIỆN GIAN LẬN VÀ AI-SERVICE")
    run.bold = True
    apply_font(run)

    para(doc, f"Dự án: {ROOT}")
    para(doc, f"Ngày tạo: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    para(
        doc,
        "Tài liệu này thay thế kiểu phân tích từng dòng bằng phân tích theo hàm/lớp để dễ theo dõi hơn. Mỗi hàm có vai trò, đầu vào, đầu ra, luồng xử lý chính, ghi chú và code gốc.",
    )

    para(doc, "Cách đọc tài liệu", "Heading 1")
    for line in [
        "Đọc phần tổng quan file để biết file nằm ở đâu trong pipeline phát hiện gian lận.",
        "Đọc bảng thành phần để tìm nhanh hàm/lớp cần giải thích.",
        "Ở từng hàm, phần quan trọng nhất là Vai trò, Luồng xử lý chính và Ghi chú khi bảo vệ.",
        "Code gốc được giữ lại theo block hàm/lớp; với hàm quá dài, tài liệu lược bớt đoạn giữa để Word dễ đọc.",
    ]:
        para(doc, f"- {line}")

    para(doc, "Phạm vi file", "Heading 1")
    summary = doc.add_table(rows=1, cols=4)
    summary.style = "Table Grid"
    header(summary, ["Nhóm", "File", "Số dòng", "Vai trò"])
    for group, rel, role in FILES:
        path = ROOT / rel
        count = len(read_lines(path)) if path.exists() else 0
        row = summary.add_row().cells
        cell_text(row[0], group)
        cell_text(row[1], rel)
        cell_text(row[2], count)
        cell_text(row[3], role)

    para(doc, "Luồng tổng quát theo service", "Heading 1")
    for line in [
        "AI service, đặc biệt `proctor.py`, phân tích frame camera và hành vi để tạo signal ban đầu.",
        "`FraudSignalService` chuẩn hóa và lưu signal thành dữ liệu nghiệp vụ có severity, confidence, riskImpact và evidence.",
        "`RiskScoringService` gom signal theo nhóm, áp dedup/cap, tính score và RiskLevel.",
        "`FraudWarningService` tạo warning dễ đọc cho giáo viên; `FraudAnalysisService` làm phân tích hậu kiểm theo similarity, timing, IP, thống kê và behavior.",
        "`MLRiskScoringService` và `FraudModelTrainer` là lớp mở rộng cho feature/anomaly/model.",
    ]:
        para(doc, f"- {line}")

    for group, rel, role in FILES:
        add_file_section(doc, group, rel, role)

    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(OUT)
