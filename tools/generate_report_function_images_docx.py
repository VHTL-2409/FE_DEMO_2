from __future__ import annotations

import re
import subprocess
import tempfile
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
DIAGRAM_ROOT = ROOT / "docs" / "uml-diagrams"
PLANTUML_ROOT = DIAGRAM_ROOT / "plantuml_2"
IMAGE_ROOT = PLANTUML_ROOT / "_images"
MANIFEST = IMAGE_ROOT / "FIGURE_NAMES.md"
DIAGRAM_DIR = ROOT / "docs" / "diagrams"
RENDER_SCRIPT = DIAGRAM_ROOT / "render-plantuml.ps1"
PLANTUML_JAR = ROOT / "plantuml-1.2026.2.jar"
OUTPUT_DOCX = ROOT / "baocao" / "Phu-luc-hinh-anh-theo-chuc-nang-A4.docx"

A4_PORTRAIT = (21.0, 29.7)
A4_LANDSCAPE = (29.7, 21.0)
MARGIN_CM = 1.0
CAPTION_SPACE_IN = 0.46
SECTION_HEADING_SPACE_IN = 0.58
SUBHEADING_SPACE_IN = 0.82
RENDER_SCALE = 2

ROLE_LABELS = {
    "Common": "Chung",
    "Student": "Sinh viên",
    "Teacher": "Giảng viên",
    "Admin": "Quản trị viên",
}

DOMAIN_LABELS = {
    "Auth": "Xác thực",
    "Class": "Lớp học",
    "Exam": "Kỳ thi",
    "Results": "Kết quả",
    "Analytics": "Thống kê",
    "Dashboard": "Dashboard",
    "User": "Người dùng",
    "Question": "Câu hỏi",
    "Proctoring": "Giám sát gian lận",
}

TITLE_OVERRIDES = {
    "00-overview/Architecture-FE-DEMO-v2.png": "Kiến trúc tổng thể FE_DEMO",
    "00-overview/SystemOverview.png": "Tổng quan hệ thống",
    "00-overview/UseCase-Tong-Quat-FE-DEMO.png": "Sơ đồ Use Case tổng quát FE_DEMO",
    "00-overview/PermissionMatrix.png": "Ma trận phân quyền",
}


@dataclass(frozen=True)
class ManifestRow:
    rel_image: str
    title: str
    source: str


@dataclass(frozen=True)
class FigureItem:
    title: str
    rel_image: str | None = None
    source_image: Path | None = None


@dataclass(frozen=True)
class Subsection:
    title: str
    figures: list[FigureItem]


@dataclass(frozen=True)
class Section:
    number: int
    title: str
    subsections: list[Subsection]


def read_manifest() -> dict[str, ManifestRow]:
    pattern = re.compile(r"^\|\s*(\d+)\s*\|\s*`([^`]+)`\s*\|\s*(.*?)\s*\|\s*`([^`]+)`\s*\|")
    rows: dict[str, ManifestRow] = {}
    for line in MANIFEST.read_text(encoding="utf-8").splitlines():
        match = pattern.match(line)
        if not match:
            continue
        image_path = ROOT / match.group(2)
        rel_image = image_path.relative_to(IMAGE_ROOT).as_posix()
        rows[rel_image] = ManifestRow(
            rel_image=rel_image,
            title=match.group(3).strip(),
            source=match.group(4),
        )
    if not rows:
        raise RuntimeError(f"No figure rows found in {MANIFEST}")
    return rows


def clean_title(title: str) -> str:
    title = re.sub(r"^Hình\s*-\s*", "", title).strip()
    title = re.sub(r"\s+", " ", title)
    return title


def title_from_path(rel_image: str, manifest_rows: dict[str, ManifestRow]) -> str:
    if rel_image in manifest_rows:
        return clean_title(manifest_rows[rel_image].title)
    if rel_image in TITLE_OVERRIDES:
        return TITLE_OVERRIDES[rel_image]

    stem = Path(rel_image).stem
    match = re.fullmatch(r"(Common|Student|Teacher|Admin)-([A-Za-z]+)-UseCase", stem)
    if match:
        role = ROLE_LABELS.get(match.group(1), match.group(1))
        domain = DOMAIN_LABELS.get(match.group(2), match.group(2))
        return f"{role} - {domain} - Sơ đồ Use Case"

    match = re.fullmatch(r"(Student|Teacher|Admin)-UseCase", stem)
    if match:
        role = ROLE_LABELS.get(match.group(1), match.group(1))
        return f"{role} - Sơ đồ Use Case tổng quan"

    match = re.fullmatch(r"Component-Map-(Student|Teacher|Admin)", stem)
    if match:
        role = ROLE_LABELS.get(match.group(1), match.group(1))
        return f"{role} - Bản đồ component"

    match = re.fullmatch(r"Hinh-(Student|Teacher|Admin)-Activity-Flow", stem)
    if match:
        role = ROLE_LABELS.get(match.group(1), match.group(1))
        return f"{role} - Biểu đồ hoạt động luồng tổng quan"

    return re.sub(r"[-_]+", " ", stem).strip()


def image_sort_key(path: Path) -> tuple[int, int, str]:
    name = path.name.lower()
    part = re.search(r"part-(\d+)([a-z]?)", name)
    part_no = int(part.group(1)) if part else 0
    part_letter = part.group(2) if part else ""
    part_offset = ord(part_letter) - ord("a") + 1 if part_letter else 0
    part_rank = part_no * 10 + part_offset

    if "usecase" in name or "use-case" in name:
        priority = 0
    elif "component-map" in name:
        priority = 1
    elif "activity-flow" in name:
        priority = 2
    elif "activity-overview" in name or re.search(r"activity\.png$", name):
        priority = 3
    elif "activity-part" in name:
        priority = 4
    elif "sequence-overview" in name or re.search(r"sequence\.png$", name):
        priority = 5
    elif "sequence-part" in name:
        priority = 6
    else:
        priority = 9

    return (priority, part_rank, name)


def plantuml_figure(rel_image: str, manifest_rows: dict[str, ManifestRow]) -> FigureItem:
    path = IMAGE_ROOT / rel_image
    if not path.exists():
        raise FileNotFoundError(path)
    return FigureItem(title=title_from_path(rel_image, manifest_rows), rel_image=rel_image)


def collect_dir(rel_dir: str, manifest_rows: dict[str, ManifestRow]) -> list[FigureItem]:
    path = IMAGE_ROOT / rel_dir
    if not path.exists():
        raise FileNotFoundError(path)
    figures = [
        p.relative_to(IMAGE_ROOT).as_posix()
        for p in sorted(path.glob("*.png"), key=image_sort_key)
    ]
    return [plantuml_figure(rel, manifest_rows) for rel in figures]


def collect_list(rel_images: list[str], manifest_rows: dict[str, ManifestRow]) -> list[FigureItem]:
    return [plantuml_figure(rel, manifest_rows) for rel in rel_images]


def ai_figure(rel_path: str, title: str) -> FigureItem:
    path = DIAGRAM_DIR / rel_path
    if not path.exists():
        raise FileNotFoundError(path)
    return FigureItem(title=title, source_image=path)


def build_sections(manifest_rows: dict[str, ManifestRow]) -> list[Section]:
    return [
        Section(
            number=1,
            title="TỔNG QUAN HỆ THỐNG",
            subsections=[
                Subsection(
                    title="1.1. Kiến trúc và bối cảnh hệ thống",
                    figures=collect_list(
                        [
                            "00-overview/Architecture-FE-DEMO-v2.png",
                            "00-overview/SystemOverview.png",
                            "00-overview/UseCase-Tong-Quat-FE-DEMO.png",
                            "00-overview/PermissionMatrix.png",
                        ],
                        manifest_rows,
                    ),
                ),
            ],
        ),
        Section(
            number=2,
            title="XÁC THỰC CHUNG",
            subsections=[
                Subsection(
                    title="2.1. Đăng nhập, đăng ký, quên mật khẩu, đổi mật khẩu, đăng xuất",
                    figures=collect_dir("00-common/auth", manifest_rows),
                ),
            ],
        ),
        Section(
            number=3,
            title="SINH VIÊN",
            subsections=[
                Subsection(
                    title="3.1. Tổng quan chức năng sinh viên",
                    figures=collect_list(
                        [
                            "01-student/Student-UseCase.png",
                            "01-student/Component-Map-Student.png",
                            "01-student/Hinh-Student-Activity-Flow.png",
                        ],
                        manifest_rows,
                    ),
                ),
                Subsection(
                    title="3.2. Đăng nhập, đăng ký, quên mật khẩu, đăng xuất",
                    figures=collect_dir("01-student/auth", manifest_rows),
                ),
                Subsection(
                    title="3.3. Tham gia kỳ thi",
                    figures=collect_dir("01-student/exam", manifest_rows),
                ),
                Subsection(
                    title="3.4. Quản lý lớp học",
                    figures=collect_dir("01-student/class", manifest_rows),
                ),
                Subsection(
                    title="3.5. Phân tích kết quả",
                    figures=collect_dir("01-student/results", manifest_rows),
                ),
            ],
        ),
        Section(
            number=4,
            title="GIẢNG VIÊN",
            subsections=[
                Subsection(
                    title="4.1. Tổng quan chức năng giảng viên",
                    figures=collect_list(
                        [
                            "02-teacher/Teacher-UseCase.png",
                            "02-teacher/Component-Map-Teacher.png",
                            "02-teacher/Hinh-Teacher-Activity-Flow.png",
                        ],
                        manifest_rows,
                    ),
                ),
                Subsection(
                    title="4.2. Đăng nhập, đăng ký, quên mật khẩu, đăng xuất",
                    figures=collect_dir("02-teacher/auth", manifest_rows),
                ),
                Subsection(
                    title="4.3. Dashboard và phân tích thống kê",
                    figures=collect_dir("02-teacher/analytics", manifest_rows),
                ),
                Subsection(
                    title="4.4. Quản lý lớp học",
                    figures=collect_dir("02-teacher/class", manifest_rows),
                ),
                Subsection(
                    title="4.5. Quản lý kỳ thi",
                    figures=collect_dir("02-teacher/exam", manifest_rows),
                ),
                Subsection(
                    title="4.6. Quản lý câu hỏi và import",
                    figures=collect_dir("02-teacher/question", manifest_rows),
                ),
                Subsection(
                    title="4.7. Giám sát kỳ thi",
                    figures=collect_dir("02-teacher/proctoring", manifest_rows),
                ),
            ],
        ),
        Section(
            number=5,
            title="QUẢN TRỊ VIÊN",
            subsections=[
                Subsection(
                    title="5.1. Tổng quan chức năng quản trị viên",
                    figures=collect_list(
                        [
                            "03-admin/Admin-UseCase.png",
                            "03-admin/Component-Map-Admin.png",
                            "03-admin/Hinh-Admin-Activity-Flow.png",
                        ],
                        manifest_rows,
                    ),
                ),
                Subsection(
                    title="5.2. Xác thực quản trị viên",
                    figures=collect_dir("03-admin/auth", manifest_rows),
                ),
                Subsection(
                    title="5.3. Dashboard quản trị",
                    figures=collect_dir("03-admin/dashboard", manifest_rows),
                ),
                Subsection(
                    title="5.4. Quản lý người dùng",
                    figures=collect_dir("03-admin/user", manifest_rows),
                ),
            ],
        ),
        Section(
            number=6,
            title="MÔ-ĐUN GIÁM SÁT VÀ AI",
            subsections=[
                Subsection(
                    title="6.1. Kiến trúc và luồng tổng thể",
                    figures=[
                        ai_figure("architecture.png", "Kiến trúc hệ thống giám sát và AI"),
                        ai_figure("client_capture.png", "Kiến trúc client thu thập tín hiệu"),
                        ai_figure("backend_processing.png", "Xử lý tín hiệu ở backend và AI"),
                        ai_figure("event_flow.png", "Luồng sự kiện realtime"),
                    ],
                ),
                Subsection(
                    title="6.2. Thu thập và suy diễn tín hiệu",
                    figures=[
                        ai_figure("00_overview.png", "Tổng quan pipeline giám sát"),
                        ai_figure("01_event_collection.png", "Thu thập sự kiện giám sát"),
                        ai_figure("signals.png", "Tín hiệu suy diễn phục vụ phát hiện gian lận"),
                        ai_figure("02_fraud_detection.png", "Quy trình phát hiện gian lận"),
                    ],
                ),
                Subsection(
                    title="6.3. Chấm điểm rủi ro và realtime",
                    figures=[
                        ai_figure("03_risk_scoring.png", "Cơ chế chấm điểm rủi ro"),
                        ai_figure("04_realtime_monitoring.png", "Theo dõi realtime và cảnh báo"),
                        ai_figure("risk_scoring.png", "Cơ chế chấm điểm rủi ro chi tiết"),
                    ],
                ),
                Subsection(
                    title="6.4. Can thiệp và tín hiệu bổ trợ",
                    figures=[
                        ai_figure("05_teacher_intervention.png", "Giáo viên can thiệp trong phiên thi"),
                        ai_figure("06_duplicate_ip.png", "Phát hiện trùng địa chỉ IP"),
                        ai_figure("07_answer_similarity.png", "Phân tích độ tương đồng câu trả lời"),
                        ai_figure("08_device_fingerprint.png", "Dấu vân tay thiết bị"),
                    ],
                ),
            ],
        ),
    ]


def add_scale_directive(text: str) -> str:
    if re.search(r"(?m)^scale\s+\d+", text):
        return text

    lines = text.splitlines()
    insert_at = 1
    for index, line in enumerate(lines[:20]):
        stripped = line.strip()
        if stripped.startswith("title "):
            insert_at = index + 1
            break
        if stripped.startswith("!theme"):
            insert_at = index + 1

    lines.insert(insert_at, f"scale {RENDER_SCALE}")
    return "\n".join(lines) + "\n"


def prepare_scaled_sources(rows: dict[str, ManifestRow], input_root: Path) -> None:
    for row in rows.values():
        source = ROOT / row.source
        if not source.exists():
            raise FileNotFoundError(source)
        rel_source = source.relative_to(PLANTUML_ROOT)
        target = input_root / rel_source
        target.parent.mkdir(parents=True, exist_ok=True)
        text = source.read_text(encoding="utf-8-sig", errors="ignore")
        target.write_text(add_scale_directive(text), encoding="utf-8", newline="\n")


def render_scaled_images(input_root: Path, output_root: Path) -> None:
    if not PLANTUML_JAR.exists():
        raise FileNotFoundError(PLANTUML_JAR)

    input_arg = input_root.relative_to(DIAGRAM_ROOT).as_posix()
    output_arg = output_root.relative_to(DIAGRAM_ROOT).as_posix()
    command = [
        "powershell",
        "-ExecutionPolicy",
        "Bypass",
        "-File",
        str(RENDER_SCRIPT),
        "-InputDir",
        input_arg,
        "-OutputDir",
        output_arg,
        "-Format",
        "png",
        "-PlantUmlJar",
        str(PLANTUML_JAR),
    ]
    subprocess.run(command, cwd=ROOT, check=True)


def set_run_font(run, *, bold: bool = False, italic: bool = False, size: float = 11) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if run._element.rPr is not None and run._element.rPr.rFonts is not None:
        for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
            run._element.rPr.rFonts.set(qn(f"w:{attr}"), "Times New Roman")


def set_style_defaults(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)

    for style_name, size, bold in [
        ("Heading 1", 15, True),
        ("Heading 2", 12, True),
        ("Caption", 10, False),
    ]:
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = bold
        if style_name == "Caption":
            style.font.italic = False


def set_section(section, orientation: WD_ORIENT) -> None:
    if orientation == WD_ORIENT.LANDSCAPE:
        width_cm, height_cm = A4_LANDSCAPE
    else:
        width_cm, height_cm = A4_PORTRAIT

    section.orientation = orientation
    section.page_width = Cm(width_cm)
    section.page_height = Cm(height_cm)
    section.top_margin = Cm(MARGIN_CM)
    section.bottom_margin = Cm(MARGIN_CM)
    section.left_margin = Cm(MARGIN_CM)
    section.right_margin = Cm(MARGIN_CM)
    section.header_distance = Cm(0.5)
    section.footer_distance = Cm(0.5)


def available_box(
    orientation: WD_ORIENT,
    *,
    reserve_section_heading: bool,
    reserve_subheading: bool,
) -> tuple[float, float]:
    if orientation == WD_ORIENT.LANDSCAPE:
        width_cm, height_cm = A4_LANDSCAPE
    else:
        width_cm, height_cm = A4_PORTRAIT

    max_width_in = (width_cm - MARGIN_CM * 2) / 2.54
    max_height_in = (height_cm - MARGIN_CM * 2) / 2.54 - CAPTION_SPACE_IN
    if reserve_section_heading:
        max_height_in -= SECTION_HEADING_SPACE_IN
    if reserve_subheading:
        max_height_in -= SUBHEADING_SPACE_IN
    return max_width_in, max_height_in


def fit_size(
    width_px: int,
    height_px: int,
    orientation: WD_ORIENT,
    *,
    reserve_section_heading: bool,
    reserve_subheading: bool,
) -> tuple[float, float]:
    max_width_in, max_height_in = available_box(
        orientation,
        reserve_section_heading=reserve_section_heading,
        reserve_subheading=reserve_subheading,
    )
    scale = min(max_width_in / width_px, max_height_in / height_px)
    return width_px * scale, height_px * scale


def choose_orientation(
    width_px: int,
    height_px: int,
    *,
    reserve_section_heading: bool,
    reserve_subheading: bool,
) -> WD_ORIENT:
    portrait = fit_size(
        width_px,
        height_px,
        WD_ORIENT.PORTRAIT,
        reserve_section_heading=reserve_section_heading,
        reserve_subheading=reserve_subheading,
    )
    landscape = fit_size(
        width_px,
        height_px,
        WD_ORIENT.LANDSCAPE,
        reserve_section_heading=reserve_section_heading,
        reserve_subheading=reserve_subheading,
    )
    return WD_ORIENT.LANDSCAPE if landscape[0] * landscape[1] > portrait[0] * portrait[1] else WD_ORIENT.PORTRAIT


def add_section_heading(document: Document, title: str) -> None:
    paragraph = document.add_paragraph(style="Heading 1")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(title)
    set_run_font(run, bold=True, size=15)


def add_subsection_heading(document: Document, title: str) -> None:
    paragraph = document.add_paragraph(style="Heading 2")
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(title)
    set_run_font(run, bold=True, size=12)


def add_figure_page(
    document: Document,
    image_path: Path,
    caption_no: str,
    title: str,
    *,
    section_title: str | None,
    subsection_title: str | None,
) -> None:
    with Image.open(image_path) as image:
        width_px, height_px = image.size

    reserve_section_heading = section_title is not None
    reserve_subheading = subsection_title is not None
    orientation = choose_orientation(
        width_px,
        height_px,
        reserve_section_heading=reserve_section_heading,
        reserve_subheading=reserve_subheading,
    )
    set_section(document.sections[-1], orientation)
    width_in, _ = fit_size(
        width_px,
        height_px,
        orientation,
        reserve_section_heading=reserve_section_heading,
        reserve_subheading=reserve_subheading,
    )

    if section_title:
        add_section_heading(document, section_title)
    if subsection_title:
        add_subsection_heading(document, subsection_title)

    image_paragraph = document.add_paragraph()
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.paragraph_format.space_before = Pt(0)
    image_paragraph.paragraph_format.space_after = Pt(2)
    image_paragraph.paragraph_format.keep_together = True
    run = image_paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width_in))

    caption = document.add_paragraph(style="Caption")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(0)
    caption.paragraph_format.space_after = Pt(0)
    caption.paragraph_format.line_spacing = 1.0
    caption.paragraph_format.keep_together = True
    run = caption.add_run(f"Hình {caption_no}. {title}")
    set_run_font(run, size=10)


def iter_figures(sections: list[Section]):
    for section in sections:
        for subsection in section.subsections:
            for figure in subsection.figures:
                yield figure


def selected_manifest_rows(sections: list[Section], manifest_rows: dict[str, ManifestRow]) -> dict[str, ManifestRow]:
    selected: dict[str, ManifestRow] = {}
    for figure in iter_figures(sections):
        if figure.rel_image and figure.rel_image in manifest_rows:
            selected[figure.rel_image] = manifest_rows[figure.rel_image]
    return selected


def resolve_image_path(figure: FigureItem, rendered_root: Path, converted_ai: dict[Path, Path]) -> Path:
    if figure.rel_image:
        rendered = rendered_root / figure.rel_image
        if rendered.exists():
            return rendered
        original = IMAGE_ROOT / figure.rel_image
        if original.exists():
            return original
        raise FileNotFoundError(original)

    if not figure.source_image:
        raise RuntimeError(f"Figure has no image source: {figure.title}")
    if figure.source_image.suffix.lower() == ".svg":
        raise RuntimeError(f"SVG requires a PNG alternative: {figure.source_image}")
    return figure.source_image


def add_title_page(document: Document, section_count: int, figure_count: int) -> None:
    set_section(document.sections[0], WD_ORIENT.PORTRAIT)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(120)
    run = paragraph.add_run("PHỤ LỤC HÌNH ẢNH THEO CHỨC NĂNG")
    set_run_font(run, bold=True, size=17)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(12)
    run = paragraph.add_run("Căn theo các chức năng chính trong BaoCaoNCKH_final_Long2026.docx")
    set_run_font(run, size=12)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(8)
    run = paragraph.add_run(f"{section_count} nhóm chức năng, {figure_count} hình ảnh, khổ A4")
    set_run_font(run, size=11)


def build_document(sections: list[Section], rendered_root: Path, converted_ai: dict[Path, Path]) -> Path:
    document = Document()
    set_style_defaults(document)
    document.core_properties.title = "Phụ lục hình ảnh theo chức năng FE_DEMO"
    document.core_properties.subject = "PlantUML và sơ đồ AI theo nhóm chức năng báo cáo"

    total_figures = sum(len(subsection.figures) for section in sections for subsection in section.subsections)
    add_title_page(document, len(sections), total_figures)

    first_figure = True
    section_figure_counts: dict[int, int] = {}
    for section in sections:
        section_figure_counts[section.number] = 0
        first_in_section = True
        for subsection in section.subsections:
            for index, figure in enumerate(subsection.figures):
                if first_figure:
                    document.add_section(WD_SECTION_START.NEW_PAGE)
                    first_figure = False
                else:
                    document.add_section(WD_SECTION_START.NEW_PAGE)

                section_figure_counts[section.number] += 1
                caption_no = f"{section.number}.{section_figure_counts[section.number]}"
                image_path = resolve_image_path(figure, rendered_root, converted_ai)
                section_title = f"{section.number}. {section.title}" if first_in_section else None
                subsection_title = subsection.title if index == 0 else None
                add_figure_page(
                    document,
                    image_path,
                    caption_no,
                    figure.title,
                    section_title=section_title,
                    subsection_title=subsection_title,
                )
                first_in_section = False

    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    try:
        document.save(OUTPUT_DOCX)
        return OUTPUT_DOCX
    except PermissionError:
        alternate = OUTPUT_DOCX.with_name(f"{OUTPUT_DOCX.stem}-new.docx")
        document.save(alternate)
        return alternate


def main() -> None:
    manifest_rows = read_manifest()
    sections = build_sections(manifest_rows)
    rows_to_render = selected_manifest_rows(sections, manifest_rows)

    with tempfile.TemporaryDirectory(prefix=".tmp-report-function-images-", dir=DIAGRAM_ROOT) as temp_dir:
        temp_root = Path(temp_dir)
        input_root = temp_root / "input"
        rendered_root = temp_root / "output"

        prepare_scaled_sources(rows_to_render, input_root)
        render_scaled_images(input_root, rendered_root)
        output = build_document(sections, rendered_root, {})

    total = sum(len(subsection.figures) for section in sections for subsection in section.subsections)
    print(f"created {output}")
    print(f"sections {len(sections)}")
    print(f"figures {total}")


if __name__ == "__main__":
    main()
