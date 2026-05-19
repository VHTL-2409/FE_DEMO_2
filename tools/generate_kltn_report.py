from __future__ import annotations

import json
from pathlib import Path
from typing import Iterable
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE_PATH = ROOT / "baocao" / "3. Bao-Cao-Khoa-Luan-Tran-Cong-Son-1921142446.docx"
OUTPUT_PATH = ROOT / "baocao" / "Bao-cao-Khoa-luan-Thi-truc-tuyen-giam-sat-gian-lan-thong-minh.docx"

SCREENSHOT_DIR = ROOT / "baocao" / "report-assets" / "screenshots"
DIAGRAM_DIR = ROOT / "docs" / "diagrams"
UML_DIR = ROOT / "docs" / "uml-diagrams" / "plantuml_word_compact_png"


def read_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def read_package_version(path: Path) -> str:
    return read_json(path).get("version", "unknown")


def read_pom_versions(path: Path) -> dict[str, str]:
    ns = {"m": "http://maven.apache.org/POM/4.0.0"}
    root = ET.fromstring(path.read_text(encoding="utf-8", errors="ignore"))
    parent_version = root.findtext("./m:parent/m:version", default="unknown", namespaces=ns)
    java_version = root.findtext("./m:properties/m:java.version", default="unknown", namespaces=ns)
    artifact = root.findtext("./m:artifactId", default="unknown", namespaces=ns)
    project_version = root.findtext("./m:version", default="unknown", namespaces=ns)
    return {
        "spring_boot": parent_version,
        "java": java_version,
        "artifact": artifact,
        "project_version": project_version,
    }


def read_requirements(path: Path) -> list[str]:
    lines = []
    for raw in read_text(path).splitlines():
        line = raw.strip()
        if not line or line.startswith("#"):
            continue
        lines.append(line)
    return lines


def clear_body(document: Document) -> None:
    body = document._body._element
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def copy_section_metrics(dst, src) -> None:
    dst.top_margin = src.top_margin
    dst.bottom_margin = src.bottom_margin
    dst.left_margin = src.left_margin
    dst.right_margin = src.right_margin
    dst.header_distance = src.header_distance
    dst.footer_distance = src.footer_distance
    dst.page_width = src.page_width
    dst.page_height = src.page_height
    dst.gutter = src.gutter
    try:
        dst.different_first_page_header_footer = src.different_first_page_header_footer
        dst.odd_and_even_pages_header_footer = src.odd_and_even_pages_header_footer
    except Exception:
        pass


def set_core_props(document: Document) -> None:
    props = document.core_properties
    props.title = "Bao cao KLTN FE_DEMO"
    props.subject = "Thiết kế hệ thống thi trực tuyến với cơ chế giám sát và phát hiện gian lận thông minh"
    props.author = "OpenAI"
    props.comments = "Generated from local FE_DEMO project sources"
    props.keywords = "KLTN, FE_DEMO, proctoring, Vue, Spring Boot, FastAPI"


def set_style_defaults(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(13)

    for style_name, size, bold, italic, align in [
        ("Heading 1", 16, True, False, WD_ALIGN_PARAGRAPH.CENTER),
        ("Heading 2", 14, True, False, WD_ALIGN_PARAGRAPH.LEFT),
        ("Heading 3", 13, True, False, WD_ALIGN_PARAGRAPH.LEFT),
        ("Caption", 12, False, True, WD_ALIGN_PARAGRAPH.CENTER),
    ]:
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.italic = italic
        try:
            style.paragraph_format.alignment = align
        except Exception:
            pass


def set_paragraph_format(paragraph, *, align=None, first_line_cm: float | None = 0.75, after_pt: float = 6, before_pt: float = 0, line_spacing: float = 1.5) -> None:
    fmt = paragraph.paragraph_format
    fmt.line_spacing = line_spacing
    fmt.space_after = Pt(after_pt)
    fmt.space_before = Pt(before_pt)
    if first_line_cm is not None:
        fmt.first_line_indent = Cm(first_line_cm)
    if align is not None:
        paragraph.alignment = align


def set_run_font(run, *, bold: bool = False, italic: bool = False, size: int = 13) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    rfonts = run._element.rPr.rFonts if run._element.rPr is not None and run._element.rPr.rFonts is not None else None
    if rfonts is not None:
        for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
            rfonts.set(qn(f"w:{attr}"), "Times New Roman")


def add_run_text(paragraph, text: str, *, bold: bool = False, italic: bool = False, size: int = 13) -> None:
    run = paragraph.add_run(text)
    set_run_font(run, bold=bold, italic=italic, size=size)


def add_paragraph(document: Document, text: str = "", *, style: str = "Normal", align=None, first_line_cm: float | None = 0.75, bold: bool = False, italic: bool = False, size: int = 13) -> object:
    paragraph = document.add_paragraph(style=style)
    set_paragraph_format(paragraph, align=align, first_line_cm=first_line_cm)
    if text:
        add_run_text(paragraph, text, bold=bold, italic=italic, size=size)
    return paragraph


def add_heading(document: Document, text: str, level: int, *, first: bool = False) -> object:
    if not first:
        document.add_page_break()
    paragraph = document.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.page_break_before = False
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    add_run_text(paragraph, text, bold=True, size=16 if level == 1 else 14 if level == 2 else 13)
    return paragraph


def add_field(paragraph, instruction: str, placeholder: str = "") -> None:
    run = paragraph.add_run()
    r = run._r
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    r.append(begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    r.append(instr)

    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    r.append(sep)

    result = paragraph.add_run(placeholder)
    set_run_font(result)

    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def add_toc(document: Document) -> None:
    paragraph = document.add_paragraph()
    set_paragraph_format(paragraph, align=WD_ALIGN_PARAGRAPH.LEFT, first_line_cm=None, after_pt=0)
    add_field(paragraph, 'TOC \\o "1-3" \\h \\z \\u')


def add_figure_list(document: Document) -> None:
    paragraph = document.add_paragraph()
    set_paragraph_format(paragraph, align=WD_ALIGN_PARAGRAPH.LEFT, first_line_cm=None, after_pt=0)
    add_field(paragraph, 'TOC \\h \\z \\c "Hinh"')


def add_table_list(document: Document) -> None:
    paragraph = document.add_paragraph()
    set_paragraph_format(paragraph, align=WD_ALIGN_PARAGRAPH.LEFT, first_line_cm=None, after_pt=0)
    add_field(paragraph, 'TOC \\h \\z \\c "Bang"')


def add_caption(document: Document, chapter_no: int, caption_no: int, text: str, *, kind: str = "Hinh", visible_label: str = "Hình", restart: bool = False) -> object:
    paragraph = document.add_paragraph(style="Caption")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.keep_with_next = True
    prefix = f"{visible_label} {chapter_no}."
    add_run_text(paragraph, prefix, bold=True, italic=False, size=12)
    instruction = f"SEQ {kind}"
    if restart:
        instruction += " \\r 1"
    add_field(paragraph, instruction, placeholder=str(caption_no))
    add_run_text(paragraph, f". {text}", bold=False, italic=False, size=12)
    return paragraph


def add_image(document: Document, image_path: Path, *, width_in: float = 6.2, align=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    if not image_path.exists():
        raise FileNotFoundError(image_path)
    paragraph = document.add_paragraph()
    paragraph.alignment = align
    paragraph.paragraph_format.keep_together = True
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width_in))


def add_figure(document: Document, chapter_no: int, caption_no: int, image_path: Path, caption: str, *, restart: bool = False, width_in: float = 6.2) -> None:
    add_image(document, image_path, width_in=width_in)
    add_caption(document, chapter_no, caption_no, caption, kind="Hinh", visible_label="Hình", restart=restart)


def add_table_caption(document: Document, chapter_no: int, caption_no: int, caption: str, *, restart: bool = False) -> object:
    return add_caption(document, chapter_no, caption_no, caption, kind="Bang", visible_label="Bảng", restart=restart)


def set_table_style(table) -> None:
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True


def add_table(document: Document, headers: list[str], rows: list[list[str]], *, caption: str | None = None, chapter_no: int | None = None, caption_no: int | None = None, restart: bool = False) -> object | None:
    if caption and chapter_no is not None and caption_no is not None:
        add_table_caption(document, chapter_no, caption_no, caption, restart=restart)
    table = document.add_table(rows=1, cols=len(headers))
    set_table_style(table)
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].text = header
        for p in hdr[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                set_run_font(r, bold=True, size=11)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            for p in cells[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    set_run_font(r, size=11)
    document.add_paragraph()
    return table


def add_text_block(document: Document, paragraphs: Iterable[str], *, first_line_cm: float = 0.75) -> None:
    for para in paragraphs:
        paragraph = add_paragraph(document, para, first_line_cm=first_line_cm)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_bold_lead(document: Document, label: str, text: str) -> None:
    paragraph = document.add_paragraph()
    set_paragraph_format(paragraph, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_run_text(paragraph, label, bold=True)
    add_run_text(paragraph, text)


def build_intro(document: Document) -> None:
    add_heading(document, "LỜI CẢM ƠN", 1, first=True)
    add_text_block(document, [
        "Em xin gửi lời cảm ơn tới giảng viên hướng dẫn [GIẢNG VIÊN HƯỚNG DẪN] đã kiên nhẫn góp ý, định hướng và hỗ trợ em trong suốt quá trình thực hiện khóa luận.",
        "Em cũng trân trọng cảm ơn thầy cô trong khoa, bạn bè và gia đình đã tạo điều kiện về thời gian, tài liệu và tinh thần để em hoàn thành báo cáo này.",
        "Do thời gian và kinh nghiệm còn hạn chế, báo cáo khó tránh khỏi thiếu sót. Em rất mong nhận được nhận xét để tiếp tục hoàn thiện hệ thống và nội dung trình bày.",
    ])

    add_heading(document, "LỜI CAM ĐOAN", 1)
    add_text_block(document, [
        "Em xin cam đoan đây là kết quả nghiên cứu và triển khai của riêng em trên dự án FE_DEMO. Các nội dung phân tích, thiết kế, hiện thực và kiểm thử trong báo cáo đều dựa trên mã nguồn, tài liệu nội bộ và quá trình chạy thực tế của hệ thống.",
        "Những thông tin sử dụng từ tài liệu bên ngoài đều được chọn lọc và ghi nhận ở phần tài liệu tham khảo. Em chịu trách nhiệm hoàn toàn trước nhà trường về tính trung thực của báo cáo.",
    ])
    add_paragraph(document, "Sinh viên: [HỌ VÀ TÊN SINH VIÊN]", first_line_cm=None)
    add_paragraph(document, "MSSV: [MSSV]", first_line_cm=None)
    add_paragraph(document, "Lớp: [LỚP]", first_line_cm=None)
    add_paragraph(document, "Giảng viên hướng dẫn: [GIẢNG VIÊN HƯỚNG DẪN]", first_line_cm=None)
    add_paragraph(document, "Ngày .... tháng .... năm 2026", align=WD_ALIGN_PARAGRAPH.RIGHT, first_line_cm=None)
    add_paragraph(document, "Ký tên", align=WD_ALIGN_PARAGRAPH.RIGHT, first_line_cm=None)

    add_heading(document, "DANH MỤC HÌNH ẢNH", 1)
    add_figure_list(document)

    add_heading(document, "DANH MỤC BẢNG BIỂU", 1)
    add_table_list(document)

    add_heading(document, "MỤC LỤC", 1)
    add_toc(document)


def build_opening(document: Document) -> None:
    add_heading(document, "MỞ ĐẦU", 1)
    add_text_block(document, [
        "Đề tài hướng tới xây dựng một hệ thống thi trực tuyến có khả năng giám sát chủ động, phát hiện gian lận theo thời gian thực và hỗ trợ giảng viên theo dõi tình trạng làm bài của người học trên một giao diện thống nhất.",
        "Bối cảnh thực tế cho thấy nhiều nền tảng thi trực tuyến chỉ dừng ở mức phát đề và nộp bài, trong khi nhu cầu kiểm soát hành vi bất thường, đồng bộ trạng thái camera/microphone, ghi nhận dấu hiệu gian lận và cảnh báo tức thời ngày càng quan trọng.",
        "Từ yêu cầu đó, FE_DEMO được phát triển như một hệ thống thực nghiệm tương đối đầy đủ, có frontend Vue 3/Vite, backend Spring Boot 4, dịch vụ AI FastAPI và cơ sở dữ liệu PostgreSQL để mô phỏng một quy trình thi, giám sát và phân tích rủi ro sát với thực tế.",
    ])
    add_paragraph(document, "Mục tiêu của khóa luận là thiết kế và triển khai luồng thi trực tuyến hoàn chỉnh cho ba vai trò chính: sinh viên, giảng viên và quản trị viên. Hệ thống phải đáp ứng việc đăng nhập, tạo đề, nhập câu hỏi, tham gia thi, theo dõi trực tiếp, chấm điểm rủi ro, xử lý cảnh báo và quản trị dữ liệu lớp học.", first_line_cm=0.75)
    add_paragraph(document, "Phạm vi thực hiện tập trung vào các chức năng cốt lõi của dự án thật đang chạy trên repo: xác thực người dùng, quản lý đề thi, làm bài, giám sát camera/mic, nhập câu hỏi từ nhiều định dạng, quản lý lớp học, dashboard quản trị và các dịch vụ AI hỗ trợ OCR, phân tích ảnh, sinh câu hỏi, đánh giá tự luận và chấm điểm rủi ro.", first_line_cm=0.75)
    add_paragraph(document, "Báo cáo được trình bày theo ba chương: Chương I giới thiệu công nghệ và môi trường phát triển; Chương II phân tích và thiết kế hệ thống; Chương III mô tả cài đặt, kiểm thử và các kết quả chạy thật trên môi trường local.", first_line_cm=0.75)


def build_chapter_1(document: Document, versions: dict[str, str], requirements: list[str], project_info: dict[str, str]) -> None:
    add_heading(document, "CHƯƠNG I: TỔNG QUAN VỀ CÔNG CỤ VÀ MÔI TRƯỜNG PHÁT TRIỂN", 1)
    add_heading(document, "1.1. Tổng quan dự án", 2)
    add_text_block(document, [
        "FE_DEMO là hệ thống thi trực tuyến có cơ chế giám sát gian lận thông minh, được tổ chức theo kiến trúc nhiều tầng. Frontend chịu trách nhiệm giao diện cho sinh viên, giảng viên và quản trị viên; backend cung cấp API nghiệp vụ, xác thực và xử lý miền; dịch vụ AI xử lý camera, OCR, gợi ý nội dung và đánh giá rủi ro; PostgreSQL lưu toàn bộ trạng thái thi, dữ liệu câu hỏi, lớp học và log giám sát.",
        "Mã nguồn cho thấy dự án không chỉ dừng ở chức năng thi cơ bản mà còn bao gồm theo dõi realtime qua WebSocket, kiểm tra camera/microphone, ghi nhận sự kiện gian lận, tổng hợp điểm rủi ro và các tính năng hỗ trợ giảng viên can thiệp tức thời.",
    ])

    add_heading(document, "1.2. Công nghệ frontend", 2)
    add_text_block(document, [
        f"Frontend nằm trong thư mục `demo/`, khai báo package version {project_info['fe_package_version']} và sử dụng Vue {project_info['vue_version']} cùng Vite {project_info['vite_version']}. Bộ công cụ này phù hợp với ứng dụng SPA vì tốc độ build nhanh, hỗ trợ HMR tốt và dễ cấu trúc theo router, store và composable.",
        "Các thư viện chính gồm Pinia cho state management, Vue Router cho định tuyến, PrimeVue cho thành phần giao diện, ECharts cho biểu đồ, KaTeX cho hiển thị công thức toán học, lucide-vue-next cho biểu tượng, xlsx cho import/export dữ liệu và STOMP/SockJS cho kết nối WebSocket.",
    ])
    add_table(document, [
        "Thành phần",
        "Phiên bản / ghi chú",
        "Vai trò",
    ], [
        ["demo/package.json", project_info["fe_package_version"], "Khai báo gói frontend"],
        ["Vue", project_info["vue_version"], "Framework giao diện"],
        ["Vite", project_info["vite_version"], "Bundler và dev server"],
        ["Pinia", project_info["pinia_version"], "Quản lý state"],
        ["Vue Router", project_info["vue_router_version"], "Định tuyến trang"],
        ["PrimeVue", project_info["primevue_version"], "UI component"],
        ["ECharts", project_info["echarts_version"], "Biểu đồ thống kê"],
        ["xlsx", project_info["xlsx_version"], "Xử lý file Excel"],
    ], caption="Công nghệ frontend và vai trò", chapter_no=1, caption_no=1, restart=True)

    add_heading(document, "1.3. Công nghệ backend", 2)
    add_text_block(document, [
        f"Backend đặt trong thư mục `BE/`, dùng Spring Boot {versions['spring_boot']} và Java {versions['java']}. Hệ thống được xây dựng theo mô hình REST API, kết hợp Spring Security, OAuth2 client, JPA, WebSocket, validation, mail và JWT để phục vụ đồng thời xác thực, nghiệp vụ thi và giám sát realtime.",
        "Các thư viện hỗ trợ như Apache POI, PDFBox, Commons CSV, Commons IO và Caffeine được dùng để xử lý file, import câu hỏi, sinh tài liệu, caching và một số tác vụ tiện ích khác.",
    ])
    add_table(document, [
        "Thành phần",
        "Phiên bản / ghi chú",
        "Vai trò",
    ], [
        ["Spring Boot", versions["spring_boot"], "Khung backend chính"],
        ["Java", versions["java"], "Ngôn ngữ triển khai"],
        ["JJWT", "0.12.6", "Phát hành/kiểm tra token"],
        ["Apache POI", "5.3.0", "Đọc ghi Office file"],
        ["PDFBox", "3.0.3", "Xử lý PDF"],
        ["Commons CSV", "1.13.0", "Đọc CSV import"],
        ["Commons IO", "2.20.0", "Tiện ích I/O"],
        ["Caffeine", "không khóa version", "Cache/rate limit"],
        ["PostgreSQL driver", "runtime", "Kết nối CSDL"],
    ], caption="Công nghệ backend và vai trò", chapter_no=1, caption_no=2)

    add_heading(document, "1.4. Dịch vụ AI", 2)
    add_text_block(document, [
        "Dịch vụ AI nằm trong `ai-service/`, triển khai bằng FastAPI và Uvicorn, dùng OpenCV, NumPy, Pillow, PyMuPDF, pytesseract và slowapi. File requirements cho thấy openai được khóa theo dải >=1.0.0,<2.0.0, phù hợp với các tác vụ sinh câu hỏi, chấm tự luận, phân tích hành vi và gợi ý học tập.",
        "Module `proctor.py` xử lý ảnh camera, phát hiện khuôn mặt, mắt, hướng nhìn, che khuất, spoofing và các trạng thái bất thường khác. Các module `generator.py`, `evaluator.py`, `analytics.py`, `chat.py` và `openai_client.py` mở rộng thêm năng lực sinh nội dung, đánh giá, dự đoán và chat.",
    ])
    add_table(document, [
        "Thành phần",
        "Ghi chú",
        "Vai trò",
    ], [[
        "FastAPI", "framework API", "Dịch vụ web AI"
    ], [
        "Uvicorn", "ASGI server", "Chạy FastAPI"
    ], [
        "OpenCV", "opencv-python-headless", "Xử lý ảnh và video"
    ], [
        "NumPy", "numpy", "Tính toán ma trận"
    ], [
        "PyMuPDF", "PyMuPDF", "OCR/tài liệu PDF"
    ], [
        "pytesseract", "pytesseract", "Trích xuất ký tự"
    ], [
        "openai", ">=1.0.0,<2.0.0", "Kết nối mô hình OpenAI"
    ]], caption="Thành phần dịch vụ AI", chapter_no=1, caption_no=3)

    add_heading(document, "1.5. Môi trường triển khai", 2)
    add_text_block(document, [
        "File docker-compose cho thấy hệ thống có thể chạy bằng Docker Compose với các service: PostgreSQL 16, backend, ai-service, python-parser và frontend. Cổng local chính dùng trong kiểm thử là 4173 cho frontend preview, 8082 cho backend và 8090 cho dịch vụ AI.",
        "Cấu trúc này giúp tách biệt trách nhiệm rõ ràng, thuận tiện cho việc phát triển, kiểm thử và mô phỏng triển khai thực tế. Khi cần chạy bản production, repository cũng đã có `docker-compose.prod.yml` với cấu hình tối ưu hóa hơn cho môi trường máy chủ.",
    ])
    add_table(document, [
        "Service",
        "Image / công nghệ",
        "Vai trò",
    ], [
        ["db", "postgres:16", "Lưu dữ liệu nghiệp vụ"],
        ["be", "Spring Boot 4.0.3", "API nghiệp vụ và auth"],
        ["ai-service", "FastAPI", "OCR, proctoring, AI"],
        ["python-parser", "Python parser", "Phân tích file import"],
        ["fe", "Vue 3 + Vite", "Giao diện người dùng"],
    ], caption="Môi trường chạy cục bộ và vai trò service", chapter_no=1, caption_no=4)


def build_chapter_2(document: Document, project_info: dict[str, str]) -> None:
    add_heading(document, "CHƯƠNG II: PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG", 1)
    add_heading(document, "2.1. Tác nhân và yêu cầu nghiệp vụ", 2)
    add_text_block(document, [
        "Hệ thống phục vụ ba nhóm người dùng chính: sinh viên, giảng viên và quản trị viên. Sinh viên cần tham gia thi, xem kết quả và quản lý lớp học; giảng viên cần tạo đề, import câu hỏi, mở phiên thi, giám sát trực tiếp và xử lý cảnh báo; quản trị viên cần quản lý tài khoản, lớp học và giám sát vận hành toàn hệ thống.",
        "Ngoài ra, hệ thống còn có các tác nhân kỹ thuật như backend, AI service, python-parser và WebSocket broker để đảm bảo chuỗi xử lý từ nhập liệu, phát đề, làm bài đến giám sát realtime diễn ra liên tục.",
    ])
    add_table(document, [
        "Tác nhân",
        "Nhiệm vụ chính",
        "Màn hình/miền tương tác",
    ], [
        ["Sinh viên", "Đăng nhập, vào phòng chờ, làm bài, nộp bài, xem kết quả", "/login, /student/*"],
        ["Giảng viên", "Tạo đề, import câu hỏi, tạo lớp, giám sát, cảnh báo", "/teacher/*"],
        ["Quản trị viên", "Quản lý user, exam, class, dashboard", "/admin/*"],
        ["Dịch vụ AI", "OCR, camera analysis, chat, sinh câu hỏi", "/api/v1/ai/*, /proctor/*"],
        ["Backend", "Auth, exams, monitoring, import, admin", "/api/*"],
    ], caption="Tác nhân và miền nghiệp vụ", chapter_no=2, caption_no=1, restart=True)

    add_heading(document, "2.2. Yêu cầu chức năng", 2)
    add_table(document, [
        "Nhóm chức năng",
        "Mô tả",
    ], [
        ["Xác thực", "Đăng ký, đăng nhập, refresh token, OAuth2 Google, quên mật khẩu, xác minh email"],
        ["Thi sinh viên", "Vào phòng chờ, kiểm tra camera/mic, làm bài, lưu nháp, nộp bài, xem kết quả"],
        ["Thiết kế đề", "Tạo/sửa/xóa đề thi, nhân bản đề, mở phiên thi, thiết lập trộn câu hỏi/đáp án"],
        ["Import câu hỏi", "Import từ Excel, CSV, PDF, DOCX, ảnh, Azota và preview dữ liệu"],
        ["Lớp học", "Tạo lớp, quản lý danh sách sinh viên, import danh sách, join bằng mã lớp"],
        ["Giám sát", "Theo dõi realtime, gửi warning/pause/resume/invalidate, risk score, timeline"],
        ["Quản trị", "Dashboard, quản lý user, lớp, đề thi, kiểm tra trạng thái hệ thống"],
    ], caption="Yêu cầu chức năng chính", chapter_no=2, caption_no=2)

    add_heading(document, "2.3. Yêu cầu phi chức năng", 2)
    add_table(document, [
        "Tiêu chí",
        "Kỳ vọng",
    ], [
        ["Thời gian thực", "Trạng thái camera, cảnh báo và điểm rủi ro phải cập nhật gần như ngay lập tức qua WebSocket."],
        ["An toàn", "Token hóa, phân quyền theo vai trò, giới hạn truy cập API và ghi audit log."],
        ["Mở rộng", "Tách FE/BE/AI/Parser/DB để có thể scale độc lập."],
        ["Khả dụng", "Có health check, smoke flow, log service và cấu hình Docker."],
        ["Khả năng truy vết", "Lưu monitoring_events, fraud_signals, risk_scores, audit_logs và import_jobs."],
    ], caption="Yêu cầu phi chức năng", chapter_no=2, caption_no=3)

    add_heading(document, "2.4. Kiến trúc tổng thể", 2)
    add_text_block(document, [
        "Kiến trúc FE_DEMO chia thành bốn lớp chính. Lớp giao diện Vue chịu trách nhiệm thu nhận thao tác, camera/mic và hiển thị dashboard. Lớp backend Spring Boot xử lý xác thực, nghiệp vụ thi, lớp học, import và giám sát. Lớp AI FastAPI đảm nhiệm phân tích ảnh, OCR và sinh nội dung. PostgreSQL lưu trạng thái lâu dài, còn WebSocket dùng cho sự kiện realtime.",
        "Luồng dữ liệu chính đi từ trình duyệt đến backend, sau đó sang AI service khi cần phân tích ảnh hoặc nội dung, rồi quay về backend để ghi điểm rủi ro, lưu sự kiện và phát thông báo về các màn hình giám sát.",
    ])

    for idx, (img, title, width) in enumerate([
        (DIAGRAM_DIR / "architecture.png", "Kiến trúc tổng thể hệ thống", 6.35),
        (UML_DIR / "00-overview" / "SystemContextCompact.png", "Bối cảnh hệ thống và tác nhân", 6.1),
        (UML_DIR / "00-overview" / "ModuleOverviewCompact.png", "Tổng quan module chức năng", 6.1),
        (UML_DIR / "00-overview" / "UserRoleMatrixCompact.png", "Ma trận vai trò người dùng", 6.1),
    ], start=1):
        add_figure(document, 2, idx, img, title, restart=(idx == 1), width_in=width)

    add_heading(document, "2.5. Thiết kế luồng giám sát và chấm điểm rủi ro", 2)
    add_text_block(document, [
        "Trên màn hình thi, frontend thu thập trạng thái tab, fullscreen, clipboard, camera và microphone. Khi phát hiện tín hiệu bất thường, FE gửi event về backend. Backend hợp nhất các nguồn tín hiệu, tính điểm rủi ro theo trọng số, ghi fraud_signals và monitoring_events, sau đó đẩy lại trạng thái lên topic WebSocket để giảng viên quan sát.",
        "Dữ liệu camera có thể được gửi sang AI service để phân tích mặt, mắt, hướng nhìn, che khuất và dấu hiệu spoofing. Kết quả AI được chuẩn hóa về các mức OK, WARNING, CRITICAL hoặc NO_CAMERA trước khi hiển thị trên dashboard giám sát.",
    ])

    for idx, (img, title, width) in enumerate([
        (DIAGRAM_DIR / "client_capture.png", "Thu thập trạng thái phía trình duyệt", 6.25),
        (DIAGRAM_DIR / "backend_processing.png", "Xử lý tín hiệu ở backend và AI", 6.25),
        (DIAGRAM_DIR / "event_flow.png", "Chuỗi sự kiện từ phát hiện đến xử lý", 6.25),
        (DIAGRAM_DIR / "risk_scoring.png", "Cơ chế chấm điểm rủi ro", 6.25),
        (DIAGRAM_DIR / "signals.png", "Tổng hợp các nhóm tín hiệu gian lận", 6.25),
        (UML_DIR / "12-realtime" / "WebsocketSequenceCompact.png", "Trình tự realtime WebSocket", 6.15),
    ], start=1):
        add_figure(document, 2, 4 + idx, img, title, restart=False, width_in=width)

    add_heading(document, "2.6. Thiết kế phân hệ xác thực và phân quyền", 2)
    add_text_block(document, [
        "Phân hệ xác thực dùng JWT cho API nội bộ và OAuth2 client cho đăng nhập Google. Sau khi đăng nhập, backend gắn vai trò vào token và frontend kiểm tra quyền truy cập theo route meta. Người dùng chưa có vai trò sẽ được chuyển đến màn chọn role.",
        "Các controller chính gồm `AuthController`, `ProfileController`, `HealthController` và các lớp helper/service liên quan tới token, email xác minh, reset mật khẩu và làm mới phiên đăng nhập.",
    ])
    for idx, (img, title) in enumerate([
        (UML_DIR / "01-authentication" / "AuthUseCaseCompact.png", "Use case xác thực",),
        (UML_DIR / "01-authentication" / "AuthLoginSequenceCompact.png", "Trình tự đăng nhập",),
    ], start=1):
        add_figure(document, 2, 10 + idx, img, title, restart=False, width_in=6.05)

    add_heading(document, "2.7. Thiết kế phân hệ đề thi, câu hỏi và import", 2)
    add_text_block(document, [
        "Phân hệ đề thi cho phép giảng viên tạo đề, chỉnh sửa, sao chép, tạo phiên thi và cấu hình random câu hỏi hoặc random đáp án. Phân hệ câu hỏi hỗ trợ câu hỏi trắc nghiệm, tự luận, import từ nhiều định dạng và xem trước nội dung trước khi confirm.",
        "Luồng import có thể đi qua `QuestionController`, `ImportController`, `ExamImportController` hoặc `AiController` tùy loại dữ liệu. Với file phức tạp, hệ thống dùng parser service để trích xuất cấu trúc đề và đẩy vào preview trước khi lưu chính thức.",
    ])
    for idx, (img, title) in enumerate([
        (UML_DIR / "02-exam-management" / "ExamUseCaseCompact.png", "Use case quản lý đề thi",),
        (UML_DIR / "02-exam-management" / "ExamCreateActivityCompact.png", "Hoạt động tạo đề thi",),
        (UML_DIR / "03-question-management" / "QuestionUseCaseCompact.png", "Use case quản lý câu hỏi",),
        (UML_DIR / "03-question-management" / "BulkImportSequenceCompact.png", "Trình tự import câu hỏi",),
    ], start=1):
        add_figure(document, 2, 12 + idx, img, title, restart=False, width_in=6.05)

    add_heading(document, "2.8. Thiết kế phân hệ sinh viên làm bài", 2)
    add_text_block(document, [
        "Sinh viên truy cập trang đăng nhập, vào dashboard, nhập mã đề, chờ kiểm tra điều kiện camera/microphone rồi mới được vào giao diện thi. Khi làm bài, hệ thống lưu draft định kỳ, theo dõi chuyển tab, mất focus, copy-paste, thoát fullscreen và các sự kiện bất thường khác.",
        "Khi nộp bài, backend lưu kết quả, tính tổng điểm và đồng thời giữ lại trạng thái bài thi để giảng viên hoặc sinh viên có thể xem lại lịch sử làm bài, báo cáo chi tiết và kết quả chấm.",
    ])
    for idx, (img, title) in enumerate([
        (UML_DIR / "04-student-exam" / "StudentExamUseCaseCompact.png", "Use case sinh viên làm bài",),
        (UML_DIR / "04-student-exam" / "TakeExamActivityCompact.png", "Hoạt động làm bài của sinh viên",),
        (UML_DIR / "04-student-exam" / "TakeExamSequencePart1.png", "Trình tự bắt đầu làm bài",),
    ], start=1):
        add_figure(document, 2, 16 + idx, img, title, restart=False, width_in=6.05)

    add_heading(document, "2.9. Thiết kế giám sát proctoring", 2)
    add_text_block(document, [
        "Phân hệ giám sát là lõi của đề tài. Nó kết hợp dữ liệu camera, tín hiệu trình duyệt, trạng thái bài thi, sự kiện theo thời gian thực và AI để nhận diện gian lận. Giảng viên có thể phát warning, pause, resume hoặc invalidate bài thi; hệ thống cũng có cơ chế cảnh báo tự động khi điểm rủi ro vượt ngưỡng.",
        "Các bảng `fraud_signals`, `risk_scores`, `monitoring_events`, `exam_events` và `audit_logs` tạo thành lớp lưu vết đầy đủ, giúp truy ngược từng sự kiện bất thường theo attempt và exam.",
    ])
    for idx, (img, title) in enumerate([
        (UML_DIR / "05-proctoring" / "ProctorUseCaseCompact.png", "Use case giám sát",),
        (UML_DIR / "05-proctoring" / "ProctorActivityCompact.png", "Hoạt động giám sát",),
        (UML_DIR / "05-proctoring" / "ProctorMonitoringSequenceCompact.png", "Trình tự theo dõi realtime",),
        (UML_DIR / "05-proctoring" / "ViolationDetectionSequenceCompact.png", "Phát hiện vi phạm",),
        (UML_DIR / "05-proctoring" / "ViolationStateCompact.png", "Trạng thái vi phạm",),
    ], start=1):
        add_figure(document, 2, 19 + idx, img, title, restart=False, width_in=6.05)

    add_heading(document, "2.10. Thiết kế quản lý lớp học và quản trị", 2)
    add_text_block(document, [
        "Giảng viên có thể tạo lớp, sinh mã lớp, thêm sinh viên thủ công, import danh sách sinh viên từ file và xem danh sách lớp gắn với từng đề thi. Quản trị viên có dashboard riêng để thống kê người dùng, đề thi và lớp học, đồng thời kiểm tra trạng thái hệ thống.",
        "Hai màn hình lớp học và quản trị tách riêng theo route nên frontend dễ áp điều kiện phân quyền, còn backend có các controller và service độc lập cho từng domain.",
    ])
    for idx, (img, title) in enumerate([
        (UML_DIR / "06-class-management" / "ClassUseCaseCompact.png", "Use case lớp học",),
        (UML_DIR / "06-class-management" / "ClassSequenceCompact.png", "Trình tự quản lý lớp",),
        (UML_DIR / "10-class-diagram" / "ClassEntitiesOverview.png", "Tổng quan class diagram",),
        (UML_DIR / "10-class-diagram" / "ClassClass.png", "Mô hình ClassEntity",),
        (UML_DIR / "10-class-diagram" / "ClassExam.png", "Mô hình Exam",),
        (UML_DIR / "10-class-diagram" / "ClassQuestion.png", "Mô hình Question",),
        (UML_DIR / "10-class-diagram" / "ClassSubmission.png", "Mô hình Submission",),
        (UML_DIR / "08-admin" / "AdminUseCaseCompact.png", "Use case quản trị",),
        (UML_DIR / "08-admin" / "AdminSequenceCompact.png", "Trình tự quản trị",),
        (UML_DIR / "11-deployment" / "DeploymentCompact.png", "Triển khai service",),
        (UML_DIR / "11-deployment" / "ComponentArchitectureOverview.png", "Kiến trúc component",),
        (UML_DIR / "13-database" / "ERDatabaseOverview.png", "Tổng quan ERD",),
        (UML_DIR / "13-database" / "EROverview.png", "Quan hệ dữ liệu tổng thể",),
        (UML_DIR / "13-database" / "ERExam.png", "ER nhánh đề thi",),
        (UML_DIR / "13-database" / "ERQuestion.png", "ER nhánh câu hỏi",),
        (UML_DIR / "13-database" / "ERSubmission.png", "ER nhánh bài nộp",),
    ], start=1):
        add_figure(document, 2, 24 + idx, img, title, restart=False, width_in=6.05)

    add_heading(document, "2.11. Thiết kế API và bản đồ endpoint", 2)
    add_table(document, [
        "Nhóm API",
        "Endpoint tiêu biểu",
        "Chức năng",
    ], [
        ["/api/auth", "/login, /refresh, /logout, /register", "Xác thực và tài khoản"],
        ["/api/exams", "/exams, /exams/{id}/sessions", "Đề thi và phiên thi"],
        ["/api/teacher/classes", "/teacher/classes, /students/import", "Lớp học giảng viên"],
        ["/api/student/classes", "/student/classes/join", "Lớp học sinh viên"],
        ["/api/v1/import", "/upload, /preview/{jobId}, /confirm/{jobId}", "Import dữ liệu chung"],
        ["/api/v1/exam-import", "/upload, /preview/{sessionId}", "Import đề thi từ file"],
        ["/api/v1/proctor", "/start, /signals, /sessions/{attemptId}/events/batch", "Giám sát realtime"],
        ["/api/attempts/{attemptId}/monitoring", "/risk, /timeline, /warning", "Theo dõi attempt"],
        ["/api/v1/fraud", "/analyze/exam/{examId}, /ml-risk/exam/{examId}", "Phân tích gian lận"],
        ["/api/admin", "/dashboard/stats, /users/*, /classes, /exams", "Quản trị hệ thống"],
        ["/api/v1/ai", "/generate/from-topic, /evaluate/essay, /chat", "Dịch vụ AI"],
    ], caption="Bản đồ nhóm endpoint", chapter_no=2, caption_no=30)

    add_table(document, [
        "Ngưỡng điểm",
        "Mức rủi ro",
        "Ý nghĩa vận hành",
    ], [
        ["0-29", "SAFE", "Không có cảnh báo nổi bật"],
        ["30-49", "WATCH", "Theo dõi tăng cường"],
        ["50-69", "SUSPICIOUS", "Có dấu hiệu gian lận rõ"],
        ["70-100", "CRITICAL", "Can thiệp trực tiếp hoặc vô hiệu hóa"],
    ], caption="Ngưỡng đánh giá rủi ro", chapter_no=2, caption_no=31)

    add_table(document, [
        "Nhóm dữ liệu",
        "Bảng chính",
        "Vai trò",
    ], [
        ["Người dùng", "users, roles, refresh_tokens", "Tài khoản và phân quyền"],
        ["Thi", "exams, exam_attempts, exam_events", "Đề và phiên thi"],
        ["Bài làm", "answers, submissions", "Lưu bài nộp"],
        ["Giám sát", "monitoring_events, fraud_signals, risk_scores, audit_logs", "Lưu vết và cảnh báo"],
        ["Lớp học", "classes, class_students", "Quản lý lớp và thành viên"],
        ["Import", "import_jobs, import_job_issues", "Theo dõi quá trình import"],
    ], caption="Nhóm bảng dữ liệu trọng tâm", chapter_no=2, caption_no=32)


def build_chapter_3(document: Document, exam_state: dict[str, str] | None = None) -> None:
    add_heading(document, "CHƯƠNG III: CÀI ĐẶT VÀ KIỂM THỬ", 1)
    add_heading(document, "3.1. Môi trường kiểm thử cục bộ", 2)
    add_text_block(document, [
        "Kiểm thử được thực hiện trên local stack do `start-demo-local.ps1` khởi động. Script này build FE/BE, chạy AI service, mở preview frontend và chuẩn bị dữ liệu đề thi để có thể sinh flow thật trên dashboard monitoring.",
        "Trong lần chạy gần nhất, hệ thống dùng exam ID 60 với mã thi được sinh từ backend; trạng thái được lưu trong `demo-local-state.json` để các màn monitoring và kết quả có dữ liệu đồng nhất.",
    ])
    if exam_state:
        add_table(document, [
            "Thông tin",
            "Giá trị",
        ], [
            ["Exam ID", str(exam_state.get("examId", ""))],
            ["Exam code", str(exam_state.get("examCode", ""))],
            ["Title", str(exam_state.get("title", ""))],
            ["Teacher", str(exam_state.get("teacherUsername", ""))],
            ["Student", str(exam_state.get("studentUsername", ""))],
        ], caption="Trạng thái dữ liệu local", chapter_no=3, caption_no=1, restart=True)

    add_table(document, [
        "Lệnh / script",
        "Mục đích",
    ], [
        ["start-demo-local.ps1", "Khởi động FE, BE, AI service và chuẩn bị dữ liệu demo"],
        ["smoke-demo-flow.ps1", "Chạy một flow thi và tạo dữ liệu giám sát thực"],
        ["npm run preview -- --host 0.0.0.0 --port 4173", "Chạy preview frontend"],
        ["mvn -q -DskipTests package", "Đóng gói backend"],
    ], caption="Lệnh kiểm thử chính", chapter_no=3, caption_no=2)

    add_heading(document, "3.2. Giao diện và luồng người dùng", 2)
    add_text_block(document, [
        "Các ảnh chụp màn hình bên dưới được lấy trực tiếp từ route chính của hệ thống: `/login`, `/student/*`, `/teacher/*` và `/admin/*`. Chúng cho thấy toàn bộ luồng từ xác thực đến tạo đề, thi thử, giám sát realtime và quản trị.",
    ])

    screenshot_items = [
        (SCREENSHOT_DIR / "01-login.png", "Màn hình đăng nhập"),
        (SCREENSHOT_DIR / "02-student-dashboard.png", "Dashboard sinh viên"),
        (SCREENSHOT_DIR / "03-student-join-exam.png", "Sinh viên nhập mã vào phòng thi"),
        (SCREENSHOT_DIR / "04-student-waiting-room-fake.png", "Phòng chờ với camera/mic đã xác nhận"),
        (SCREENSHOT_DIR / "05-student-exam-interface-fake.png", "Giao diện làm bài của sinh viên"),
        (SCREENSHOT_DIR / "06-teacher-dashboard.png", "Dashboard giảng viên"),
        (SCREENSHOT_DIR / "07-teacher-create-exam.png", "Màn hình tạo đề thi"),
        (SCREENSHOT_DIR / "08-teacher-monitoring.png", "Màn hình giám sát trực tiếp"),
        (SCREENSHOT_DIR / "09-teacher-student-monitoring-detail.png", "Chi tiết sinh viên đang thi"),
        (SCREENSHOT_DIR / "10-admin-dashboard.png", "Dashboard quản trị"),
        (SCREENSHOT_DIR / "11-admin-users.png", "Quản lý người dùng trong admin"),
    ]
    for idx, (img, title) in enumerate(screenshot_items, start=1):
        add_figure(document, 3, idx, img, title, restart=(idx == 1), width_in=5.95)

    add_heading(document, "3.3. Kiểm thử chức năng", 2)
    add_table(document, [
        "Mã test",
        "Kịch bản",
        "Kết quả kỳ vọng",
        "Kết quả quan sát",
    ], [
        ["TC-01", "Đăng nhập bằng tài khoản teacher1", "Nhận token và vào dashboard giảng viên", "Đạt"],
        ["TC-02", "Sinh viên join exam bằng mã thi", "Vào phòng chờ và kiểm tra thiết bị", "Đạt"],
        ["TC-03", "Làm bài và lưu nháp", "Draft được lưu và trạng thái bài thi cập nhật", "Đạt"],
        ["TC-04", "Giảng viên mở monitoring", "Xem list attempt, risk score và timeline", "Đạt"],
        ["TC-05", "Gửi signal vi phạm", "Điểm rủi ro tăng và cảnh báo xuất hiện", "Đạt"],
        ["TC-06", "Admin xem dashboard", "Thấy thống kê user/exam/class", "Đạt"],
    ], caption="Bảng kiểm thử chức năng", chapter_no=3, caption_no=3)

    add_heading(document, "3.4. Kết quả kiểm thử smoke", 2)
    add_text_block(document, [
        "Khi chạy `smoke-demo-flow.ps1`, hệ thống sinh dữ liệu monitoring thật cho một attempt và đẩy các event bất thường như tab switch, blur, paste lớn và mở devtools. Backend gom tín hiệu, tính risk score và đẩy lên dashboard giảng viên.",
        "Kết quả quan sát được trên lần chạy gần nhất cho thấy attempt được ghi nhận với điểm rủi ro ở mức cảnh báo, số lượng alert và timeline đủ để minh họa luồng giám sát thực tế trong báo cáo.",
    ])
    add_table(document, [
        "Chỉ số",
        "Giá trị ghi nhận",
    ], [
        ["Attempt ID", "164"],
        ["Risk score", "53"],
        ["Risk level", "SUSPICIOUS"],
        ["Alert count", "5"],
        ["Timeline count", "10"],
    ], caption="Kết quả smoke flow", chapter_no=3, caption_no=4)


def build_conclusion(document: Document) -> None:
    add_heading(document, "KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", 1)
    add_text_block(document, [
        "Khóa luận đã hoàn thành một hệ thống thi trực tuyến có cấu trúc rõ ràng, có giao diện tách vai trò, có backend API đủ lớn cho nghiệp vụ thật, có dịch vụ AI hỗ trợ camera/OCR/nội dung và có cơ chế giám sát realtime để phát hiện gian lận.",
        "Điểm nổi bật của FE_DEMO là kết hợp đồng thời xác thực, import, lớp học, thi, monitoring, risk scoring và dashboard vận hành trong cùng một hệ thống chạy được trên local stack và có dữ liệu kiểm thử thật.",
        "Hướng phát triển tiếp theo gồm hoàn thiện mô hình học máy cho chấm rủi ro, tăng độ ổn định của phát hiện camera/mic, tối ưu luồng đồng bộ WebSocket và bổ sung thêm báo cáo phân tích sau thi cho giảng viên và quản trị viên.",
    ])


def build_references(document: Document) -> None:
    add_heading(document, "TÀI LIỆU THAM KHẢO", 1)
    refs = [
        "1. `demo/package.json` - cấu hình frontend Vue/Vite và các dependency chính.",
        "2. `BE/pom.xml` - cấu hình Spring Boot 4, Java 17 và thư viện backend.",
        "3. `ai-service/requirements.txt` - thư viện Python cho AI service.",
        "4. `BE/src/main/resources/schema.sql` - lược đồ dữ liệu và bảng mở rộng.",
        "5. `demo/src/router.js` - bản đồ route frontend theo vai trò.",
        "6. `docs/exam-proctoring-system-analysis.md` và các tài liệu phân tích trong `docs/`.",
        "7. `docs/diagrams/` và `docs/uml-diagrams/plantuml_word_compact_png/` - sơ đồ kiến trúc, use case, sequence, ERD.",
        "8. Mã nguồn các controller, service và realtime gateway trong `BE/src/main/java/com/example/demo/`.",
        "9. Mã nguồn dịch vụ AI trong `ai-service/app/`.",
    ]
    add_text_block(document, refs, first_line_cm=0.0)


def collect_project_info() -> dict[str, str]:
    package_json = read_json(ROOT / "demo" / "package.json")
    pom = read_pom_versions(ROOT / "BE" / "pom.xml")
    requirements = read_requirements(ROOT / "ai-service" / "requirements.txt")

    def find_req(prefix: str) -> str:
        for line in requirements:
            if line.startswith(prefix):
                return line
        return "không khóa version"

    return {
        "fe_package_version": package_json.get("version", "unknown"),
        "vue_version": package_json["dependencies"].get("vue", "unknown").strip("^"),
        "vite_version": package_json["devDependencies"].get("vite", "unknown").strip("^"),
        "pinia_version": package_json["dependencies"].get("pinia", "unknown").strip("^"),
        "vue_router_version": package_json["dependencies"].get("vue-router", "unknown").strip("^"),
        "primevue_version": package_json["dependencies"].get("primevue", "unknown").strip("^"),
        "echarts_version": package_json["dependencies"].get("echarts", "unknown").strip("^"),
        "xlsx_version": package_json["dependencies"].get("xlsx", "unknown").strip("^"),
        "spring_boot": pom["spring_boot"],
        "java": pom["java"],
        "project_version": pom["project_version"],
        "fastapi": find_req("fastapi"),
        "uvicorn": find_req("uvicorn"),
        "opencv": find_req("opencv-python-headless"),
        "numpy": find_req("numpy"),
        "pymupdf": find_req("PyMuPDF"),
        "openai": find_req("openai"),
    }


def main() -> None:
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(TEMPLATE_PATH)

    source = Document(TEMPLATE_PATH)
    template_section = source.sections[0]

    document = Document(TEMPLATE_PATH)
    clear_body(document)
    copy_section_metrics(document.sections[0], template_section)
    set_core_props(document)
    set_style_defaults(document)

    project_info = collect_project_info()
    versions = read_pom_versions(ROOT / "BE" / "pom.xml")
    requirements = read_requirements(ROOT / "ai-service" / "requirements.txt")
    exam_state = read_json(ROOT / "demo-local-state.json") if (ROOT / "demo-local-state.json").exists() else None

    build_intro(document)
    build_opening(document)
    build_chapter_1(document, versions, requirements, project_info)
    build_chapter_2(document, project_info)
    build_chapter_3(document, exam_state)
    build_conclusion(document)
    build_references(document)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_PATH)
    print(f"wrote {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
