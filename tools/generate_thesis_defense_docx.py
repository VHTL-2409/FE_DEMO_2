from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "baocao" / "Phan_tich_core_thuat_toan_va_cau_hoi_bao_ve_KLTN.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(11)


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(18)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(level=level)
    p.style.font.name = "Times New Roman"
    p.style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = True


def add_p(doc: Document, text: str = "", bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r2 = p.add_run(text[len(bold_prefix):])
        runs = [r1, r2]
    else:
        runs = [p.add_run(text)]
    for run in runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(12)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(item)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(12)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(item)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(12)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True)
        set_cell_shading(hdr[i], "D9EAF7")
        if widths:
            hdr[i].width = Inches(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            if widths:
                cells[i].width = Inches(widths[i])
    doc.add_paragraph()


def add_code_ref_table(doc: Document) -> None:
    rows = [
        ["Frontend", "demo/src/router.js", "Khai báo route, guard phân quyền, phục hồi query phiên thi.", "Bảo vệ được vì sao tách route admin/teacher/student và kiểm tra role trước khi vào trang."],
        ["Frontend", "demo/src/services/apiClient.js", "Gửi API, tự gắn JWT, refresh token trước khi hết hạn, xử lý lỗi chuẩn.", "Giải thích cơ chế xác thực phía client và lý do không lặp code fetch ở từng component."],
        ["Frontend", "demo/src/composables/useProctoringSession.js", "Batch event giám sát và heartbeat 15 giây.", "Nêu được cách giảm spam request nhưng vẫn giữ dữ liệu realtime đủ gần."],
        ["Frontend", "demo/src/composables/useExamMonitoring.js", "Kết nối STOMP theo exam/attempt, fallback polling khi mất WebSocket.", "Trả lời được câu hỏi về realtime và khả năng chịu lỗi mạng."],
        ["Backend", "BE/src/main/java/com/example/demo/service/SubmissionService.java", "Chuẩn bị, bắt đầu, lưu nháp, nộp bài, auto-submit, kiểm tra IP.", "Đây là lõi nghiệp vụ làm bài thi trực tuyến."],
        ["Backend", "BE/src/main/java/com/example/demo/service/RiskScoringService.java", "Tính điểm rủi ro theo signal, category, dedup window, cap và level.", "Đây là thuật toán chính cho phát hiện gian lận dựa trên sự kiện."],
        ["Backend", "BE/src/main/java/com/example/demo/service/MonitoringService.java", "Ghi event, timeline, pause/resume/invalidate, audit và trạng thái camera.", "Giải thích được vai trò giám thị và kiểm soát phiên thi."],
        ["Backend", "BE/src/main/java/com/example/demo/service/AiAssistService.java", "Bridge backend với AI service, chuẩn hóa signal, lưu evidence, phát realtime.", "Nêu được cách AI được tích hợp vào hệ thống nghiệp vụ."],
        ["AI service", "ai-service/app/proctor.py", "Phân tích frame: face detection, eye/gaze, brightness, blur, spoofing/deepfake.", "Trả lời được các câu hỏi về thuật toán xử lý ảnh và giới hạn của AI."],
        ["Parser", "python_parser/app/parser_router.py", "Chọn parser theo profile/score, fallback nếu parser chính lỗi.", "Nêu được cách import đề tự động hoạt động với nhiều template."],
    ]
    add_table(
        doc,
        ["Tầng", "File chính", "Vai trò", "Ý nghĩa khi bảo vệ"],
        rows,
        [1.0, 2.3, 2.5, 2.5],
    )


def build_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    styles["Normal"].font.size = Pt(12)

    add_title(doc, "PHÂN TÍCH CORE, THUẬT TOÁN, CODE VÀ BỘ CÂU HỎI BẢO VỆ KHÓA LUẬN")
    add_title(doc, "Dự án: Hệ thống thi trực tuyến có giám sát và phát hiện gian lận")
    add_p(doc, "Tài liệu này tổng hợp các thành phần lõi của dự án, mô tả thuật toán theo đúng cấu trúc code hiện tại và chuẩn bị bộ câu hỏi bảo vệ khóa luận tốt nghiệp. Nội dung được viết theo hướng có thể dùng trực tiếp khi thuyết trình, trả lời hội đồng và rà soát lại code trước buổi bảo vệ.")

    add_heading(doc, "1. Tổng quan dự án", 1)
    add_p(doc, "Dự án là hệ thống thi trực tuyến gồm nhiều vai trò: quản trị viên, giáo viên và học sinh. Hệ thống hỗ trợ tạo đề, nhập đề từ file, tổ chức phiên thi, học sinh làm bài, tự động chấm điểm, giám sát gian lận thời gian thực, phân tích AI camera và lưu vết để giáo viên review.")
    add_bullets(doc, [
        "Frontend: Vue 3, Vite, Pinia, Vue Router, Tailwind/PrimeVue, STOMP/SockJS cho realtime.",
        "Backend: Java 17, Spring Boot, Spring Security, JPA/Hibernate, WebSocket, PostgreSQL, JWT/refresh token, OAuth2 Google.",
        "AI service: Python FastAPI, OpenCV, PIL, NumPy, OCR, phân tích frame camera, deepfake/spoofing heuristic.",
        "Parser service: Python FastAPI, PyMuPDF, python-docx, parser router theo template để nhập đề PDF/DOCX.",
        "Triển khai: Dockerfile từng service, docker-compose cho môi trường dev/prod, cấu hình Nginx mẫu.",
    ])

    add_heading(doc, "2. Kiến trúc tổng thể", 1)
    add_p(doc, "Kiến trúc được chia theo mô hình nhiều service. Frontend chỉ phụ trách trải nghiệm người dùng, lưu trạng thái giao diện và gửi sự kiện. Backend là lõi nghiệp vụ, kiểm soát quyền truy cập, trạng thái bài thi, điểm số, rủi ro và realtime. AI service và parser service được tách riêng vì đây là các tác vụ nặng, có thư viện Python đặc thù và có thể scale độc lập.")
    add_table(doc, ["Thành phần", "Trách nhiệm", "Công nghệ", "Dữ liệu ra/vào"], [
        ["Frontend Vue", "Giao diện admin/teacher/student, làm bài, giám sát, import đề, dashboard.", "Vue 3, Pinia, Router, STOMP", "HTTP JSON, multipart upload, WebSocket payload."],
        ["Backend Spring Boot", "Nghiệp vụ chính, bảo mật, phân quyền, database, realtime, risk scoring.", "Spring Boot, JPA, Security, WebSocket", "Entity, DTO, JWT, event, audit log."],
        ["AI service", "OCR, frame analysis, identity verification, behavior analysis.", "FastAPI, OpenCV, PIL, NumPy", "Base64 image, metadata, signals/warnings."],
        ["Python parser", "Parse PDF/DOCX đề thi theo template, tách câu hỏi/đáp án/metadata.", "FastAPI, PyMuPDF, python-docx", "File đề thi, ParseResponse, report, questions."],
        ["Database", "Lưu user, role, exam, question, attempt, answer, signal, warning, import session.", "PostgreSQL", "Bảng nghiệp vụ và JSONB cho payload linh hoạt."],
    ], [1.3, 2.6, 1.7, 2.2])

    add_p(doc, "Luồng tổng quát: giáo viên tạo/import đề thi -> phân công cho lớp -> học sinh vào phòng chờ -> bắt đầu bài thi -> frontend gửi answer draft, monitoring event, heartbeat, camera frame -> backend lưu attempt/answer/signal -> risk scoring cập nhật mức rủi ro -> WebSocket đẩy cảnh báo tới giáo viên -> giáo viên review/pause/resume/invalidate -> học sinh nộp bài hoặc hệ thống auto-submit khi hết giờ.")

    add_heading(doc, "3. Các core code chính", 1)
    add_code_ref_table(doc)

    add_heading(doc, "4. Core nghiệp vụ: xác thực và phân quyền", 1)
    add_p(doc, "Hệ thống dùng JWT access token kết hợp refresh token. Frontend lưu token ở localStorage, `apiClient.js` kiểm tra token sắp hết hạn trước mỗi request và gọi `/api/auth/refresh` để làm mới. Nếu refresh thất bại, client gọi `invalidateSession()` và điều hướng về đăng nhập.")
    add_bullets(doc, [
        "Backend có `AuthController` cho register, login, refresh, logout, forgot/reset password, verify email và Google OAuth2.",
        "Router frontend dùng `beforeEach` để chặn route theo `requiresAuth`, `adminOnly`, `teacherOnly`, `studentOnly`.",
        "Backend không tin frontend: service vẫn kiểm tra quyền bằng role và quan hệ sở hữu dữ liệu, ví dụ giáo viên chỉ quản lý đề mình tạo, học sinh chỉ vào bài thi thuộc lớp được phép.",
        "Cách thiết kế này đáp ứng nguyên tắc defense in depth: UI chặn sớm để UX tốt, backend chặn cuối để bảo mật thật.",
    ])

    add_heading(doc, "5. Core nghiệp vụ: tạo đề, lớp học và import câu hỏi", 1)
    add_p(doc, "`ExamService` là service trung tâm cho đề thi. Khi tạo đề, service validate thời gian bắt đầu/kết thúc, sinh mã đề 8 ký tự bằng `SecureRandom`, gán lớp hợp lệ, lưu các cấu hình giám sát như tab switch, fullscreen, copy paste, duplicate IP, AI proctoring, shuffle câu hỏi/đáp án.")
    add_p(doc, "Import câu hỏi có hai hướng. Với CSV/XLSX, backend dùng `QuizParserEngine`, `CsvQuestionParser`, `XlsxQuestionParser`. Với PDF/DOCX template phức tạp, backend dùng `ExamImportService` gọi Python parser service theo cơ chế async session.")
    add_bullets(doc, [
        "CSV/XLSX: đọc header, normalize tên cột tiếng Việt/tiếng Anh, parse content/options/correctAnswer/scoreWeight/difficulty.",
        "PDF/DOCX: upload file -> tạo `ExamImportSession` -> chạy parse async -> lưu report/parseResult JSON -> giáo viên preview -> confirm để tạo `Question` entities.",
        "Parser service chọn template theo điểm nhận dạng, sau đó validate kết quả và fallback parser nếu parser chính lỗi.",
        "Import session chỉ lưu kết quả parse và metadata; file gốc được đọc thành bytes để xử lý, giảm rủi ro lưu file thừa.",
    ])

    add_heading(doc, "6. Core nghiệp vụ: luồng học sinh làm bài", 1)
    add_p(doc, "`SubmissionService` điều khiển vòng đời bài làm. Hai bước quan trọng là `prepareAttempt` và `startAttempt`. `prepareAttempt` cho phép học sinh vào phòng chờ, tạo attempt trạng thái `WAITING`. `startAttempt` chuyển attempt sang `IN_PROGRESS`, ghi thời gian bắt đầu, heartbeat đầu tiên, fingerprint/session token và phát realtime cho giáo viên.")
    add_numbered(doc, [
        "Học sinh nhập mã đề hoặc chọn bài trong dashboard.",
        "Backend gọi `validateExamJoinable` hoặc `validateExamAvailability` để kiểm tra đề active, chưa kết thúc, đã tới giờ thi.",
        "Nếu chưa bắt đầu, attempt ở `WAITING`; nếu bấm bắt đầu, attempt chuyển `IN_PROGRESS`.",
        "Trong khi làm bài, frontend lưu nháp qua `draft-answers`; backend tăng save count, lưu answer hiện tại.",
        "Khi nộp, backend kiểm tra status, deadline, lưu answer cuối cùng, tính điểm, set `SUBMITTED`, phát realtime.",
        "Nếu quá hạn, `autoSubmitFromDraft` tính điểm từ nháp và set `AUTO_SUBMITTED`.",
    ])
    add_p(doc, "Một điểm bảo vệ quan trọng là kiểm tra IP nhất quán. `enforceIpConsistency` so sánh IP ban đầu với IP hiện tại; nếu thay đổi trong khi làm bài, attempt bị set `STOPPED`, ghi audit log và thông báo realtime. Đây là biện pháp chống chia sẻ phiên hoặc chuyển thiết bị/mạng bất thường, tuy chưa phải bằng chứng tuyệt đối về gian lận.")

    add_heading(doc, "7. Thuật toán chấm điểm", 1)
    add_p(doc, "Chấm điểm trực tiếp trong `SubmissionHelper.calculateScore` và phân tích nâng cao trong `GradingService`. Bài trắc nghiệm được chấm bằng cách so sánh `selectedAnswer` với `Question.correctAnswer`, cộng `scoreWeight` khi đúng. Kết quả có thể quy đổi theo tổng điểm hoặc giữ raw score tùy cách cấu hình đề.")
    add_p(doc, "`GradingService` bổ sung thống kê lớp và phân tích tương đối: điểm trung bình lớp, độ lệch chuẩn, percentile, rank, theta IRT xấp xỉ và reliability. Phần IRT đang dùng mô hình 1PL/Rasch đơn giản theo logit của tỉ lệ điểm.")
    add_table(doc, ["Công thức", "Ý nghĩa", "Cách dùng trong hệ thống"], [
        ["p = score / maxScore", "Tỉ lệ điểm của học sinh.", "Chuẩn hóa điểm trước khi tính theta."],
        ["theta = log(p/(1-p)) - log(p_class/(1-p_class))", "Năng lực tương đối so với trung bình lớp.", "Theta > 0 nghĩa là cao hơn lớp; theta < 0 nghĩa là thấp hơn lớp."],
        ["variance = sum((score - mean)^2) / n", "Độ phân tán điểm trong lớp.", "Dùng để tính reliability và đánh giá độ phân hóa đề."],
        ["reliability = variance / (variance + errorVariance)", "Độ tin cậy xấp xỉ của kết quả.", "Nếu lớp quá ít bài nộp, hệ thống dùng giá trị mặc định."],
    ], [2.0, 2.2, 3.1])
    add_p(doc, "Khi bảo vệ, cần nói rõ đây là IRT/phân tích thống kê ở mức ứng dụng, không phải mô hình đo lường giáo dục đầy đủ. Điểm mạnh là minh bạch, dễ giải thích và đủ dùng cho dashboard; hạn chế là chưa ước lượng tham số câu hỏi a/b/c bằng dữ liệu lớn.")

    add_heading(doc, "8. Core giám sát realtime", 1)
    add_p(doc, "Giám sát được triển khai theo mô hình event-driven. Frontend phát hiện hành vi như rời tab, blur, thoát fullscreen, copy/paste, heartbeat trễ, trạng thái camera/mic và gửi event batch về backend. Backend ghi event, chuyển thành fraud signal, tính lại risk score và đẩy thông báo qua WebSocket/STOMP.")
    add_bullets(doc, [
        "`useProctoringSession.js`: gom event trong cửa sổ mặc định 1500ms để giảm số request.",
        "Heartbeat mặc định 15000ms, kèm trạng thái mạng, fingerprint và telemetry.",
        "`useExamMonitoring.js`: giáo viên subscribe `/topic/exams/{examId}/alerts` và từng `/topic/attempts/{attemptId}/proctor-actions`.",
        "Nếu WebSocket mất kết nối, frontend chuyển sang polling mỗi 8 giây để đồng bộ danh sách attempt.",
        "`MonitoringService`: cho phép giáo viên warning, pause, resume, invalidate; đồng thời ghi audit/timeline.",
    ])

    add_heading(doc, "9. Thuật toán risk scoring", 1)
    add_p(doc, "`RiskScoringService` là thuật toán trọng tâm cho phát hiện gian lận. Hệ thống không kết luận trực tiếp “gian lận”, mà tính điểm rủi ro dựa trên tập signal đã ghi nhận. Cách tiếp cận này phù hợp bảo vệ khóa luận vì minh bạch, có threshold rõ, dễ kiểm thử và giáo viên vẫn là người review cuối.")
    add_table(doc, ["Nhóm signal", "Ví dụ signal", "Dedup window", "Cap điểm", "Ý nghĩa"], [
        ["SCREEN_LEAVE", "TAB_SWITCH, WINDOW_BLUR, EXIT_FULLSCREEN, LONG_SCREEN_LEAVE", "60 giây", "35", "Hành vi rời màn hình hoặc né fullscreen."],
        ["CLIPBOARD", "COPY_PASTE", "30 giây", "25", "Sao chép/dán bất thường."],
        ["TECHNICAL", "RIGHT_CLICK, PRINT_SCREEN, DEVTOOLS_OPEN", "120 giây", "25", "Hành vi kỹ thuật đáng nghi."],
        ["IDENTITY", "IP_CHANGED, DUPLICATE_IP, DEVICE_FINGERPRINT_CHANGED", "Theo session", "30", "Bất thường định danh, thiết bị hoặc mạng."],
        ["HEARTBEAT", "HEARTBEAT_STALE, NETWORK_INSTABILITY, SESSION_RECOVERY", "Không window cụ thể", "10", "Mất kết nối hoặc phiên không ổn định."],
        ["VISUAL_IDENTITY", "FACE_NOT_DETECTED, MULTIPLE_FACES, GAZE_OFF_SCREEN, DEEPFAKE", "30 giây", "40", "Bất thường từ AI camera."],
    ], [1.3, 2.1, 1.2, 0.9, 2.0])
    add_numbered(doc, [
        "Lấy toàn bộ `FraudSignal` của attempt theo thời gian tăng dần.",
        "Chuẩn hóa tên signal bằng `FraudSignalTypeNormalizer`.",
        "Map signal vào category bằng bảng `SIGNAL_CATEGORY`.",
        "Với mỗi category, sort theo `createdAt` và chia cluster theo cửa sổ dedup.",
        "Mỗi cluster chỉ lấy điểm cao nhất để tránh spam cùng một lỗi trong thời gian ngắn.",
        "Cộng điểm các cluster rồi giới hạn bởi cap của category.",
        "Tính `behaviorScore = min(70, screenLeave + clipboard + technical + heartbeat)`.",
        "Tính `totalRisk = min(100, behaviorScore + identity + visualIdentity)`.",
        "Map điểm sang level: CLEAN, SUSPICIOUS, HIGH_RISK, CRITICAL.",
        "Nếu điểm đủ cao, đồng bộ `ProctorFlag`, lưu snapshot và gửi realtime cho giáo viên.",
    ])
    add_table(doc, ["Khoảng điểm", "RiskLevel", "Ý nghĩa vận hành", "Hành động gợi ý"], [
        ["0-19", "CLEAN", "Chưa có dấu hiệu đáng kể.", "Tiếp tục giám sát."],
        ["20-49", "SUSPICIOUS", "Có dấu hiệu đáng ngờ nhưng chưa nghiêm trọng.", "Review attempt khi cần."],
        ["50-74", "HIGH_RISK", "Nguy cơ cao, cần giáo viên xem xét.", "Tạo flag và đưa vào danh sách review."],
        ["75-100", "CRITICAL", "Nguy cơ nghiêm trọng.", "Flag ưu tiên cao, giáo viên can thiệp."],
    ], [1.2, 1.2, 2.6, 2.2])
    add_p(doc, "Ví dụ giải thích nhanh: nếu học sinh chuyển tab 3 lần trong 30 giây, hệ thống không cộng 30 điểm ngay mà gom vào một cluster 60 giây và lấy max 10 điểm. Nếu sau hơn 60 giây lại chuyển tab, cluster mới được tính thêm. Điều này giúp giảm false positive do thao tác lặp hoặc trình duyệt nháy focus.")

    add_heading(doc, "10. AI camera và xác minh danh tính", 1)
    add_p(doc, "AI camera gồm hai lớp: frontend/backend bridge và Python AI service. Frontend gửi frame base64 kèm metadata. Backend `AiAssistService.analyzeFrame` kiểm tra payload, publish frame cho giáo viên, gọi AI service nếu bật, chuẩn hóa signal, lưu evidence image khi có cảnh báo, ghi `FraudSignal` và `FraudWarning`, rồi đẩy realtime.")
    add_p(doc, "Trong `ai-service/app/proctor.py`, `ProctorAnalyzer.analyze_frame` thực hiện pipeline xử lý ảnh: decode base64 -> phát hiện khuôn mặt -> theo dõi mắt/gaze -> đo brightness/variance -> kiểm tra spoofing/deepfake heuristic -> kiểm tra occlusion -> sinh signals/warnings.")
    add_numbered(doc, [
        "Decode ảnh base64 bằng PIL và chuyển sang RGB/gray.",
        "Face detection ưu tiên OpenCV DNN Caffe; nếu không sẵn sàng hoặc fail thì fallback Haar Cascade.",
        "Merge các bounding box trùng bằng IoU/center-close để tránh đếm trùng khuôn mặt.",
        "Eye/gaze tracking dùng MediaPipe FaceMesh nếu có, fallback OpenCV face module hoặc heuristic.",
        "Frame quality dựa vào độ sáng trung bình và phương sai ảnh xám: quá tối, quá sáng, mờ.",
        "Spoofing/deepfake kết hợp model nếu có và heuristic như printed photo, screen replay, flat image.",
        "Sinh signal chuẩn hóa: `signal_type`, `severity`, `confidence`, `evidence`.",
    ])
    add_table(doc, ["Signal AI", "Nguyên nhân có thể", "Mức độ thường gặp", "Cách giải thích khi bảo vệ"], [
        ["FACE_NOT_DETECTED", "Không thấy mặt, camera lệch, ánh sáng yếu.", "HIGH", "Không kết luận gian lận, chỉ là tín hiệu cần kiểm tra."],
        ["MULTIPLE_FACES", "Có nhiều người trong khung hình.", "CRITICAL", "Rủi ro cao vì vi phạm quy định phòng thi."],
        ["GAZE_OFF_SCREEN", "Mắt nhìn ra ngoài màn hình đủ lâu.", "HIGH/MEDIUM", "Có ngưỡng thời gian và số frame để giảm nhiễu."],
        ["VERY_LOW_LIGHTING", "Ánh sáng thấp, ảnh khó xác thực.", "HIGH", "Yêu cầu môi trường đủ sáng để giám sát."],
        ["DEEPFAKE/SCREEN_REPLAY", "Nghi ngờ giả mạo camera.", "CRITICAL", "Là cảnh báo review, cần giáo viên xem evidence."],
    ], [1.5, 2.1, 1.2, 2.6])
    add_p(doc, "Xác minh danh tính dùng `verifyIdentity`: yêu cầu ảnh giấy tờ và selfie, kiểm tra attempt/student, gọi AI `/identity/verify`, normalize response, lưu `StudentIdentityCheck`, lưu evidence và bridge signal về risk scoring. Khi AI không khả dụng, hệ thống fallback `NEEDS_REVIEW`, không cho ra kết quả xác minh giả.")

    add_heading(doc, "11. Parser nhập đề PDF/DOCX", 1)
    add_p(doc, "`python_parser/app/parser_router.py` chứa `ParserRouter`, có registry nhiều parser template. Với PDF, service xây dựng `PdfProfile`, chấm điểm từng parser qua `can_handle`, chọn parser điểm cao nhất, parse, validate và fallback nếu lỗi. Với DOCX, router đọc sample bằng `DocxReader`, chấm riêng template 04/05 dựa trên pattern trong nội dung.")
    add_bullets(doc, [
        "Template 01: đề toán có công thức hoặc layout phức tạp được rebuild.",
        "Template 02: trắc nghiệm sạch, cấu trúc rõ.",
        "Template 03: toán có bảng đáp án hoặc answer grid.",
        "Template 04/05: DOCX tiếng Việt hoặc dạng database câu hỏi.",
        "Template 06: đề tiếng Anh.",
    ])
    add_p(doc, "Điểm đáng bảo vệ là hệ thống không hard-code một parser duy nhất. Router dùng scoring và fallback nên linh hoạt với nhiều mẫu đề. Kết quả parse có `ParseReport` gồm warning, error, confidence, parseTimeMs để giáo viên biết chất lượng import.")

    add_heading(doc, "12. Realtime, audit và khả năng chịu lỗi", 1)
    add_p(doc, "Realtime không chỉ để hiển thị đẹp mà là phần nghiệp vụ: giáo viên cần thấy học sinh vào phòng chờ, bắt đầu, nộp bài, risk update, camera signal, warning, pause/resume/stop. Backend dùng `RealtimeNotificationService` và gateway để publish payload. Frontend có fallback polling để tránh phụ thuộc tuyệt đối vào WebSocket.")
    add_bullets(doc, [
        "Audit log ghi lại hành động quan trọng: thay đổi IP, giáo viên pause/resume/invalidate, warning.",
        "Timeline attempt ghép nhiều nguồn: exam event, monitoring event, fraud signal, fraud warning, risk snapshot.",
        "Các read endpoint dùng `recomputeRiskSkipAutoActions` để tránh side effect khi chỉ xem dữ liệu.",
        "Transaction synchronization afterCommit giúp chỉ phát realtime sau khi database commit thành công.",
    ])

    add_heading(doc, "13. Cơ sở dữ liệu và entity quan trọng", 1)
    add_table(doc, ["Entity", "Vai trò", "Quan hệ chính"], [
        ["User, Role, StudentProfile, TeacherProfile", "Tài khoản, vai trò và hồ sơ.", "User có nhiều Role; profile theo vai trò."],
        ["ClassEntity, ClassStudent", "Lớp học và danh sách học sinh.", "Giáo viên sở hữu lớp; học sinh tham gia lớp."],
        ["Exam, Question, Assignment", "Đề thi, câu hỏi, phân công.", "Exam có nhiều Question, có thể gán lớp hoặc publish."],
        ["ExamAttempt, Answer", "Phiên làm bài và câu trả lời.", "Attempt thuộc Exam và Student; Answer thuộc Attempt và Question."],
        ["MonitoringEvent, ExamEvent", "Sự kiện giám sát và lịch sử attempt.", "Gắn với attempt, dùng timeline."],
        ["FraudSignal, FraudWarning, ProctorFlag", "Tín hiệu, cảnh báo và flag review.", "Gắn với attempt/exam, dùng risk scoring."],
        ["RiskScoreLog", "Snapshot điểm rủi ro.", "Lưu breakdown theo thời gian."],
        ["ExamImportSession, ExamQuestionRender", "Phiên import đề và ảnh crop câu hỏi.", "Gắn user/exam, parse result JSON."],
        ["StudentIdentityCheck", "Kết quả xác minh danh tính.", "Gắn attempt/student, evidence refs."],
    ], [1.8, 2.5, 3.0])

    add_heading(doc, "14. Điểm mạnh, hạn chế và hướng phát triển", 1)
    add_table(doc, ["Nhóm", "Điểm mạnh", "Hạn chế hiện tại", "Hướng phát triển"], [
        ["Giám sát", "Event-driven, realtime, có fallback polling, có audit/timeline.", "Một số signal phụ thuộc trình duyệt nên có thể bị hạn chế bởi OS/browser.", "Thêm native lockdown browser hoặc extension nếu cần mức bảo mật cao hơn."],
        ["Risk scoring", "Minh bạch, giải thích được, có dedup/cap giảm spam.", "Trọng số rule-based, chưa học từ dữ liệu lịch sử.", "Huấn luyện model ML từ dữ liệu đã review để hiệu chỉnh trọng số."],
        ["AI camera", "Có pipeline face/eye/gaze/quality/spoofing và evidence.", "Camera AI có false positive khi ánh sáng kém hoặc camera yếu.", "Thêm calibration đầu bài, benchmark dataset, threshold theo thiết bị."],
        ["Import đề", "Hỗ trợ nhiều template, fallback, report confidence.", "File đề quá dị biệt vẫn cần chỉnh thủ công.", "Thêm UI mapping câu hỏi và học template mới từ correction."],
        ["Chấm điểm", "Tự động, có thống kê lớp và IRT xấp xỉ.", "IRT hiện là approximation, chưa estimate item parameters đầy đủ.", "Thu thập dữ liệu nhiều kỳ thi để áp dụng IRT/CTT đầy đủ."],
        ["Bảo mật", "JWT, refresh token, role guard, backend authorization.", "Token lưu localStorage có rủi ro XSS nếu app bị chèn script.", "Chuyển refresh token sang HttpOnly cookie và tăng CSP."],
    ], [1.2, 2.0, 2.0, 2.2])

    add_heading(doc, "15. Kịch bản demo đề xuất khi bảo vệ", 1)
    add_numbered(doc, [
        "Đăng nhập giáo viên, tạo lớp và thêm học sinh.",
        "Tạo đề thi hoặc import đề từ PDF/DOCX, xem preview câu hỏi.",
        "Cấu hình giám sát: tab switch, fullscreen, copy/paste, duplicate IP, AI camera.",
        "Đăng nhập học sinh, vào phòng chờ, bắt đầu làm bài.",
        "Trong lúc làm bài, thử chuyển tab/thoát fullscreen/copy paste để tạo signal.",
        "Mở dashboard giáo viên xem realtime alert, risk score, timeline.",
        "Thử camera frame bất thường nếu môi trường demo cho phép.",
        "Giáo viên gửi warning hoặc pause/resume attempt.",
        "Học sinh nộp bài, xem điểm, giáo viên xem báo cáo kết quả.",
    ])

    add_heading(doc, "16. Bộ câu hỏi bảo vệ khóa luận", 1)
    add_p(doc, "Các câu hỏi dưới đây được chia theo nhóm. Mỗi câu có gợi ý trả lời ngắn để luyện phản xạ. Khi bảo vệ, không nên đọc nguyên văn; nên trả lời theo cấu trúc: ý chính -> code/module liên quan -> lý do thiết kế -> hạn chế/hướng cải tiến nếu bị hỏi sâu.")

    question_groups = [
        ("A. Tổng quan và kiến trúc", [
            ("Mục tiêu chính của đề tài là gì?", "Xây dựng hệ thống thi trực tuyến có quản lý đề, lớp, làm bài, chấm điểm, giám sát realtime và hỗ trợ phát hiện gian lận bằng rule-based risk scoring kết hợp AI camera."),
            ("Vì sao chia thành frontend, backend, AI service và parser service?", "Backend giữ nghiệp vụ và dữ liệu; frontend giữ UI; AI/parser dùng Python và thư viện đặc thù nên tách service để dễ scale, dễ deploy và tránh làm nặng backend Java."),
            ("Điểm mới của hệ thống so với thi online thông thường là gì?", "Không chỉ làm bài và chấm điểm, hệ thống có phòng chờ, realtime monitoring, risk scoring, camera AI, timeline evidence và import đề tự động."),
            ("Luồng dữ liệu chính từ học sinh tới giáo viên là gì?", "Học sinh gửi answer, monitoring event, heartbeat, camera frame -> backend lưu và tính risk -> WebSocket publish alert -> dashboard giáo viên cập nhật."),
            ("Nếu một service phụ như AI service bị lỗi thì hệ thống có dừng không?", "Không. Backend có fallback: frame vẫn được publish/ack, identity có `NEEDS_REVIEW`, parser có fallback hoặc báo report lỗi; bài thi vẫn vận hành phần lõi."),
        ]),
        ("B. Backend Spring Boot", [
            ("Backend dùng mô hình phân lớp như thế nào?", "Controller nhận request, Service xử lý nghiệp vụ, Repository truy cập DB, DTO định nghĩa dữ liệu vào/ra, Entity map bảng."),
            ("Tại sao logic không đặt trực tiếp trong Controller?", "Để controller mỏng, dễ test, tránh lặp logic, transaction và authorization nằm ở service rõ ràng hơn."),
            ("Cách hệ thống kiểm tra quyền giáo viên với đề thi?", "Service kiểm tra role teacher/admin và `exam.createdBy.id` phải trùng actor, hoặc admin được phép rộng hơn."),
            ("Vì sao vẫn cần backend authorization dù frontend đã guard route?", "Frontend có thể bị bypass bằng gọi API trực tiếp; backend là lớp bảo vệ cuối cùng và đáng tin cậy."),
            ("Transaction được dùng ở đâu và vì sao?", "Các thao tác tạo attempt, nộp bài, pause/resume, confirm import cần atomic để DB không bị trạng thái nửa vời."),
        ]),
        ("C. Xác thực và phân quyền", [
            ("JWT và refresh token hoạt động ra sao?", "Access token dùng gọi API; khi sắp hết hạn frontend gọi refresh để lấy token mới; nếu refresh fail thì logout/invalidate session."),
            ("Các vai trò trong hệ thống là gì?", "Admin quản trị toàn hệ thống; Teacher quản lý lớp/đề/giám sát; Student làm bài, xem lịch sử, tham gia lớp."),
            ("Google OAuth2 được dùng để làm gì?", "Hỗ trợ đăng nhập nhanh qua tài khoản Google, giảm ma sát người dùng nhưng vẫn map về user/role nội bộ."),
            ("Rủi ro khi lưu token localStorage là gì?", "Dễ bị đánh cắp nếu có XSS. Hướng cải tiến là HttpOnly cookie, CSP và kiểm soát script."),
            ("Nếu user chưa có role thì xử lý sao?", "Router đưa về `/select-role`; backend cũng cần đảm bảo API phù hợp role mới được gọi."),
        ]),
        ("D. Tạo đề và quản lý lớp", [
            ("Mã đề được sinh như thế nào?", "`ExamService` dùng `SecureRandom`, ký tự A-Z/0-9, dài 8, lặp tới khi không trùng DB."),
            ("Vì sao cần validate startTime < endTime?", "Tránh đề có lịch thi vô nghĩa, gây lỗi join/start và tính deadline."),
            ("Cách gán đề cho lớp?", "Exam có `classEntity`/`className`; service kiểm tra lớp thuộc giáo viên và học sinh phải enrolled mới truy cập được."),
            ("Shuffle câu hỏi/đáp án dùng để làm gì?", "Giảm khả năng học sinh nhìn bài nhau theo cùng thứ tự câu/đáp án."),
            ("Tại sao có cả đề luyện tập?", "Cho học sinh tự tạo bài luyện hoặc hệ thống sinh đề luyện, tách với đề thi chính bằng flag/practice title."),
        ]),
        ("E. Làm bài và nộp bài", [
            ("Khác nhau giữa `prepareAttempt` và `startAttempt`?", "`prepareAttempt` tạo/khôi phục phòng chờ `WAITING`; `startAttempt` chuyển sang `IN_PROGRESS`, ghi startedAt và bắt đầu giám sát."),
            ("Nếu học sinh refresh trang khi đang thi thì sao?", "Frontend lưu query context; backend nếu attempt đang `IN_PROGRESS`/`PAUSED` sẽ trả lại attempt hiện tại thay vì tạo mới."),
            ("Auto-submit hoạt động thế nào?", "Nếu quá deadline, backend lưu đáp án hiện tại nếu có, tính điểm từ draft và set status `AUTO_SUBMITTED`."),
            ("Tại sao attempt bị STOPPED khi IP thay đổi?", "Đó là kiểm soát tính nhất quán phiên thi. IP khác có thể là đổi mạng hoặc chia sẻ phiên, nên hệ thống dừng để giáo viên review."),
            ("Draft answer có ý nghĩa gì?", "Giảm mất bài khi mạng/browser lỗi và làm cơ sở auto-submit khi hết giờ."),
        ]),
        ("F. Chấm điểm và phân tích kết quả", [
            ("Cách chấm câu trắc nghiệm?", "So sánh selectedAnswer với correctAnswer; đúng thì cộng scoreWeight, sai thì 0."),
            ("ScoreWeight dùng để làm gì?", "Cho phép câu khó hoặc quan trọng có trọng số điểm cao hơn."),
            ("IRT theta trong hệ thống có ý nghĩa gì?", "Là chỉ số năng lực tương đối so với điểm trung bình lớp bằng logit transform; dùng tham khảo, không thay thế điểm chính thức."),
            ("Percentile/rank được tính để làm gì?", "Giúp giáo viên/học sinh hiểu vị trí tương đối trong lớp."),
            ("Hạn chế của chấm tự động là gì?", "Tốt cho trắc nghiệm; câu tự luận cần AI/rubric hoặc giáo viên chấm thủ công để đảm bảo công bằng."),
        ]),
        ("G. Giám sát và realtime", [
            ("Frontend gửi event giám sát thế nào?", "`useProctoringSession` gom event vào batch 1500ms, gửi kèm browser context, fingerprint, timestamp."),
            ("Heartbeat có tác dụng gì?", "Cho biết phiên còn sống, trạng thái mạng/camera/mic, phát hiện stale/network instability."),
            ("Nếu WebSocket lỗi thì dashboard giáo viên có mất dữ liệu không?", "Không hoàn toàn; frontend chuyển sang polling 8 giây để reconcile danh sách attempt."),
            ("Vì sao dùng WebSocket/STOMP thay vì chỉ polling?", "WebSocket giảm độ trễ, phù hợp cảnh báo realtime; polling dùng làm fallback chịu lỗi."),
            ("Giáo viên có những hành động nào khi thấy vi phạm?", "Warning, pause, resume, invalidate/stop attempt; tất cả được ghi event/audit."),
        ]),
        ("H. Risk scoring", [
            ("Risk scoring giải quyết bài toán gì?", "Tổng hợp nhiều tín hiệu giám sát thành điểm rủi ro dễ hiểu để giáo viên ưu tiên review."),
            ("Vì sao dùng rule-based thay vì ML hoàn toàn?", "Dễ giải thích, phù hợp dữ liệu ban đầu ít, hội đồng/giáo viên hiểu được vì sao bị cảnh báo."),
            ("Dedup window là gì?", "Cửa sổ gom các signal cùng category; trong một cluster chỉ lấy điểm cao nhất để tránh spam tăng điểm."),
            ("Tại sao mỗi category có cap điểm?", "Không để một nhóm hành vi duy nhất làm điểm tăng vô hạn, giúp cân bằng giữa các nguồn rủi ro."),
            ("CLEAN/SUSPICIOUS/HIGH_RISK/CRITICAL được xác định ra sao?", "Theo threshold: 0-19, 20-49, 50-74, 75-100."),
            ("Risk score có phải kết luận gian lận không?", "Không. Nó là chỉ báo rủi ro và bằng chứng hỗ trợ; giáo viên vẫn review cuối cùng."),
            ("VISUAL_IDENTITY khác IDENTITY thế nào?", "IDENTITY là thiết bị/IP/phiên; VISUAL_IDENTITY là tín hiệu camera như mặt, gaze, deepfake."),
            ("Khi nào tạo ProctorFlag?", "Khi score đạt mức high risk trở lên theo threshold, backend đồng bộ flag open để giáo viên review."),
            ("Tại sao behaviorScore bị giới hạn 70?", "Để hành vi trình duyệt không một mình đẩy tổng điểm tuyệt đối; identity/visual vẫn có vai trò riêng."),
            ("False positive xử lý thế nào?", "Dedup, cap, threshold, evidence/timeline và giáo viên review giúp giảm kết luận sai."),
        ]),
        ("I. AI camera", [
            ("Pipeline AI camera gồm những bước nào?", "Decode ảnh -> detect face -> eye/gaze -> frame quality -> spoof/deepfake -> occlusion -> sinh signal."),
            ("Face detection dùng thuật toán gì?", "Ưu tiên OpenCV DNN Caffe SSD face detector; fallback Haar Cascade nếu DNN không sẵn sàng."),
            ("Làm sao tránh đếm trùng khuôn mặt?", "Merge bounding box bằng IoU hoặc khoảng cách tâm gần nhau."),
            ("Brightness/blur đo bằng gì?", "Brightness là mean của ảnh grayscale; blur/quality dùng variance, thấp quá thì ảnh mờ/kém."),
            ("Gaze off-screen có đáng tin tuyệt đối không?", "Không. Nó cần nhiều frame/thời gian và chỉ là signal rủi ro, không phải bằng chứng duy nhất."),
            ("Khi AI service không khả dụng thì sao?", "Backend trả trạng thái AI_UNAVAILABLE/NEEDS_REVIEW, vẫn publish frame và không làm sập bài thi."),
            ("Evidence image dùng để làm gì?", "Lưu ảnh đại diện khi có cảnh báo để giáo viên xem lại, tăng tính giải trình."),
            ("Deepfake/spoofing phát hiện như thế nào?", "Có thể dùng model nếu sẵn; nếu không dùng heuristic như flat image, screen replay, printed photo."),
            ("Tại sao cần camera calibration đầu bài?", "Để kiểm tra camera/mic/ánh sáng trước khi thi, giảm cảnh báo sai trong lúc làm bài."),
            ("Hạn chế của AI camera là gì?", "Phụ thuộc chất lượng camera, ánh sáng, góc mặt, CPU; cần review người thật."),
        ]),
        ("J. Xác minh danh tính", [
            ("Identity verification yêu cầu dữ liệu gì?", "Ảnh giấy tờ và ảnh selfie, kèm attemptId/studentId."),
            ("Nếu giấy tờ không đọc được thì sao?", "AI/Backend trả status cần review và có thể sinh signal `IDENTITY_DOCUMENT_UNREADABLE`."),
            ("Tại sao lưu StudentIdentityCheck?", "Để có lịch sử xác minh, trạng thái, evidence refs và phục vụ review sau thi."),
            ("Vì sao không tự động cho VERIFIED khi AI fail?", "Để tránh bỏ lọt rủi ro; fallback an toàn là NEEDS_REVIEW."),
            ("Bảo mật dữ liệu định danh cần lưu ý gì?", "Giới hạn quyền truy cập evidence, lưu tối thiểu cần thiết, mã hóa/ẩn thông tin nhạy cảm nếu triển khai thật."),
        ]),
        ("K. Import đề và parser", [
            ("ParserRouter chọn parser thế nào?", "Build profile/sample, gọi `can_handle` từng parser, sort điểm giảm dần, chọn cao nhất."),
            ("Fallback parser dùng khi nào?", "Khi parser chính lỗi hoặc không parse ra câu hỏi hợp lệ."),
            ("ParseReport có vai trò gì?", "Báo số câu, số đáp án, warning/error, confidence và thời gian parse để giáo viên kiểm tra chất lượng."),
            ("Vì sao PDF/DOCX parse bằng Python mà không làm trong Java?", "Python có PyMuPDF/python-docx và xử lý text/layout linh hoạt hơn; tách service giảm phức tạp backend."),
            ("Nếu file đề không đúng template thì sao?", "Parser có thể trả 0 câu hoặc warning; frontend cho giáo viên review/chỉnh file hoặc nhập thủ công."),
            ("Import confirm làm gì?", "Chuyển parsed questions sang `Question` entity, gán vào exam hiện có hoặc tạo exam mới."),
            ("Cách xử lý công thức toán?", "Parser có module math/latex normalization để giữ nội dung toán tốt hơn trong latexContent/latexOptions khi có."),
            ("Tại sao cần preview trước khi import chính thức?", "Vì parse tự động có thể sai, giáo viên cần kiểm tra trước khi lưu vào đề thi."),
        ]),
        ("L. Database và dữ liệu", [
            ("Vì sao dùng JSONB cho một số trường?", "Signal evidence, parse result, options có cấu trúc linh hoạt; JSONB lưu được payload biến đổi mà không sửa schema liên tục."),
            ("Quan hệ ExamAttempt và Answer?", "Một attempt có nhiều answer; mỗi answer gắn với một question để chấm và báo cáo."),
            ("RiskScoreLog khác riskScore trên ExamAttempt?", "`ExamAttempt.riskScore` là trạng thái hiện tại; `RiskScoreLog` là lịch sử snapshot theo thời gian."),
            ("MonitoringEvent và FraudSignal khác nhau?", "MonitoringEvent là event thô/lịch sử; FraudSignal là tín hiệu đã chuẩn hóa để tính rủi ro."),
            ("Tại sao cần AuditLog?", "Ghi dấu hành động quan trọng để truy vết, giải trình khi có tranh chấp."),
        ]),
        ("M. Kiểm thử, bảo mật, triển khai", [
            ("Dự án có test những phần nào?", "Frontend có Vitest cho composables/utils; parser có pytest cho template/text normalizer; backend nên bổ sung unit/integration test cho service chính."),
            ("Rate limit event giám sát để làm gì?", "Chống spam event làm nặng backend và tăng điểm rủi ro sai."),
            ("Docker-compose giúp gì?", "Chạy đồng bộ frontend/backend/AI/parser/database dễ hơn, giảm lỗi môi trường."),
            ("Nginx dùng ở đâu?", "Reverse proxy frontend/backend, cấu hình host/SSL/port khi deploy."),
            ("Nếu triển khai thật cần bổ sung gì?", "HTTPS, secure cookie, CSP, log monitoring, backup DB, giới hạn upload, mã hóa evidence, chính sách dữ liệu cá nhân."),
        ]),
    ]

    q_index = 1
    for group_title, questions in question_groups:
        add_heading(doc, group_title, 2)
        for question, answer in questions:
            add_p(doc, f"Câu {q_index}. {question}", bold_prefix=f"Câu {q_index}.")
            add_p(doc, f"Gợi ý trả lời: {answer}", bold_prefix="Gợi ý trả lời:")
            q_index += 1

    add_heading(doc, "17. Câu hỏi phản biện khó và cách trả lời", 1)
    hard_questions = [
        ("Hệ thống có đảm bảo phát hiện 100% gian lận không?", "Không hệ thống online nào đảm bảo 100%. Đề tài đặt mục tiêu giảm rủi ro và tăng khả năng phát hiện bằng nhiều lớp: trình duyệt, heartbeat, IP/fingerprint, AI camera, realtime review và audit."),
        ("AI camera có thể sai, vậy dùng làm gì?", "AI không phải người phán quyết mà là bộ lọc cảnh báo. Evidence, timeline và giáo viên review giúp cân bằng giữa tự động hóa và tính công bằng."),
        ("Rule-based risk scoring có chủ quan không?", "Có yếu tố thiết kế trọng số, nhưng ưu điểm là minh bạch và có thể hiệu chỉnh. Khi có dữ liệu review thật, hệ thống có thể học lại trọng số hoặc kết hợp ML."),
        ("Nếu học sinh dùng điện thoại thứ hai thì hệ thống có biết không?", "Không thể phát hiện chắc chắn bằng web app nếu thiết bị ngoài không nằm trong camera/gaze. Hệ thống có thể phát hiện gián tiếp qua gaze, face turned away, audio/speaking, nhưng cần quy chế phòng thi và giám thị."),
        ("Vì sao không khóa hoàn toàn tab/trình duyệt?", "Web browser không cho ứng dụng web kiểm soát tuyệt đối vì lý do bảo mật. Hệ thống phát hiện và ghi nhận hành vi thay vì cưỡng chế hoàn toàn; muốn khóa cứng cần lockdown browser/native app."),
        ("Dữ liệu camera và giấy tờ có nhạy cảm không?", "Có. Khi triển khai thật cần policy rõ, giới hạn thời gian lưu, phân quyền evidence, mã hóa storage và tuân thủ quy định bảo vệ dữ liệu cá nhân."),
        ("Parser đề thi có thể sai thì có ảnh hưởng chất lượng đề không?", "Có thể, nên hệ thống có preview/report/confidence và yêu cầu giáo viên confirm trước khi lưu thành đề chính thức."),
        ("Nếu mất mạng khi đang thi thì sao?", "Heartbeat phát hiện gián đoạn; draft answer giảm mất dữ liệu; khi kết nối lại có session recovery. Nếu quá hạn hệ thống auto-submit từ nháp."),
        ("Tại sao không dùng microservice toàn bộ?", "Dự án tách những phần cần thiết: AI/parser vì khác stack và nặng. Backend nghiệp vụ giữ monolith module hóa để đơn giản triển khai và bảo trì trong phạm vi khóa luận."),
        ("Điểm khoa học của đề tài nằm ở đâu?", "Ở mô hình giám sát đa tín hiệu, thuật toán risk scoring có dedup/cap/threshold giải thích được, pipeline AI camera và cơ chế import đề template-based có scoring/fallback."),
    ]
    add_table(doc, ["Câu hỏi khó", "Cách trả lời gợi ý"], [[q, a] for q, a in hard_questions], [3.0, 4.5])

    add_heading(doc, "18. Checklist ôn bảo vệ", 1)
    add_bullets(doc, [
        "Nắm 5 file backend chính: `ExamService`, `SubmissionService`, `RiskScoringService`, `MonitoringService`, `AiAssistService`.",
        "Nắm 3 file frontend chính: `router.js`, `apiClient.js`, `useProctoringSession.js` hoặc `useExamMonitoring.js`.",
        "Nắm 2 file Python chính: `ai-service/app/proctor.py`, `python_parser/app/parser_router.py`.",
        "Thuộc luồng: tạo đề -> phòng chờ -> bắt đầu -> làm bài -> monitoring -> risk -> warning/pause -> nộp -> report.",
        "Giải thích được công thức risk scoring và ví dụ dedup cluster.",
        "Chuẩn bị demo có dữ liệu mẫu để tạo ít nhất một alert realtime.",
        "Khi bị hỏi về hạn chế, trả lời thẳng: AI có false positive, browser không khóa tuyệt đối, parser phụ thuộc template, token localStorage cần cải tiến.",
        "Kết thúc bằng hướng phát triển: ML weight tuning, lockdown browser, calibration AI, bảo mật evidence, thêm test tự động.",
    ])

    add_heading(doc, "19. Tóm tắt 1 phút để mở đầu thuyết trình", 1)
    add_p(doc, "Đề tài của em xây dựng hệ thống thi trực tuyến có giám sát và hỗ trợ phát hiện gian lận. Hệ thống gồm frontend Vue, backend Spring Boot, AI service Python và parser service Python. Giáo viên có thể tạo hoặc import đề, phân công cho lớp và theo dõi phòng thi realtime. Học sinh làm bài trên giao diện web, hệ thống tự lưu nháp, tự nộp khi hết giờ và chấm điểm tự động. Điểm trọng tâm của đề tài là cơ chế giám sát đa tín hiệu: hành vi trình duyệt, heartbeat, IP/fingerprint và AI camera được chuẩn hóa thành signal, sau đó thuật toán risk scoring tính điểm rủi ro theo nhóm, có dedup window và giới hạn điểm để giảm cảnh báo nhiễu. Kết quả không thay thế quyết định của giáo viên mà cung cấp dashboard, timeline và evidence để giáo viên review công bằng hơn.")

    add_heading(doc, "20. Phụ lục: mốc code nên mở khi bị hỏi", 1)
    add_table(doc, ["Chủ đề", "File và mốc hàm/dòng đã khảo sát"], [
        ["Risk scoring", "RiskScoringService.java: `recomputeRisk` khoảng dòng 231, `computeCategoryScores` khoảng dòng 357, `computeClusteredScore` khoảng dòng 391, `resolveLevel` khoảng dòng 531."],
        ["Làm bài/nộp bài", "SubmissionService.java: `prepareAttempt` khoảng dòng 75, `startAttempt` khoảng dòng 115, `submitAttempt` khoảng dòng 308, `autoSubmitFromDraft` khoảng dòng 807, `enforceIpConsistency` khoảng dòng 914."],
        ["AI bridge", "AiAssistService.java: `analyzeFrame` khoảng dòng 145, `verifyIdentity` khoảng dòng 245, `safeBridgeAiSignals` khoảng dòng 652, `clusterAiCameraSignalsForRecording` khoảng dòng 828."],
        ["AI image processing", "proctor.py: `analyze_frame` khoảng dòng 321, `_detect_face_locations` khoảng dòng 2129, `_detect_faces_dnn` khoảng dòng 2149, `analyze_behavior` khoảng dòng 2411."],
        ["Parser router", "parser_router.py: `ParserRouter` khoảng dòng 43, `route` khoảng dòng 55, `_route_docx` khoảng dòng 119, `_score_all_parsers` khoảng dòng 236, `_fallback_parse` khoảng dòng 258."],
    ], [2.0, 5.5])

    return doc


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
