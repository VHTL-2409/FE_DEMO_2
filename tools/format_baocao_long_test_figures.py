from __future__ import annotations

import re
from copy import deepcopy
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
INPUT_DOCX = ROOT / "baocao" / "baocao_long_test.docx"
OUTPUT_DOCX = ROOT / "baocao" / "baocao_long_test_format_anh_clean_textportrait_v3.docx"

EMU_PER_INCH = 914400
A4_PORTRAIT_CM = (21.0, 29.7)
A4_LANDSCAPE_CM = (29.7, 21.0)
MARGIN_CM = 1.0
CAPTION_RESERVE_IN = 0.55

TWIPS_PER_INCH = 1440
EMU_PER_CM = EMU_PER_INCH / 2.54
TWIPS_PER_CM = TWIPS_PER_INCH / 2.54


@dataclass
class FigureBlock:
    start: int
    caption: int | None
    wide: bool
    has_caption: bool
    placeholder_caption: bool
    caption_text: str
    generated_text: str
    insert_before: bool = True
    chapter: int = 1
    caption_body: str = ""
    display_caption: str = ""
    anchor: int = 0
    anchor_is_title: bool = False
    restore_portrait_after: bool = False

    @property
    def end(self) -> int:
        return self.caption if self.caption is not None else self.start


def clean_text(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def paragraph_has_drawing(paragraph) -> bool:
    return bool(paragraph._p.xpath(".//w:drawing"))


def drawing_count(paragraph) -> int:
    return len(paragraph._p.xpath(".//w:drawing"))


def is_caption(text: str) -> bool:
    value = clean_text(text)
    return value.startswith(("Hình", "Hinh", "Hìbh"))


def is_placeholder_caption(text: str) -> bool:
    value = clean_text(text)
    return is_caption(value) and (len(value) < 16 or "…" in value or ".." in value)


def is_diagram_group_title(text: str) -> bool:
    value = clean_text(text).lower()
    if not value:
        return False
    return bool(
        re.match(r"^(đặc\s*tả\s*use\s*case|đặc\s*tả\s*usecase|sơ\s*đồ\s*tuần\s*tự|sơ\s*đồ\s*hoạt\s*động)\b", value)
    )


def normalize_caption(text: str) -> str:
    value = clean_text(text)
    value = value.replace("Hìbh", "Hình")
    if value.startswith("Hinh"):
        value = "Hình" + value[4:]
    value = re.sub(r"\s+\.", ".", value)
    value = re.sub(r"\s+,", ",", value)
    value = re.sub(r"\s+-\s+", " - ", value)
    value = value.replace("tuần tạm", "tuần tự tạm")
    value = value.replace("quản lí", "quản lý")
    value = value.replace("chỉnh sửa , xóa , tạo", "chỉnh sửa, xóa, tạo")
    value = value.replace("xóa -  tạo", "xóa - tạo")
    return value.strip()


def roman_to_int(value: str) -> int | None:
    mapping = {"I": 1, "V": 5, "X": 10, "L": 50, "C": 100, "D": 500, "M": 1000}
    total = 0
    prev = 0
    for char in reversed(value.upper()):
        current = mapping.get(char)
        if current is None:
            return None
        if current < prev:
            total -= current
        else:
            total += current
            prev = current
    return total if total > 0 else None


def parse_chapter_number(text: str) -> int | None:
    value = clean_text(text).upper()
    match = re.search(r"CHƯƠNG\s+([IVXLCDM]+|\d+)", value)
    if not match:
        return None
    token = match.group(1)
    if token.isdigit():
        return int(token)
    return roman_to_int(token)


def nearest_chapter(paragraphs, index: int) -> int:
    for i in range(index, -1, -1):
        chapter = parse_chapter_number(paragraphs[i].text)
        if chapter is not None:
            return chapter
    return 1


def strip_caption_number(text: str) -> str:
    value = normalize_caption(text)
    value = re.sub(r"^H(?:ình|inh|ìbh)\s+\d+(?:\.\d+)*\.?\s*", "", value, flags=re.IGNORECASE)
    return clean_text(value)


def caption_key(text: str) -> str:
    return re.sub(r"\s+", " ", clean_text(text).rstrip(".")).lower()


def apply_part_suffix(text: str, part_index: int) -> str:
    value = clean_text(text)
    if value.endswith("."):
        value = value[:-1].rstrip()
    return f"{value} (phần {part_index})."


def build_display_caption(chapter: int, number: int, body: str) -> str:
    body = clean_text(body)
    if not body:
        body = "Hình minh họa"
    if not body.endswith("."):
        body += "."
    return f"Hình {chapter}.{number}. {body}"


def prune_empty_paragraphs(document: Document) -> int:
    removed = 0
    for paragraph in reversed(document.paragraphs):
        if paragraph_has_drawing(paragraph):
            continue
        if paragraph._p.xpath("./w:pPr/w:sectPr"):
            continue
        if clean_text(paragraph.text):
            continue
        paragraph._p.getparent().remove(paragraph._p)
        removed += 1
    return removed


def format_diagram_group_titles(document: Document) -> None:
    for paragraph in document.paragraphs:
        if is_diagram_group_title(paragraph.text):
            set_title_keep_with_next(paragraph)


def split_title_text_from_image_paragraphs(document: Document) -> int:
    split_count = 0
    for paragraph in reversed(document.paragraphs):
        if not paragraph_has_drawing(paragraph):
            continue
        text = clean_text(paragraph.text)
        if not text or not is_diagram_group_title(text):
            continue

        parent = paragraph._p.getparent()
        insert_at = parent.index(paragraph._p)
        new_p = OxmlElement("w:p")
        ppr = deepcopy(paragraph._p.find(qn("w:pPr"))) if paragraph._p.find(qn("w:pPr")) is not None else None
        if ppr is not None:
            new_p.append(ppr)
        add_run(new_p, text, bold=True, size_half_points=22)
        set_title_keep_with_next(new_p)
        parent.insert(insert_at, new_p)
        remove_text_runs_keep_drawings(paragraph)
        split_count += 1
    return split_count


def nearest_heading(paragraphs, index: int) -> str:
    for i in range(index - 1, -1, -1):
        paragraph = paragraphs[i]
        if paragraph.style and paragraph.style.name.startswith("Heading"):
            text = clean_text(paragraph.text)
            if text:
                return text
    return "chức năng hệ thống"


def strip_heading_number(text: str) -> str:
    return re.sub(r"^\d+(\.\d+)*\.?\s*", "", clean_text(text)).strip()


def diagram_kind_before(paragraphs, index: int) -> str:
    start = max(0, index - 5)
    for i in range(index - 1, start - 1, -1):
        text = clean_text(paragraphs[i].text).lower()
        if "tuần tự" in text:
            return "Sơ đồ tuần tự"
        if "hoạt động" in text:
            return "Sơ đồ hoạt động"
        if "usecase" in text or "use case" in text:
            return "Sơ đồ Use Case"
    return "Sơ đồ minh họa"


def generated_caption(paragraphs, index: int) -> str:
    heading = strip_heading_number(nearest_heading(paragraphs, index))
    kind = diagram_kind_before(paragraphs, index)
    return f"Hình minh họa - {kind} {heading}."


def paragraph_has_real_text_between(paragraphs, start: int, end: int) -> bool:
    for i in range(start, end):
        if i < 0 or i >= len(paragraphs):
            continue
        paragraph = paragraphs[i]
        if paragraph_has_drawing(paragraph):
            return True
        text = clean_text(paragraph.text)
        if text and not is_caption(text):
            return True
    return False


def get_extent(drawing) -> tuple[int, int] | None:
    extents = drawing.xpath(".//wp:extent")
    if not extents:
        return None
    extent = extents[0]
    try:
        return int(extent.get("cx")), int(extent.get("cy"))
    except (TypeError, ValueError):
        return None


def set_extent(drawing, cx: int, cy: int) -> None:
    for extent in drawing.xpath(".//wp:extent"):
        extent.set("cx", str(cx))
        extent.set("cy", str(cy))
    for extent in drawing.xpath(".//a:ext"):
        extent.set("cx", str(cx))
        extent.set("cy", str(cy))


def choose_wide(paragraph) -> bool:
    max_aspect = 0.0
    for drawing in paragraph._p.xpath(".//w:drawing"):
        extent = get_extent(drawing)
        if not extent:
            continue
        cx, cy = extent
        if cy:
            max_aspect = max(max_aspect, cx / cy)
    return max_aspect >= 1.22


def fit_drawings(paragraph, wide: bool) -> None:
    if wide:
        page_w_cm, page_h_cm = A4_LANDSCAPE_CM
    else:
        page_w_cm, page_h_cm = A4_PORTRAIT_CM
    max_w = int((page_w_cm - MARGIN_CM * 2) * EMU_PER_CM)
    max_h = int((page_h_cm - MARGIN_CM * 2) * EMU_PER_CM - CAPTION_RESERVE_IN * EMU_PER_INCH)

    drawings = paragraph._p.xpath(".//w:drawing")
    if len(drawings) > 1:
        max_h = int(max_h / len(drawings))

    for drawing in drawings:
        extent = get_extent(drawing)
        if not extent:
            continue
        cx, cy = extent
        if cx <= 0 or cy <= 0:
            continue
        aspect = cx / cy
        new_w = min(max_w, int(max_h * aspect))
        new_h = int(new_w / aspect)
        if new_h > max_h:
            new_h = max_h
            new_w = int(new_h * aspect)
        set_extent(drawing, new_w, new_h)


def ensure_ppr(p):
    ppr = p.find(qn("w:pPr"))
    if ppr is None:
        ppr = OxmlElement("w:pPr")
        p.insert(0, ppr)
    return ppr


def remove_existing(ppr, local_names: set[str]) -> None:
    for child in list(ppr):
        if child.tag.rsplit("}", 1)[-1] in local_names:
            ppr.remove(child)


def set_paragraph_image_format(paragraph) -> None:
    ppr = ensure_ppr(paragraph._p)
    remove_existing(ppr, {"jc", "spacing", "keepNext", "keepLines"})

    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    ppr.append(jc)

    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "60")
    ppr.append(spacing)

    keep_next = OxmlElement("w:keepNext")
    ppr.append(keep_next)
    keep_lines = OxmlElement("w:keepLines")
    ppr.append(keep_lines)


def set_title_keep_with_next(paragraph) -> None:
    p = paragraph._p if hasattr(paragraph, "_p") else paragraph
    ppr = ensure_ppr(p)
    remove_existing(ppr, {"keepNext", "keepLines"})
    ppr.append(OxmlElement("w:keepNext"))
    ppr.append(OxmlElement("w:keepLines"))


def add_run(parent, text: str, *, bold: bool = False, size_half_points: int = 22) -> None:
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")

    rfonts = OxmlElement("w:rFonts")
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), "Times New Roman")
    rpr.append(rfonts)

    size = OxmlElement("w:sz")
    size.set(qn("w:val"), str(size_half_points))
    rpr.append(size)

    size_cs = OxmlElement("w:szCs")
    size_cs.set(qn("w:val"), str(size_half_points))
    rpr.append(size_cs)

    if bold:
        rpr.append(OxmlElement("w:b"))
        rpr.append(OxmlElement("w:bCs"))

    run.append(rpr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    run.append(t)
    parent.append(run)


def style_caption_paragraph_element(p, text: str | None = None, *, replace_text: bool = True) -> None:
    if text is not None and replace_text:
        ppr = deepcopy(p.find(qn("w:pPr"))) if p.find(qn("w:pPr")) is not None else None
        for child in list(p):
            p.remove(child)
        if ppr is not None:
            p.append(ppr)
        add_run(p, text, bold=True, size_half_points=22)

    ppr = ensure_ppr(p)
    remove_existing(ppr, {"jc", "spacing", "keepLines"})

    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    ppr.append(jc)

    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "120")
    ppr.append(spacing)

    ppr.append(OxmlElement("w:keepLines"))

    for run in p.xpath(".//w:r"):
        rpr = run.find(qn("w:rPr"))
        if rpr is None:
            rpr = OxmlElement("w:rPr")
            run.insert(0, rpr)
        if rpr.find(qn("w:b")) is None:
            rpr.append(OxmlElement("w:b"))
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.insert(0, rfonts)
        for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
            rfonts.set(qn(f"w:{attr}"), "Times New Roman")
        for tag in ("w:sz", "w:szCs"):
            size = rpr.find(qn(tag))
            if size is None:
                size = OxmlElement(tag)
                rpr.append(size)
            size.set(qn("w:val"), "22")


def remove_text_runs_keep_drawings(paragraph) -> None:
    for run in list(paragraph._p.xpath("./w:r")):
        if run.xpath(".//w:drawing"):
            continue
        paragraph._p.remove(run)


def split_multi_drawing_paragraphs(document: Document) -> int:
    split_count = 0
    for paragraph in reversed(document.paragraphs):
        drawings = paragraph._p.xpath("./w:r//w:drawing")
        if len(drawings) <= 1:
            continue

        parent = paragraph._p.getparent()
        insert_at = parent.index(paragraph._p)
        original_ppr = paragraph._p.find(qn("w:pPr"))

        for drawing in drawings:
            new_p = OxmlElement("w:p")
            if original_ppr is not None:
                new_p.append(deepcopy(original_ppr))

            new_r = OxmlElement("w:r")
            run = drawing
            while run is not None and run.tag != qn("w:r"):
                run = run.getparent()
            if run is not None:
                rpr = run.find(qn("w:rPr"))
                if rpr is not None:
                    new_r.append(deepcopy(rpr))
            new_r.append(deepcopy(drawing))
            new_p.append(new_r)
            parent.insert(insert_at, new_p)
            insert_at += 1
            split_count += 1

        parent.remove(paragraph._p)
    return split_count


def make_caption_paragraph(text: str):
    p = OxmlElement("w:p")
    style_caption_paragraph_element(p, text, replace_text=True)
    return p


def section_break_paragraph(*, wide: bool):
    p = OxmlElement("w:p")
    ppr = OxmlElement("w:pPr")
    sect = OxmlElement("w:sectPr")

    type_el = OxmlElement("w:type")
    type_el.set(qn("w:val"), "nextPage")
    sect.append(type_el)

    if wide:
        width_cm, height_cm = A4_LANDSCAPE_CM
        orient = "landscape"
    else:
        width_cm, height_cm = A4_PORTRAIT_CM
        orient = None

    pg_sz = OxmlElement("w:pgSz")
    pg_sz.set(qn("w:w"), str(round(width_cm * TWIPS_PER_CM)))
    pg_sz.set(qn("w:h"), str(round(height_cm * TWIPS_PER_CM)))
    if orient:
        pg_sz.set(qn("w:orient"), orient)
    sect.append(pg_sz)

    pg_mar = OxmlElement("w:pgMar")
    margin = str(round(MARGIN_CM * TWIPS_PER_CM))
    pg_mar.set(qn("w:top"), margin)
    pg_mar.set(qn("w:right"), margin)
    pg_mar.set(qn("w:bottom"), margin)
    pg_mar.set(qn("w:left"), margin)
    pg_mar.set(qn("w:header"), str(round(0.5 * TWIPS_PER_CM)))
    pg_mar.set(qn("w:footer"), str(round(0.5 * TWIPS_PER_CM)))
    pg_mar.set(qn("w:gutter"), "0")
    sect.append(pg_mar)

    ppr.append(sect)
    p.append(ppr)
    return p


def insert_after(anchor, new_element):
    parent = anchor.getparent()
    parent.insert(parent.index(anchor) + 1, new_element)


def insert_before(anchor, new_element):
    parent = anchor.getparent()
    parent.insert(parent.index(anchor), new_element)


def build_blocks(document: Document) -> list[FigureBlock]:
    paragraphs = document.paragraphs
    blocks: list[FigureBlock] = []
    for index, paragraph in enumerate(paragraphs):
        if not paragraph_has_drawing(paragraph):
            continue

        chapter = nearest_chapter(paragraphs, index)
        caption_index: int | None = None
        caption_text = ""
        current_text = clean_text(paragraph.text)
        if is_caption(current_text):
            caption_index = index
            caption_text = current_text
        else:
            limit = min(index + 8, len(paragraphs))
            for candidate_index in range(index + 1, limit):
                candidate = paragraphs[candidate_index]
                candidate_text = clean_text(candidate.text)
                if is_caption(candidate_text):
                    caption_index = candidate_index
                    caption_text = candidate_text
                    break
                if paragraph_has_drawing(candidate):
                    break
                if candidate.style and candidate.style.name.startswith("Heading") and candidate_index > index + 1:
                    break

        has_caption = caption_index is not None
        body_source = generated_caption(paragraphs, index)
        if has_caption:
            body_source = normalize_caption(caption_text)
            if is_placeholder_caption(caption_text):
                body_source = generated_caption(paragraphs, index)

        anchor_index = index
        anchor_is_title = False
        search_start = max(0, index - 20)
        for candidate_index in range(index - 1, search_start - 1, -1):
            candidate = paragraphs[candidate_index]
            candidate_text = clean_text(candidate.text)
            if paragraph_has_drawing(candidate):
                break
            if not candidate_text:
                continue
            if is_diagram_group_title(candidate_text):
                anchor_index = candidate_index
                anchor_is_title = True

        blocks.append(
            FigureBlock(
                start=index,
                caption=caption_index,
                wide=choose_wide(paragraph),
                has_caption=has_caption,
                placeholder_caption=is_placeholder_caption(caption_text) if has_caption else False,
                caption_text=caption_text,
                generated_text=generated_caption(paragraphs, index),
                chapter=chapter,
                caption_body=strip_caption_number(body_source),
                anchor=anchor_index,
                anchor_is_title=anchor_is_title,
            )
        )

    blocks_with_breaks: list[FigureBlock] = []
    previous_end: int | None = None
    for index, block in enumerate(blocks):
        insert_before_flag = True
        if previous_end is not None:
            insert_before_flag = paragraph_has_real_text_between(paragraphs, previous_end + 1, block.anchor)
        next_anchor = blocks[index + 1].anchor if index + 1 < len(blocks) else len(paragraphs)
        restore_portrait_after = paragraph_has_real_text_between(paragraphs, block.end + 1, next_anchor)
        blocks_with_breaks.append(
            FigureBlock(
                start=block.start,
                caption=block.caption,
                wide=block.wide,
                has_caption=block.has_caption,
                placeholder_caption=block.placeholder_caption,
                caption_text=block.caption_text,
                generated_text=block.generated_text,
                insert_before=insert_before_flag,
                chapter=block.chapter,
                caption_body=block.caption_body,
                anchor=block.anchor,
                anchor_is_title=block.anchor_is_title,
                restore_portrait_after=restore_portrait_after,
            )
        )
        previous_end = block.end
    return blocks_with_breaks


def assign_display_captions(blocks: list[FigureBlock]) -> None:
    frequencies: dict[tuple[int, str], int] = {}
    for block in blocks:
        key = (block.chapter, caption_key(block.caption_body))
        frequencies[key] = frequencies.get(key, 0) + 1

    chapter_counts: dict[int, int] = {}
    part_counts: dict[tuple[int, str], int] = {}
    for block in blocks:
        chapter_counts[block.chapter] = chapter_counts.get(block.chapter, 0) + 1
        body = block.caption_body or block.generated_text
        key = (block.chapter, caption_key(body))
        if frequencies.get(key, 0) > 1:
            part_counts[key] = part_counts.get(key, 0) + 1
            body = apply_part_suffix(body, part_counts[key])
        block.display_caption = build_display_caption(block.chapter, chapter_counts[block.chapter], body)


def format_document(input_path: Path, output_path: Path) -> tuple[int, int]:
    document = Document(input_path)
    prune_empty_paragraphs(document)
    split_multi_drawing_paragraphs(document)
    split_title_text_from_image_paragraphs(document)
    format_diagram_group_titles(document)
    blocks = build_blocks(document)
    assign_display_captions(blocks)
    generated_count = 0

    for block in reversed(blocks):
        paragraphs = document.paragraphs
        image_paragraph = paragraphs[block.start]
        set_paragraph_image_format(image_paragraph)
        fit_drawings(image_paragraph, True)
        if block.anchor_is_title:
            anchor_paragraph = paragraphs[block.anchor]
            set_title_keep_with_next(anchor_paragraph)

        caption_element = None
        if block.has_caption and block.caption is not None:
            caption_paragraph = paragraphs[block.caption]
            replacement = block.display_caption
            if block.placeholder_caption:
                generated_count += 1

            if block.caption == block.start and paragraph_has_drawing(caption_paragraph):
                remove_text_runs_keep_drawings(caption_paragraph)
                caption_element = make_caption_paragraph(replacement)
                insert_after(caption_paragraph._p, caption_element)
            else:
                caption_element = caption_paragraph._p
                style_caption_paragraph_element(caption_element, replacement, replace_text=True)
        else:
            caption_element = make_caption_paragraph(block.display_caption)
            insert_after(image_paragraph._p, caption_element)
            generated_count += 1

        insert_after(caption_element, section_break_paragraph(wide=True))
        anchor_paragraph = paragraphs[block.anchor]
        if block.insert_before:
            insert_before(anchor_paragraph._p, section_break_paragraph(wide=False))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return len(blocks), generated_count


def main() -> None:
    block_count, generated_count = format_document(INPUT_DOCX, OUTPUT_DOCX)
    print(f"created {OUTPUT_DOCX}")
    print(f"figure blocks {block_count}")
    print(f"generated captions {generated_count}")


if __name__ == "__main__":
    main()
