from __future__ import annotations

from datetime import datetime
from pathlib import Path
import re

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Phan_tich_service_phat_hien_gian_lan_va_ai_service.docx"
FONT = "Times New Roman"
FONT_SIZE = Pt(13)


FILES = [
    (
        "Backend service phát hiện gian lận",
        "BE/src/main/java/com/example/demo/service/FraudSignalTypeNormalizer.java",
        "Chuẩn hóa tên signal gian lận để các service dùng cùng một bộ thuật ngữ.",
    ),
    (
        "Backend service phát hiện gian lận",
        "BE/src/main/java/com/example/demo/service/FraudSignalService.java",
        "Ghi nhận FraudSignal từ event hoặc tín hiệu hệ thống, xác định severity, confidence, riskImpact và evidence.",
    ),
    (
        "Backend service phát hiện gian lận",
        "BE/src/main/java/com/example/demo/service/RiskScoringService.java",
        "Tính điểm rủi ro theo nhóm tín hiệu, cửa sổ chống trùng, giới hạn điểm, cấp rủi ro và proctor flag.",
    ),
    (
        "Backend service phát hiện gian lận",
        "BE/src/main/java/com/example/demo/service/FraudWarningService.java",
        "Sinh cảnh báo gian lận từ signal hoặc pattern, chống trùng cảnh báo và hỗ trợ giáo viên review.",
    ),
    (
        "Backend service phát hiện gian lận",
        "BE/src/main/java/com/example/demo/service/FraudAnalysisService.java",
        "Phân tích hậu kiểm: giống đáp án, thời gian làm bài, thống kê điểm, hành vi, IP và pattern nghi vấn.",
    ),
    (
        "Backend service phát hiện gian lận",
        "BE/src/main/java/com/example/demo/service/MLRiskScoringService.java",
        "Trích xuất feature, tính anomaly score và kết hợp điểm rule-based với điểm rủi ro theo ML.",
    ),
    (
        "Backend service phát hiện gian lận",
        "BE/src/main/java/com/example/demo/service/FraudModelTrainer.java",
        "Huấn luyện và đánh giá mô hình gian lận từ dữ liệu attempt và fraud signal.",
    ),
    (
        "Backend service phát hiện gian lận",
        "BE/src/main/java/com/example/demo/service/AnswerSimilarityService.java",
        "So sánh đáp án giữa các thí sinh để phát hiện pattern giống nhau bất thường.",
    ),
    (
        "Backend service phát hiện gian lận",
        "BE/src/main/java/com/example/demo/service/DeviceFingerprintService.java",
        "Tạo và chuẩn hóa fingerprint thiết bị để phát hiện đổi thiết bị hoặc dùng chung thiết bị/IP.",
    ),
    (
        "Backend service phát hiện gian lận",
        "BE/src/main/java/com/example/demo/service/ProctorEvidenceImageService.java",
        "Lưu ảnh bằng chứng từ camera AI và tạo URL evidence cho warning hoặc signal.",
    ),
    (
        "Backend service phát hiện gian lận",
        "BE/src/main/java/com/example/demo/service/ProctorFlagService.java",
        "Truy vấn và review flag proctoring được sinh bởi risk scoring.",
    ),
    ("AI service", "ai-service/app/main.py", "FastAPI entrypoint: expose API OCR, proctor frame/behavior, generator, evaluator, analytics và chat."),
    ("AI service", "ai-service/app/schemas.py", "Pydantic schema cho request/response AI, đặc biệt FrameAnalysis và BehaviorAnalysis."),
    ("AI service", "ai-service/app/proctor.py", "Lõi AI proctoring: decode frame, detect mặt/mắt, gaze, spoofing, quality gate và signal."),
    ("AI service", "ai-service/app/ocr.py", "OCR engine đọc text từ ảnh/PDF đầu vào cho các chức năng AI phụ trợ."),
    ("AI service", "ai-service/app/analytics.py", "AI analytics dự đoán hiệu suất, chất lượng câu hỏi và phân tích dữ liệu học tập."),
    ("AI service", "ai-service/app/evaluator.py", "AI chấm hoặc tư vấn đánh giá bài tự luận theo tiêu chí."),
    ("AI service", "ai-service/app/generator.py", "AI sinh câu hỏi từ prompt hoặc nội dung đầu vào."),
    ("AI service", "ai-service/app/chat.py", "AI chat assistant dùng model provider để trả lời hội thoại."),
    ("AI service", "ai-service/app/openai_client.py", "Wrapper tạo OpenAI client theo cấu hình môi trường."),
    ("AI service", "ai-service/app/model_key_resolver.py", "Resolve model, key và provider cho các chức năng AI theo environment."),
]


SIGNAL_TERMS = {
    "TAB_SWITCH": "rời tab",
    "WINDOW_BLUR": "mất focus cửa sổ",
    "EXIT_FULLSCREEN": "thoát fullscreen",
    "COPY_PASTE": "copy/paste",
    "RIGHT_CLICK": "click phải",
    "PRINT_SCREEN": "chụp màn hình",
    "DEVTOOLS": "mở devtools",
    "NO_CAMERA": "camera tắt",
    "NO_MIC": "micro tắt",
    "FACE_NOT_DETECTED": "không thấy mặt",
    "MULTIPLE_FACES": "nhiều mặt",
    "GAZE_AWAY": "nhìn ra ngoài",
    "SPOOFING": "nghi giả mạo camera",
    "DEVICE_CHANGE": "đổi thiết bị",
    "DUPLICATE_IP": "trùng IP",
    "IDENTITY": "nhóm định danh",
    "SCREEN_LEAVE": "nhóm rời màn hình",
    "CLIPBOARD": "nhóm clipboard",
    "TECHNICAL": "nhóm kỹ thuật",
    "HEARTBEAT": "nhóm heartbeat",
    "VISUAL_IDENTITY": "nhóm định danh bằng hình ảnh",
}


def apply_font_to_run(run, size: Pt = FONT_SIZE) -> None:
    run.font.name = FONT
    run.font.size = size
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run._element.rPr.rFonts.set(qn("w:cs"), FONT)


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def text_cell(cell, text, *, bold: bool = False, color=None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(str(text))
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    apply_font_to_run(run)


def code_cell(cell, text: str) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text if text.strip() else " ")
    apply_font_to_run(run)


def header(table, columns: list[str]) -> None:
    cells = table.rows[0].cells
    for i, name in enumerate(columns):
        text_cell(cells[i], name, bold=True, color=(255, 255, 255))
        shade(cells[i], "1F4E79")


def para(doc: Document, text: str, style: str | None = None, *, bold: bool = False):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    apply_font_to_run(run)
    return p


def read_lines(path: Path) -> list[str]:
    raw = path.read_bytes()
    for enc in ("utf-8-sig", "utf-8", "cp1258", "cp1252"):
        try:
            return raw.decode(enc).splitlines()
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace").splitlines()


def file_kind(path: Path) -> str:
    name = path.name.lower()
    if name.endswith(".java"):
        return "java"
    if name.endswith(".py"):
        return "python"
    return "text"


def symbol(line: str, kind: str):
    s = line.strip()
    if kind == "java":
        if re.match(r"(public|private|protected)?\s*(class|record|enum)\s+", s):
            return "type", s.replace("{", "").strip()
        if re.match(r"(public|private|protected)\s+.*\)\s*\{?$", s) and "(" in s:
            return "method", re.sub(r"\{\s*$", "", s)
    if kind == "python":
        m = re.match(r"(class|def|async def)\s+([A-Za-z_][\w_]*)", s)
        if m:
            return m.group(1), m.group(2)
    return None


def explain(line: str, kind: str, ctx: str) -> str:
    s = line.strip()
    if not s:
        return "Dòng trống dùng để tách các khối code, giúp phân biệt import, cấu hình, hàm và nhánh xử lý."
    if s.startswith(("//", "#", "/*", "*", "*/")):
        return "Chú thích giải thích ý đồ hoặc nghiệp vụ của khối code lân cận; khi bảo vệ có thể dùng để diễn giải thiết kế."
    for term, meaning in SIGNAL_TERMS.items():
        if term in s:
            return f"Liên quan `{term}` ({meaning}); đây là dữ liệu hoặc nhóm dữ liệu dùng để phát hiện, phân loại hoặc chấm điểm gian lận."
    if kind == "java":
        if s.startswith("package "):
            return "Khai báo package của service trong backend Spring Boot."
        if s.startswith("import "):
            low = s.lower()
            if "repository" in low:
                return "Import repository để truy vấn hoặc lưu dữ liệu gian lận, attempt, signal, warning hoặc risk log."
            if "dto" in low:
                return "Import DTO dùng làm dữ liệu đầu ra hoặc đầu vào của service phân tích."
            if "entity" in low:
                return "Import entity domain như Attempt, FraudSignal, FraudWarning, RiskLevel để xử lý nghiệp vụ gian lận."
            return "Import class hoặc thư viện cần thiết cho service."
        if s.startswith("@Service"):
            return "Đánh dấu đây là service nghiệp vụ Spring, được inject vào controller hoặc service khác."
        if s.startswith("@RequiredArgsConstructor"):
            return "Lombok tự sinh constructor cho các dependency `final`, giúp dependency injection gọn."
        if s.startswith("@Transactional"):
            return "Đảm bảo thao tác đọc/ghi database trong phương thức diễn ra nhất quán trong một transaction."
        if s.startswith("@Value"):
            return "Đọc cấu hình từ application properties hoặc environment, thường là ngưỡng điểm, threshold hoặc cửa sổ dedup."
        if s.startswith("@"):
            return "Annotation cấu hình behavior runtime như transaction, async, lifecycle hoặc binding cấu hình."
        if "Repository" in s and ("private final" in s or "final" in s):
            return "Dependency repository để service đọc hoặc ghi dữ liệu phục vụ phát hiện gian lận."
        if re.match(r"public class|class ", s):
            return "Khai báo class service chính; mọi phương thức bên trong phục vụ một phần pipeline phát hiện gian lận."
        if re.match(r"(public|private|protected)\s+.*\)\s*\{?$", s) and "(" in s:
            return "Khai báo phương thức; đây là điểm xử lý hoặc helper trong service đang phân tích."
        if ".builder()" in s:
            return "Bắt đầu dựng object, DTO hoặc entity bằng builder để tạo response, warning, signal hoặc payload có cấu trúc."
        if re.match(r"\.[A-Za-z]", s):
            return "Một bước trong chuỗi builder hoặc stream, gán field hoặc biến đổi dữ liệu trước khi tạo object cuối."
        if s.endswith(".build();") or s == ".build();":
            return "Kết thúc builder và tạo object hoàn chỉnh để lưu hoặc trả về."
        if s.startswith("if ") or s.startswith("if("):
            return "Điều kiện rẽ nhánh để kiểm tra quyền, trạng thái attempt, threshold, signal hoặc dữ liệu bất thường."
        if s.startswith("else"):
            return "Nhánh thay thế khi điều kiện trước không thỏa, giúp phân loại tình huống khác."
        if s.startswith("for ") or s.startswith("for("):
            return "Vòng lặp xử lý nhiều attempt, signal, warning, đáp án hoặc feature."
        if ".stream()" in s or "Collectors" in s or ".map(" in s or ".filter(" in s:
            return "Dùng Java Stream để lọc, gom nhóm, tính toán hoặc chuyển đổi danh sách dữ liệu gian lận."
        if ".save(" in s:
            return "Lưu entity xuống database, thường là FraudSignal, FraudWarning, model metadata, evidence hoặc flag."
        if ".find" in s or "findBy" in s:
            return "Truy vấn database để lấy dữ liệu đầu vào cho phân tích gian lận."
        if "risk" in s.lower() or "score" in s.lower() or "threshold" in s.lower():
            return "Xử lý điểm hoặc ngưỡng rủi ro; đây là phần quyết định mức nghi vấn của attempt hoặc signal."
        if "similar" in s.lower() or "answer" in s.lower():
            return "Liên quan so sánh đáp án hoặc pattern trả lời giữa các attempt để phát hiện giống bất thường."
        if "fingerprint" in s.lower() or "device" in s.lower() or "ip" in s.lower():
            return "Liên quan định danh thiết bị hoặc IP, dùng để phát hiện đổi thiết bị, trùng thiết bị hoặc dấu hiệu tổ chức gian lận."
        if "return " in s:
            return "Trả kết quả xử lý cho caller; có thể là DTO phân tích, entity đã lưu hoặc giá trị helper."
        if "throw " in s:
            return "Ném lỗi khi dữ liệu không hợp lệ, không tìm thấy hoặc không đủ điều kiện xử lý."
        if "=" in s:
            return "Gán biến trung gian để lưu kết quả tính toán, dữ liệu truy vấn hoặc cấu hình xử lý."
        return f"Dòng xử lý trong ngữ cảnh `{ctx}`, góp phần vào pipeline phát hiện, phân tích hoặc ghi nhận gian lận."
    if kind == "python":
        if s.startswith("from __future__"):
            return "Bật compatibility cho type annotation tương lai của Python."
        if s.startswith("import ") or s.startswith("from "):
            low = s.lower()
            if "cv2" in low or "pil" in low or "numpy" in low:
                return "Import thư viện xử lý ảnh hoặc ma trận, phục vụ AI camera và OCR."
            if "fastapi" in low:
                return "Import FastAPI để khai báo endpoint AI service."
            if "openai" in low:
                return "Import client hoặc model provider cho các chức năng AI sinh, đánh giá hoặc chat."
            return "Import thư viện hoặc module nội bộ dùng trong AI service."
        if s.startswith("try:"):
            return "Bắt đầu khối thử tải dependency tùy chọn; nếu thiếu thư viện, service chuyển sang fallback."
        if s.startswith("except"):
            return "Xử lý lỗi import hoặc khởi tạo để AI service không crash khi thiếu dependency."
        if re.match(r"class\s+", s):
            return "Khai báo class gom trạng thái và hành vi của một engine AI hoặc service."
        if re.match(r"(async\s+)?def\s+", s):
            return "Khai báo hàm hoặc endpoint; đây là điểm xử lý request AI hoặc helper tính toán."
        if s.startswith("@"):
            return "Decorator gắn route FastAPI, dependency, lifecycle hook hoặc metadata cho hàm/lớp."
        low = s.lower()
        if "image_base64" in s or "base64" in s:
            return "Xử lý ảnh được frontend gửi dạng base64 để AI service decode và phân tích."
        if "cv2" in s or "dnn" in low or "Cascade" in s:
            return "Gọi OpenCV, DNN hoặc Haar cascade để detect khuôn mặt, mắt hoặc đặc trưng ảnh."
        if "mediapipe" in low or "facemesh" in low or "landmark" in low:
            return "Liên quan landmark hoặc FaceMesh để nhận diện mắt, iris hoặc hướng nhìn chính xác hơn."
        if "spoof" in low:
            return "Liên quan phát hiện giả mạo camera như màn hình replay, texture bất thường hoặc ảnh không tự nhiên."
        if "gaze" in low or "eye" in low:
            return "Liên quan tracking mắt hoặc hướng nhìn để phát hiện thí sinh nhìn ra ngoài màn hình."
        if "signal" in low:
            return "Tạo, lọc hoặc chuẩn hóa signal gian lận trả về backend."
        if "risk" in low or "warning" in low:
            return "Liên quan cảnh báo hoặc đánh giá rủi ro trong response AI."
        if s.startswith("if ") or s.startswith("elif "):
            return "Điều kiện phân nhánh theo trạng thái request, chất lượng frame, kết quả detect hoặc cấu hình."
        if s.startswith("else"):
            return "Nhánh mặc định khi các điều kiện trên không khớp."
        if s.startswith("for ") or s.startswith("while "):
            return "Vòng lặp qua ảnh, box, signal, token hoặc dữ liệu đầu vào để xử lý."
        if s.startswith("return "):
            return "Trả kết quả từ hàm; trong AI service thường là JSON response hoặc object phân tích."
        if "=" in s:
            return "Gán biến, cấu hình hoặc trạng thái trung gian dùng cho các bước AI tiếp theo."
        return f"Dòng xử lý trong ngữ cảnh `{ctx}`, phục vụ API AI, phân tích camera, OCR hoặc tác vụ AI phụ trợ."
    return "Dòng mã nguồn thuộc file đang phân tích."


def add_file(doc: Document, group: str, rel: str, role: str) -> int:
    path = ROOT / rel
    if not path.exists():
        para(doc, f"Không tìm thấy file: {rel}", "Heading 1")
        return 0
    lines = read_lines(path)
    kind = file_kind(path)
    para(doc, rel, "Heading 1")
    para(doc, f"Nhóm: {group}")
    para(doc, f"Vai trò: {role}")

    syms = []
    for i, line in enumerate(lines, 1):
        got = symbol(line, kind)
        if got:
            syms.append((i, got[0], got[1]))
    if syms:
        para(doc, "Bản đồ hàm/lớp đọc từ file", "Heading 2")
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        header(table, ["Dòng", "Loại", "Tên/khai báo"])
        for line_no, typ, name in syms[:90]:
            row = table.add_row().cells
            text_cell(row[0], line_no)
            text_cell(row[1], typ)
            text_cell(row[2], name)
        if len(syms) > 90:
            para(doc, f"File có {len(syms)} symbol; bảng hiển thị 90 symbol đầu để Word không quá nặng.")

    para(doc, "Phân tích từng dòng code", "Heading 2")
    ctx = "toàn file"
    chunk = 50
    for start in range(0, len(lines), chunk):
        end = min(start + chunk, len(lines))
        para(doc, f"Dòng {start + 1}-{end}", "Heading 3")
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        header(table, ["Dòng", "Code gốc", "Giải thích"])
        for idx in range(start, end):
            line = lines[idx]
            got = symbol(line, kind)
            if got:
                ctx = got[1]
            row = table.add_row().cells
            for cell in row:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            text_cell(row[0], idx + 1)
            code_cell(row[1], line)
            text_cell(row[2], explain(line, kind, ctx))
    return len(lines)


def build_doc() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    for style_name in ("Normal", "Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[style_name]
        style.font.name = FONT
        style.font.size = FONT_SIZE
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style._element.rPr.rFonts.set(qn("w:cs"), FONT)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("PHÂN TÍCH SERVICE PHÁT HIỆN GIAN LẬN VÀ AI-SERVICE")
    run.bold = True
    apply_font_to_run(run)

    para(doc, f"Dự án: {ROOT}")
    para(doc, f"Ngày tạo: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    para(
        doc,
        "Tài liệu này được tạo lại theo phạm vi mới: chỉ đọc và phân tích các service backend phát hiện gian lận cùng các file trong ai-service/app.",
    )

    para(doc, "Phạm vi file", "Heading 1")
    summary = doc.add_table(rows=1, cols=4)
    summary.style = "Table Grid"
    header(summary, ["Nhóm", "File", "Số dòng", "Vai trò"])
    counts = []
    for group, rel, role in FILES:
        path = ROOT / rel
        count = len(read_lines(path)) if path.exists() else 0
        counts.append((group, rel, role, count))
        row = summary.add_row().cells
        text_cell(row[0], group)
        text_cell(row[1], rel)
        text_cell(row[2], count)
        text_cell(row[3], role)

    para(doc, "Cách đọc nhanh", "Heading 1")
    for line in [
        "Backend service phát hiện gian lận bắt đầu từ FraudSignalService: event hoặc signal hệ thống được chuẩn hóa thành FraudSignal.",
        "RiskScoringService là nơi giải thích điểm rủi ro: signal score, category, dedup window, cap điểm, RiskLevel và ProctorFlag.",
        "FraudWarningService và FraudAnalysisService là lớp cảnh báo và hậu kiểm: gom warning, review, phân tích similarity, timing, IP, statistics và behavior.",
        "MLRiskScoringService và FraudModelTrainer là phần mở rộng ML: trích feature, anomaly score, train và evaluate model.",
        "ai-service/app/proctor.py là lõi camera AI: decode ảnh, detect mặt/mắt, gaze, spoofing, quality và trả signal cho backend.",
    ]:
        para(doc, line)

    for group, rel, role, _ in counts:
        add_file(doc, group, rel, role)

    para(doc, "Tóm tắt bảo vệ", "Heading 1")
    for line in [
        "Luồng gian lận về mặt service: AI service sinh signal, backend FraudSignalService lưu signal, RiskScoringService tính điểm, FraudWarningService và FraudAnalysisService tổng hợp cảnh báo/hậu kiểm.",
        "Điểm mạnh của thiết kế là tách rõ signal thô, risk score, warning và ML analysis; nhờ vậy có thể giải thích từng quyết định thay vì chỉ đưa ra kết luận mơ hồ.",
        "Các bằng chứng camera được lưu qua ProctorEvidenceImageService, còn DeviceFingerprintService và AnswerSimilarityService bổ sung hai nhóm phát hiện quan trọng: định danh thiết bị và giống đáp án.",
    ]:
        para(doc, line)

    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(OUT)
