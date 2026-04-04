#!/usr/bin/env python3
"""
extract_docx.py — Extract text from DOCX files using python-docx.
Preserves paragraph structure and handles tables.

Usage:
    python extract_docx.py <input.docx> [--output <output.txt>]

    If --output is omitted, prints to stdout.
Exit codes:
    0  — success
    1  — file not found or not a DOCX
    2  — python-docx not installed
"""

import sys
import os
import argparse


def extract_text(filepath: str) -> str:
    """Extract all text from DOCX, preserving paragraph order."""
    import docx
    doc = docx.Document(filepath)
    parts = []

    for element in doc.element.body:
        tag = element.tag.split('}')[-1]  # strip namespace

        if tag == 'p':
            paragraph = next(
                (p for p in doc.paragraphs if p._element is element),
                None
            )
            if paragraph is not None:
                # Join runs to preserve inline content
                run_text = ''.join(run.text for run in paragraph.runs)
                text = run_text.strip()
                if text:
                    parts.append(text)
                else:
                    parts.append('')

        elif tag == 'tbl':
            table = next(
                (t for t in doc.tables if t._element is element),
                None
            )
            if table is not None:
                for row in table.rows:
                    row_cells = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_cells.append(cell_text)
                    if row_cells:
                        parts.append(' | '.join(row_cells))
                parts.append('')

    # Filter trailing empty lines and join
    while parts and parts[-1] == '':
        parts.pop()

    return '\n'.join(parts)


def main():
    parser = argparse.ArgumentParser(description="Extract text from DOCX using python-docx")
    parser.add_argument("input", help="Path to input DOCX file")
    parser.add_argument("--output", "-o", help="Output file path (default: stdout)")
    args = parser.parse_args()

    filepath = os.path.abspath(args.input)
    if not os.path.isfile(filepath):
        print(f"ERROR: File not found: {filepath}", file=sys.stderr)
        sys.exit(1)

    try:
        text = extract_text(filepath)
        text = text.strip()
        if not text:
            print("WARNING: No text extracted from DOCX.", file=sys.stderr)

        # Write to stdout with UTF-8 encoding (handles Windows charmap issues)
        if args.output:
            with open(args.output, "w", encoding="utf-8") as f:
                f.write(text)
        else:
            try:
                sys.stdout.reconfigure(encoding='utf-8')
            except Exception:
                pass
            # Encode to avoid charmap errors on Windows console
            try:
                sys.stdout.buffer.write(text.encode('utf-8', errors='replace'))
            except (AttributeError, TypeError):
                print(text)

        sys.exit(0)
    except ImportError as e:
        print("ERROR: python-docx not installed. Install with: pip install python-docx", file=sys.stderr)
        sys.exit(2)
    except Exception as e:
        print(f"ERROR: Failed to extract DOCX: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
